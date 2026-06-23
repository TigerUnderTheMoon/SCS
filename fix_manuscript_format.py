"""Fix manuscript format: line numbering, Statements heading, line spacing, EDGAR year, abstract trim."""
import docx
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MANUSCRIPT = r'D:\Workplace\SCS\outputs\espr_submission\manuscript_espr.docx'

doc = docx.Document(MANUSCRIPT)

# 1. Enable continuous line numbering in section properties
sectPr = doc.sections[0]._sectPr
existing = sectPr.find(qn('w:lnNumType'))
if existing is not None:
    sectPr.remove(existing)
lnNumType = parse_xml(
    '<w:lnNumType %s w:countBy="1" w:restart="continuous" w:distance="360"/>'
    % nsdecls("w")
)
sectPr.append(lnNumType)
print('1. Line numbering enabled (continuous, countBy=1)')

# 2. Change "Statements and Declarations" from Normal to Heading 1
for p in doc.paragraphs:
    if p.text.strip() == 'Statements and Declarations' and p.style and p.style.name == 'Normal':
        p.style = doc.styles['Heading 1']
        print('2. "Statements and Declarations" set to Heading 1')
        break

# 3. Set 1.5x line spacing for Normal and List Bullet paragraphs
count_spaced = 0
for p in doc.paragraphs:
    if p.style and p.style.name in ('Normal', 'List Bullet'):
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        count_spaced += 1
print('3. Set 1.5x line spacing for %d paragraphs' % count_spaced)

# 4. Fix EDGAR year 2026 -> 2025
for p in doc.paragraphs:
    if 'European Commission JointResearchCentre (2026)' in p.text:
        for run in p.runs:
            if '2026' in run.text:
                run.text = run.text.replace('2026', '2025')
        print('4. EDGAR year fixed 2026 -> 2025')
        break

# 5. Trim abstract from 228 to <=220 words
# Find abstract paragraphs (between Abstract heading and Keywords)
abstract_start = None
abstract_end = None
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == 'Heading 1' and p.text.strip() == 'Abstract':
        abstract_start = i + 1
    elif abstract_start and p.text.strip().startswith('Keywords:'):
        abstract_end = i
        break

# Get abstract text - it's in one paragraph (paragraph index abstract_start)
abstract_para = doc.paragraphs[abstract_start]
abstract_text = abstract_para.text.strip()
words = abstract_text.split()
print('5. Abstract original: %d words' % len(words))

if len(words) > 220:
    # Strategy: remove redundant phrases to get under 220
    # Original phrases to remove/shorten:
    # "which lies inside the observed SCCD range" -> remove (7 words) -> 221
    # "before operating efficiency emerges" -> "before efficiency gains materialize" (same count)
    # Let's remove "which lies inside the observed SCCD range" and slightly trim elsewhere

    new_abstract = abstract_text
    # Remove "which lies inside the observed SCCD range" (but keep sentence coherent)
    new_abstract = new_abstract.replace(
        ', and the turning point is 0.522, which lies inside the observed SCCD range.',
        ', with a turning point of 0.522.'
    )
    # Remove "Supplementary instrumental-variable estimates imply a turning point of 0.570 and are interpreted cautiously as diagnostic rather than decisive evidence."
    # -> shorten to "Supplementary instrumental-variable estimates imply a turning point of 0.570."
    new_abstract = new_abstract.replace(
        'Supplementary instrumental-variable estimates imply a turning point of 0.570 and are interpreted cautiously as diagnostic rather than decisive evidence.',
        'Supplementary instrumental-variable estimates imply a turning point of 0.570.'
    )

    new_words = new_abstract.split()
    print('   After trim: %d words' % len(new_words))

    if len(new_words) <= 220:
        # Update the paragraph
        for run in abstract_para.runs:
            run.text = ''
        abstract_para.runs[0].text = new_abstract
        print('   Abstract trimmed and updated')
    else:
        # Need more trimming
        # Remove "covering 4,544 city-year observations" -> "4,544 city-year observations"
        new_abstract2 = new_abstract.replace(
            'Using an observed panel of 284 Chinese prefecture-level and above cities from 2006 to 2021, covering 4,544 city-year observations, this study constructs',
            'Using an observed panel of 284 Chinese cities from 2006 to 2021 (4,544 city-year observations), this study constructs'
        )
        new_words2 = new_abstract2.split()
        print('   After second trim: %d words' % len(new_words2))
        if len(new_words2) <= 220:
            for run in abstract_para.runs:
                run.text = ''
            abstract_para.runs[0].text = new_abstract2
            print('   Abstract trimmed (second pass) and updated')
        else:
            print('   WARNING: still over 220, manual trim needed')
            print('   Current: %d words' % len(new_words2))

doc.save(MANUSCRIPT)
print('\nManuscript saved.')

# Verify all fixes
print('\n=== VERIFICATION ===')
doc2 = docx.Document(MANUSCRIPT)

# Verify line numbering
sectPr2 = doc2.sections[0]._sectPr
ln = sectPr2.find(qn('w:lnNumType'))
print('Line numbering: %s' % ('ENABLED' if ln is not None else 'NOT SET'))

# Verify Statements heading
for p in doc2.paragraphs:
    if p.text.strip() == 'Statements and Declarations':
        print('Statements heading style: %s' % p.style.name)
        break

# Verify EDGAR
for p in doc2.paragraphs:
    if 'European Commission JointResearchCentre' in p.text:
        if '(2025)' in p.text:
            print('EDGAR year: 2025 OK')
        elif '(2026)' in p.text:
            print('EDGAR year: STILL 2026 - FAILED')
        break

# Verify abstract word count
for i, p in enumerate(doc2.paragraphs):
    if p.style and p.style.name == 'Heading 1' and p.text.strip() == 'Abstract':
        abstract = doc2.paragraphs[i+1].text.strip()
        wc = len(abstract.split())
        print('Abstract word count: %d %s' % (wc, 'OK' if wc <= 220 else 'STILL OVER'))
        break
