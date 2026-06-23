"""Complete 33 references via Crossref API - patch page ranges and issue numbers."""
import docx
import re
import time
import json
import urllib.request
import urllib.error
import os

MANUSCRIPT = r'D:\Workplace\SCS\outputs\espr_submission\manuscript_espr.docx'
BACKUP = r'D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_pre_ref_patch.docx'

# Backup first
import shutil
if not os.path.exists(BACKUP):
    shutil.copy2(MANUSCRIPT, BACKUP)
    print('Backup created:', BACKUP)

doc = docx.Document(MANUSCRIPT)

# Find references section
ref_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'References' and p.style and 'Heading' in p.style.name:
        ref_idx = i
        break

# Extract references and DOIs
refs = []
ref_para_indices = []
for i in range(ref_idx + 1, len(doc.paragraphs)):
    t = doc.paragraphs[i].text.strip()
    if t:
        refs.append(t)
        ref_para_indices.append(i)

print(f'Found {len(refs)} references')

# Extract DOI from each reference
def extract_doi(ref_text):
    m = re.search(r'https://doi\.org/(10\.[^\s]+)', ref_text)
    if m:
        return m.group(1).rstrip('.')
    return None

def query_crossref(doi):
    """Query Crossref API for a given DOI."""
    url = f'https://api.crossref.org/works/{doi}'
    req = urllib.request.Request(url, headers={
        'User-Agent': 'SCS-Manuscript/1.0 (mailto:wangningning@bistu.edu.cn)'
    })
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            return data.get('message', {})
    except Exception as e:
        print(f'  Crossref error for {doi}: {e}')
        return None

def extract_metadata(meta):
    """Extract volume, issue, page from Crossref metadata."""
    volume = meta.get('volume', '')
    issue = meta.get('issue', '')
    page = meta.get('page', '')
    # Some Crossref records use 'article-number' instead of page
    if not page:
        page = meta.get('article-number', '')
    return volume, issue, page

# Process each reference
results = []
for idx, (ref_text, para_idx) in enumerate(zip(refs, ref_para_indices)):
    doi = extract_doi(ref_text)
    print(f'[{idx+1:02d}/33] {ref_text[:80]}...')
    print(f'  DOI: {doi}')

    if not doi:
        # EDGAR reference has no DOI - skip
        print(f'  No DOI - skipping')
        results.append({'idx': idx, 'doi': None, 'status': 'no_doi', 'ref': ref_text})
        continue

    meta = query_crossref(doi)
    if meta is None:
        print(f'  Crossref failed - keeping original')
        results.append({'idx': idx, 'doi': doi, 'status': 'api_failed', 'ref': ref_text})
        time.sleep(1)
        continue

    volume, issue, page = extract_metadata(meta)
    print(f'  vol={volume}, issue={issue}, page={page}')

    # Reconstruct the reference
    # Current format: "Authors (Year) Title. Journal Volume. https://doi.org/..."
    # Target format:  "Authors (Year) Title. Journal Volume(Issue): Pages. https://doi.org/..."

    # Split on the DOI to preserve it
    doi_match = re.search(r'(https://doi\.org/10\.[^\s]+)', ref_text)
    if not doi_match:
        results.append({'idx': idx, 'doi': doi, 'status': 'no_doi_in_text', 'ref': ref_text})
        continue

    doi_url = doi_match.group(1).rstrip('.')
    parts = ref_text.split(doi_url)
    before_doi = parts[0].rstrip().rstrip('.').rstrip()
    after_doi = parts[1].strip() if len(parts) > 1 else ''

    # before_doi is like: "An X, Yang Y, Zhang X, Zeng X (2024) Smarter and Cleaner? ... Sustainability 16"
    # We need to append (issue): page before the DOI
    suffix = ''
    if issue and page:
        suffix = f'({issue}): {page}'
    elif issue and not page:
        suffix = f'({issue})'
    elif page and not issue:
        suffix = f': {page}'

    new_ref = f'{before_doi}{suffix}. {doi_url}'
    if after_doi:
        new_ref += f' {after_doi}'

    # Update the paragraph
    para = doc.paragraphs[para_idx]
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_ref
    else:
        para.add_run(new_ref)

    results.append({
        'idx': idx, 'doi': doi, 'status': 'patched',
        'volume': volume, 'issue': issue, 'page': page,
        'new_ref': new_ref[:150]
    })
    print(f'  PATCHED: {new_ref[:100]}...')

    time.sleep(1)  # Rate limit

# Save
doc.save(MANUSCRIPT)
print(f'\nManuscript saved with patched references.')

# Summary
patched = sum(1 for r in results if r['status'] == 'patched')
failed = sum(1 for r in results if r['status'] == 'api_failed')
no_doi = sum(1 for r in results if r['status'] == 'no_doi')
print(f'\n=== SUMMARY ===')
print(f'Patched: {patched}/33')
print(f'API failed: {failed}/33')
print(f'No DOI: {no_doi}/33')

# Verify
print(f'\n=== VERIFICATION ===')
doc2 = docx.Document(MANUSCRIPT)
ref_idx2 = None
for i, p in enumerate(doc2.paragraphs):
    if p.text.strip() == 'References' and p.style and 'Heading' in p.style.name:
        ref_idx2 = i
        break
has_pages = 0
has_issue = 0
for i in range(ref_idx2 + 1, len(doc2.paragraphs)):
    t = doc2.paragraphs[i].text.strip()
    if t:
        if re.search(r'\d+\s*[-–]\s*\d+', t) or re.search(r':\s*\d+', t):
            has_pages += 1
        if re.search(r'\d+\(\d+\)', t):
            has_issue += 1
print(f'Refs with pages: {has_pages}/33')
print(f'Refs with issue: {has_issue}/33')

# Save results JSON
with open(r'D:\Workplace\SCS\outputs\espr_submission\ref_patch_results.json', 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, ensure_ascii=False)
print(f'\nDetailed results saved to ref_patch_results.json')
