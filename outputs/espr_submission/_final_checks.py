"""Final checks: association language and EC in-text citation."""
from pathlib import Path
from docx import Document

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_final.docx")
doc = Document(str(SRC))
full = "\n".join(p.text for p in doc.paragraphs)

print("=== ASSOCIATION LANGUAGE CHECK ===")
import re
# Count "association" mentions
assoc_count = len(re.findall(r"\bassociation\b", full, re.IGNORECASE))
print(f"  'association' mentions: {assoc_count}")

# Check no "causal effect" was introduced
causal_effect = re.findall(r"causal effect", full, re.IGNORECASE)
print(f"  'causal effect' mentions: {len(causal_effect)}")
if causal_effect:
    for m in causal_effect:
        print(f"    -> {m}")

# Check "associated pathway" preserved
ap_count = len(re.findall(r"associated (pathway|restructuring)", full, re.IGNORECASE))
print(f"  'associated pathway/restructuring' mentions: {ap_count}")

print()
print("=== EC IN-TEXT CITATION CHECK ===")
for i, p in enumerate(doc.paragraphs):
    if "European Commission" in p.text and "2026" in p.text:
        print(f"  [{i:03d}] {p.text[:200]}")

print()
print("=== SECTION 5.2 HEADING CHECK (should still say 'for Sustainability' - concept) ===")
for i, p in enumerate(doc.paragraphs):
    if "5.2." in p.text and p.style.name == "Heading 2":
        print(f"  [{i:03d}] {p.text}")

print()
print("=== TASK 1 REPLACEMENT CONTEXT CHECK ===")
contexts = [
    ("the main evidence", "excludes those rows"),
    ("keeps the paper's message", "data-enabled"),
    ("clear to reviewers who", "data transparency"),
    ("For readers, this matters", "endpoint"),
    ("aligns with the applied policy audience", "staged governance"),
]
for find, ctx in contexts:
    if find in full:
        idx = full.find(find)
        s = max(0, idx - 40)
        e = min(len(full), idx + len(find) + 60)
        print(f"  OK: ...{full[s:e]}...")
    else:
        print(f"  MISSING: '{find}'")
