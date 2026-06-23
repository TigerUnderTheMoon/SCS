"""
Apply cover letter revisions:
1. Add Suggested Reviewers section (4 reviewers) with placeholders;
   each entry specifies field direction, with name/affiliation/email marked [MANUAL - corresponding author to
   confirm] and a one-sentence reason.
2. Verify no green-economy / low-carbon-economy framing (no change needed if already absent).
   If found, replace with urban environmental governance / carbon emission management / sustainable urban transition.
3. Save as cover_letter_espr_revised.docx
"""

import json
from pathlib import Path
from docx import Document

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\cover_letter_espr.docx")
DST = Path(r"D:\Workplace\SCS\outputs\espr_submission\cover_letter_espr_revised.docx")
LOG = Path(r"D:\Workplace\SCS\outputs\espr_submission\_revisions_log_cover.json")

doc = Document(str(SRC))
cover_log = []

def log(loc, before, after, reason, manual=False):
    cover_log.append({"location": loc, "before_summary": before, "after_summary": after, "reason": reason, "manual": manual})

def set_para_text(p, text):
    runs = list(p.runs)
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

def find_para_signing_off():
    """Locate the 'Sincerely,' paragraph; reviewers section will be inserted before final signoff block."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() in ("sincerely,", "sincerely"):
            return i
    return -1

# Build the Suggested Reviewers section text (paragraphs).
# Insert immediately before "Sincerely," or before "We thank you for considering…"

SR_TITLE = "Suggested Reviewers"
SR_INTRO = (
"We suggest the following experts as potential reviewers. Each recommendation targets a distinct field that aligns with the "
"manuscript's scope (urban environmental policy & carbon emissions, urban spatial structure & econometrics, China environmental "
"policy & sustainable urban development, and digital economy & environmental performance). Contact details will be confirmed by "
"the corresponding author before final submission to ensure accurate information."
)
SR_R1 = (
"Reviewer 1 (Field: urban environmental policy & carbon emissions). Name: [To be confirmed by corresponding author]. Affiliation: "
"[To be confirmed - preference for a research institution with urban environmental policy and carbon-emission accounting focus]. "
"Email: [To be confirmed]. Reason: Recommended for expertise in smart-city policy evaluation and urban carbon-emission accounting, "
"which aligns with the manuscript's central question on whether smart-city coordination reduces emissions."
)
SR_R2 = (
"Reviewer 2 (Field: urban spatial structure & econometrics). Name: [To be confirmed]. Affiliation: [To be confirmed - preference "
"for an urban-economics or spatial-econometrics research group]. Email: [To be confirmed]. Reason: Recommended because the "
"manuscript's two-way fixed-effects panel with a quadratic SCCD specification, supplementary IV-2SLS evidence, and spatial "
"coordination index construction require expertise in urban spatial structure econometrics."
)
SR_R3 = (
"Reviewer 3 (Field: China environmental policy & sustainable urban development). Name: [To be confirmed]. Affiliation: [To be "
"confirmed - preference for a Chinese university or research institute specializing in sustainable urban transition]. Email: [To be "
"confirmed]. Reason: Recommended for domain knowledge of China's prefecture-level city data, regional heterogeneity in "
"eastern/central/western development stage contexts, and governance sequencing relevant to the manuscript's policy implications."
)
SR_R4 = (
"Reviewer 4 (Field: digital economy & environmental performance). Name: [To be confirmed]. Affiliation: [To be confirmed - preference "
"for a digital economy and emissions research group]. Email: [To be confirmed]. Reason: Recommended for expertise on the digital "
"economy's moderating role in urban carbon emissions, directly relevant to the manuscript's DEI-moderation finding and four-"
"functional-space SCCD framework."
)

# Find an insertion anchor: insert before the "We thank you for considering our manuscript" paragraph
anchor_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("We thank you for considering our manuscript"):
        anchor_idx = i
        break
if anchor_idx == -1:
    # fallback: insert before "Sincerely,"
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() in ("sincerely,",):
            anchor_idx = i
            break
assert anchor_idx > 0, "could not locate insertion anchor in cover letter"

anchor = doc.paragraphs[anchor_idx]

# Insert paragraphs in this order BEFORE the anchor
def add_before(anchor_p, text, style=None):
    return anchor_p.insert_paragraph_before(text, style=style)

# Add an empty separator, title, intro, then 4 reviewer blocks, then empty separator
add_before(anchor, "")
add_before(anchor, SR_TITLE)
add_before(anchor, "")
add_before(anchor, SR_INTRO)
add_before(anchor, "")
add_before(anchor, SR_R1)
add_before(anchor, "")
add_before(anchor, SR_R2)
add_before(anchor, "")
add_before(anchor, SR_R3)
add_before(anchor, "")
add_before(anchor, SR_R4)
add_before(anchor, "")

log(
    "Cover Letter - Suggested Reviewers section",
    "Cover letter had no Suggested Reviewers section (author_metadata_finalized.md notes Suggested reviewers: (none))",
    "Added Suggested Reviewers section with 4 entries covering urban environmental policy & carbon emissions, urban spatial structure & econometrics, China environmental policy & sustainable urban development, and digital economy & environmental performance. Each entry has a one-sentence reason; name/affiliation/email fields are explicitly marked as [To be confirmed] pending corresponding-author confirmation.",
    "Per task: supplement cover letter with 3-5 suggested reviewers matching the three requested reviewer-direction categories. To avoid fabricating real-person contact information without verification, name/affiliation/email are templated as placeholders flagged for manual confirmation.",
    manual=True,
)

# Validate no green/low-carbon economy phrasing (final-checks already shows green_economy=0)
REPLACE_TERMS = {
    "green economy": "urban environmental governance",
    "low-carbon economy": "carbon emission management",
    "low carbon economy": "carbon emission management",
}

full_before = "\n".join(p.text for p in doc.paragraphs)
for term, repl in REPLACE_TERMS.items():
    for p in doc.paragraphs:
        if term.lower() in p.text.lower():
            new_text = p.text.replace(term, repl).replace(term.capitalize(), repl)
            set_para_text(p, new_text)
            log(
                "Cover Letter body - framing term",
                f"Contained '{term}'",
                f"Replaced with '{repl}'",
                "Per task: avoid framework-level green/low-carbon economy wording; use urban environmental governance / carbon emission management / sustainable urban transition.",
            )

full_after = "\n".join(p.text for p in doc.paragraphs)
assert "green economy" not in full_after.lower() and "low-carbon economy" not in full_after.lower() and "low carbon economy" not in full_after.lower(), "framing term still present"

doc.save(str(DST))
LOG.write_text(json.dumps(cover_log, ensure_ascii=False, indent=2), encoding="utf-8")
print("OK", DST)
print("Paragraphs:", len(doc.paragraphs))