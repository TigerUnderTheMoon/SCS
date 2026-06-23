"""Verify cover letter."""
from docx import Document
doc = Document(r'D:\Workplace\SCS\outputs\espr_submission\cover_letter_espr_revised.docx')
print(f"paragraphs: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs):
    print(f"[{i:03d}] <{p.style.name}> {p.text[:160]}")