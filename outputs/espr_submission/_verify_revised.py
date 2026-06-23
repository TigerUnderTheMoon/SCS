"""Verify revised manuscript."""
from docx import Document
import json
doc = Document(r'D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_revised.docx')

# Count words per section
print("paragraphs:", len(doc.paragraphs))
print("tables:", len(doc.tables))

# Show all Heading 2 paragraphs
print("\n=== Heading 2 sections ===")
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name.startswith("Heading 2"):
        print(f"  [{i}] {p.text}")

# Show abstract body
i_abs = -1
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == "Heading 1" and p.text.strip().lower() == "abstract":
        i_abs = i
        break
abs_text = doc.paragraphs[i_abs+1].text
print(f"\n=== Abstract ({len(abs_text.split())} words) ===")
print(abs_text)

# Check no remaining numbered citations
import re
print("\n=== Remaining numbered citations [N] ===")
for i, p in enumerate(doc.paragraphs):
    if re.search(r'\[\d', p.text):
        print(f"  [{i}] ...{p.text[max(0,p.text.find('['))-20:p.text.find('[')+50]}...")

# Section 2 word count
def section_word_count(doc):
    total = 0
    in_sec2 = False
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1":
            if p.text.startswith("2.") or p.text == "2. Literature Review and Theoretical Framework":
                in_sec2 = True
                continue
            elif in_sec2:
                break
        if p.style and p.style.name == "Heading 2" and in_sec2:
            # we're entering a 2.x
            continue
        if in_sec2:
            total += len(p.text.split())
    return total

# Simpler: just count words between Section 2 heading and Section 3 heading
def count_section2():
    in_sec = False
    total = 0
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1":
            if "Literature Review" in p.text:
                in_sec = True
                continue
            if in_sec and p.text.startswith("3."):
                break
        if in_sec:
            total += len(p.text.split())
    return total

print(f"\n=== Section 2 word count: {count_section2()} ===")

# Dump key verification target paragraphs
import re
keys = [
    "Elevated VIF values for SCCD",
    "Electronic Supplementary Material (ESM_1",
    "agreement because",
    "treatment concept",
]
print("\n=== Key text presence ===")
full_text = "\n".join(p.text for p in doc.paragraphs)
for k in keys:
    print(f"  '{k}': {'FOUND' if k in full_text else 'MISSING'}")

# Show all references starting positions
print("\n=== References block check ===")
in_refs = False
ref_count = 0
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == "Heading 1" and "References" in p.text:
        in_refs = True
        continue
    if in_refs and p.text.strip() and "Statements" not in p.text:
        # Reset on next H1
        if p.style and p.style.name == "Heading 1":
            in_refs = False
            break
        ref_count += 1
print(f"  Reference entries (paragraphs): {ref_count}")

# Check ESM mention
for k in ["ESM_1", "ESM_2", "ESM_3"]:
    print(f"  {k}: {'FOUND' if k in full_text else 'MISSING'}")
# Tables/Figures
for k in ["Table 1", "Table 2", "Table 3", "Table 4", "Table 5", "Table 6", "Table 7",
          "Figure 1", "Figure 2", "Figure 3"]:
    print(f"  {k}: {'FOUND' if k in full_text else 'MISSING'}")