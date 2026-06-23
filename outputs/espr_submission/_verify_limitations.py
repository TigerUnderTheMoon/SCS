"""Verify Limitations 5.4 section is fully intact."""
from pathlib import Path
from docx import Document

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_final.docx")
doc = Document(str(SRC))

# Compare 5.4 between source and final
SRC_ORIG = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_revised.docx")
doc_orig = Document(str(SRC_ORIG))

def get_section_54(d):
    paras = []
    in_54 = False
    for p in d.paragraphs:
        if p.style.name == "Heading 2" and "5.4." in p.text:
            in_54 = True
            paras.append(p.text)
            continue
        if in_54:
            if p.style.name.startswith("Heading 1"):
                break
            paras.append(p.text)
    return paras

orig_54 = get_section_54(doc_orig)
final_54 = get_section_54(doc)

print("=== ORIGINAL 5.4 Limitations ===")
for i, t in enumerate(orig_54):
    print(f"  [{i}] ({len(t.split())}w) {t[:120]}...")
print(f"  Total paragraphs: {len(orig_54)}")
print(f"  Total words: {sum(len(t.split()) for t in orig_54)}")

print()
print("=== FINAL 5.4 Limitations ===")
for i, t in enumerate(final_54):
    print(f"  [{i}] ({len(t.split())}w) {t[:120]}...")
print(f"  Total paragraphs: {len(final_54)}")
print(f"  Total words: {sum(len(t.split()) for t in final_54)}")

print()
if orig_54 == final_54:
    print("MATCH: Limitations 5.4 is IDENTICAL between source and final.")
else:
    print("MISMATCH: Limitations 5.4 differs!")
    for i in range(max(len(orig_54), len(final_54))):
        o = orig_54[i] if i < len(orig_54) else "<missing>"
        f = final_54[i] if i < len(final_54) else "<missing>"
        if o != f:
            print(f"  DIFF at [{i}]:")
            print(f"    ORIG: {o[:200]}")
            print(f"    FINAL: {f[:200]}")
