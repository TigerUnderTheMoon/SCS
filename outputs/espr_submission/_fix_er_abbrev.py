"""Fix ER abbreviation first-mention definition in manuscript_espr_final.docx."""
from pathlib import Path
from docx import Document

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_final.docx")
doc = Document(str(SRC))

OLD = "Moderation models examine DEI, ER, and POLY."
NEW = "Moderation models examine DEI, environmental regulation (ER), and POLY."

fixed = False
for p in doc.paragraphs:
    if OLD in p.text:
        # Single-run paragraph; replace text preserving formatting
        runs = list(p.runs)
        if runs:
            runs[0].text = p.text.replace(OLD, NEW)
            for r in runs[1:]:
                r.text = ""
        fixed = True
        print(f"FIXED: '{OLD}'")
        print(f"     -> '{NEW}'")
        break

if not fixed:
    print("ERROR: Target text not found!")
else:
    doc.save(str(SRC))
    print(f"\nSaved: {SRC}")

    # Verify
    doc2 = Document(str(SRC))
    for p in doc2.paragraphs:
        if "environmental regulation (ER)" in p.text:
            print(f"\nVERIFIED: definition present")
            print(f"  Context: ...{p.text[:120]}...")
            break
