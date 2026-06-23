"""Dump cover_letter_espr.docx."""
from docx import Document
doc = Document(r'D:\Workplace\SCS\outputs\espr_submission\cover_letter_espr.docx')
with open(r'D:\Workplace\SCS\outputs\espr_submission\_inspect_cover.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        f.write(f'[{i:03d}] <{p.style.name}> {p.text}\n')
    f.write('\n=== TABLES ===\n')
    for ti, t in enumerate(doc.tables):
        f.write(f'\n--- Table {ti} ({len(t.rows)}x{len(t.columns)}) ---\n')
        for ri, row in enumerate(t.rows):
            f.write(f'  R{ri}: ' + ' || '.join(c.text.replace("\n"," | ") for c in row.cells) + '\n')
print('done')