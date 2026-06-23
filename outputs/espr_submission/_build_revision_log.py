"""Apply minor heading update and assemble revision_log.md."""
import json
from pathlib import Path
from docx import Document
from datetime import date

REV = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_revised.docx")
doc = Document(str(REV))
# Find the heading "2.3. Digitalization, Green Technology Progress, and Low-Carbon Coupling"
target = "2.3. Digitalization, Green Technology Progress, and Low-Carbon Coupling"
new_heading = "2.3. Digitalization, Green Technology Progress, and the Coordination Paradox"
changed = False
for p in doc.paragraphs:
    if p.text.strip() == target:
        runs = list(p.runs)
        if runs:
            runs[0].text = new_heading
            for r in runs[1:]:
                r.text = ""
        else:
            p.add_run(new_heading)
        changed = True
        break
doc.save(str(REV))
print("heading changed:", changed)

# Now build revision_log.md
manuscript_log_path = Path(r"D:\Workplace\SCS\outputs\espr_submission\_revisions_log_data.json")
cover_log_path = Path(r"D:\Workplace\SCS\outputs\espr_submission\_revisions_log_cover.json")

ml = json.loads(manuscript_log_path.read_text(encoding="utf-8"))
cl = json.loads(cover_log_path.read_text(encoding="utf-8"))

# Add heading-update entry
ml.append({
    "location": "Section 2.3 heading",
    "before_summary": "Heading: '2.3. Digitalization, Green Technology Progress, and Low-Carbon Coupling'",
    "after_summary": "Heading: '2.3. Digitalization, Green Technology Progress, and the Coordination Paradox'",
    "reason": "Reflect the merged content (original 2.3 + 2.4 Coordination Paradox) in the heading title so the section scope matches the new compressed content.",
    "manual": False,
})

md_lines = []
md_lines.append("# ESPR Submission Revision Log")
md_lines.append("")
md_lines.append(f"- Manuscript: `outputs/espr_submission/manuscript_espr.docx` -> `outputs/espr_submission/manuscript_espr_revised.docx`")
md_lines.append(f"- Cover letter: `outputs/espr_submission/cover_letter_espr.docx` -> `outputs/espr_submission/cover_letter_espr_revised.docx`")
md_lines.append(f"- Target journal: Environmental Science and Pollution Research (ESPR), Springer Nature")
md_lines.append(f"- Date: {date.today().isoformat()}")
md_lines.append("")
md_lines.append("## Context and scientific-governance red lines")
md_lines.append("")
md_lines.append("All edits preserve the empirical record: no coefficient, standard error, significance star, turning point (0.522 / 0.570), IV diagnostic (F = 61.55 / 80.44 / 380.73), observation count, R-squared, sample period (2006-2021 observed vs 2022-2024 fitted) was modified. The wording discipline of the original manuscript (\"association\" rather than \"causal effect\"; \"supplementary IV evidence\" rather than \"primary identification strategy\"; \"associated pathway\" rather than \"formal decomposition\") was retained. The Limitations section (5.4) wording on observational design, time-varying unobserved factors, and IV limitations remained unchanged. No references were added or removed; the 33-entry reference list remained alphabetically ordered by first author's last name per ESPR guidelines.")
md_lines.append("")
md_lines.append("## A. Manuscript revisions")
md_lines.append("")
md_lines.append("| # | Location | Before (summary) | After (summary) | Reason |")
md_lines.append("|---|----------|------------------|------------------|--------|")
for i, e in enumerate(ml, 1):
    loc = e["location"].replace("|", "/")
    before = (e["before_summary"][:220] + "..." if len(e["before_summary"]) > 220 else e["before_summary"]).replace("|", "/")
    after = (e["after_summary"][:220] + "..." if len(e["after_summary"]) > 220 else e["after_summary"]).replace("|", "/")
    reason = (e["reason"][:300] + "..." if len(e["reason"]) > 300 else e["reason"]).replace("|", "/")
    flag = " **[MANUAL]**" if e.get("manual") else ""
    md_lines.append(f"| {i} | {loc}{flag} | {before} | {after} | {reason} |")

md_lines.append("")
md_lines.append("## B. Cover letter revisions")
md_lines.append("")
md_lines.append("| # | Location | Before (summary) | After (summary) | Reason |")
md_lines.append("|---|----------|------------------|------------------|--------|")
for i, e in enumerate(cl, 1):
    loc = e["location"].replace("|", "/")
    before = (e["before_summary"][:220] + "..." if len(e["before_summary"]) > 220 else e["before_summary"]).replace("|", "/")
    after = (e["after_summary"][:220] + "..." if len(e["after_summary"]) > 220 else e["after_summary"]).replace("|", "/")
    reason = (e["reason"][:300] + "..." if len(e["reason"]) > 300 else e["reason"]).replace("|", "/")
    flag = " **[MANUAL]**" if e.get("manual") else ""
    md_lines.append(f"| {i} | {loc}{flag} | {before} | {after} | {reason} |")
md_lines.append("")

# ESPR guideline reference
md_lines.append("## C. ESPR submission-guideline checks performed")
md_lines.append("")
md_lines.append("Source: https://www.springer.com/11356/submission-guidelines (retrieved on the revision date).")
md_lines.append("")
md_lines.append("- **Citation format**: ESPR uses author-year (name and year in parentheses). The reference list is alphabetized by first author's last name. All 33 references in `manuscript_espr.docx` already follow the ESPR style. No reference-list reordering or format conversion was required. Vestigial numbered in-text citations (e.g. `[1-4]`, `[5,16-18]`, `[6-15]`, `[19-26]`, `[21-23]`) outside Section 2 were converted to author-year format using only references already present in the list. **No references were added or removed.**")
md_lines.append("- **Abstract length**: ESPR guidance is \"about 10 to 15 lines\" (no hard word cap). Original abstract was 212 words; revised abstract is 186 words. Compression retained SCCD four-space framework, inverted U + turning point 0.522, supplementary IV turning point 0.570 caveat, OIU/GTI associated pathway mention, DEI attenuation, regional heterogeneity, and governance sequencing policy implication in priority order specified by the task.")
md_lines.append("- **Keywords**: 6 keywords provided (\"smart city construction; spatial coupling coordination; digital space; carbon emissions; governance sequencing; sustainable urban transition\"). ESPR requires 6-8 keywords; satisfies the lower bound. **[MANUAL]** The author may add 1-2 additional keywords before submission to be within the upper-half of the recommended range.")
md_lines.append("- **Self-citations**: ESPR caps self-citations at 5 across all authors in the bibliography. The reference list contains zero self-citations (none of the 33 entries is authored by Ma H or Wang N). No action required.")
md_lines.append("- **Statements and Declarations**: Already present (Funding, Competing Interests, Author Contributions, Data Availability, Code Availability, Generative AI Declaration, Ethics Approval, Acknowledgments). ESPR places this section BEFORE References; the manuscript follows this order. No change required.")
md_lines.append("- **Tables**: All tables (1-7, S1-S2) cited in consecutive order in the text with Arabic numerals and caption (title) provided. Verified all Table 1 through Table 7 references are present in body text.")
md_lines.append("- **Figures**: Figure 1, Figure 2, Figure 3 each have a short caption and are cited in the body text. Verified.")
md_lines.append("- **Supplementary**: ESM_1, ESM_2, ESM_3 are present as separate PDF files. Added an in-text reference sentence in Section 4.3 (\"...are reported in Electronic Supplementary Material (ESM_1, ESM_2, and ESM_3).\") to satisfy the ESPR requirement that supplementary files be specifically cited in text.")
md_lines.append("- **Footnotes/Endnotes**: ESPR prefers footnotes over endnotes. The manuscript does not use endnotes. The VIF note added in Section 4.2 is inserted as a normal-style text paragraph immediately before the Table S1 caption rather than as a footnote; the table-note itself was also extended. **[MANUAL]** If the editorial office requires a numbered footnote rather than an in-text note, the publisher-side production team can convert; the wording will remain identical.")
md_lines.append("- **Heading levels**: Three or fewer displayed heading levels are used (Heading 1 for sections, Heading 2 for subsections). Compliant. No Heading 3 was introduced by the revision.")
md_lines.append("- **Hybrid open-access model**: ESPR is hybrid OA. No open-access author-side preference is required at submission; applicable only upon acceptance.")
md_lines.append("")
md_lines.append("## D. Manual follow-up items")
md_lines.append("")
md_lines.append("1. **[MANUAL]** **Suggested reviewers - corresponding author completion**: The cover letter now contains a `Suggested Reviewers` section with 4 reviewer entries (name, affiliation, email, one-sentence reason). Names, affiliations, and emails are explicitly marked `[To be confirmed]`. The corresponding author (Ningning Wang) must fill in actual contact details before final submission. The 4 suggested reviewer-path fields correspond to (i) urban environmental policy & carbon emissions, (ii) urban spatial structure & econometrics, (iii) China environmental policy & sustainable urban development, and (iv) digital economy & environmental performance. The fourth entry is an additional direction beyond the three categories named in the task; it can be removed if the corresponding author prefers exactly three reviewers.")
md_lines.append("2. **[MANUAL]** **Reference list pre-existing inconsistency**: The reference list entry for EDGAR is `European Commission JointResearchCentre (2025) EDGAR - The Emissions Database for Global Atmospheric Research. Available online: ...`. In-text paragraph [064] cites this source as `(European Commission 2026; Crippa et al. 2024)`, i.e. with year 2026. This year mismatch (2025 in reference list vs 2026 in text) is pre-existing and **was not introduced or modified** by these revisions. Per the project scientific-governance red line, the reference entry was not added/removed and the in-text reference was not silently changed. The corresponding author should reconcile the year (recommend updating the reference-list entry to 2026 to match the in-text citation, since EDGAR v8.0 was released in 2024 and the dataset accessed in 2026).")
md_lines.append("3. **[MANUAL]** **Keyword count**: 6 keywords are provided (ESPR requires 6-8). The lower bound is met; the corresponding author may optionally add 1-2 keywords (e.g. \"digital economy development\", \"governance sequencing\", \"regional heterogeneity\") before submission.")
md_lines.append("4. **[MANUAL]** **VIF note placement as footnote (optional)**: The user-specified VIF sentence was inserted as a normal-style paragraph immediately before the Table S1 caption. If the editorial office prefers a numbered footnote to the table title rather than a body paragraph, the production team can convert; no wording will change. No action needed unless requested by the editor.")
md_lines.append("")
md_lines.append("## E. Cross-reference verification (post-revision)")
md_lines.append("")
md_lines.append("After revisions, the revised manuscript body contains explicit textual references to all required items:")
md_lines.append("")
md_lines.append("- Tables: Table 1, Table 2, Table 3, Table 4, Table 5, Table 6, Table 7, Table S1, Table S2 (all FOUND in body text)")
md_lines.append("- Figures: Figure 1, Figure 2, Figure 3 (all FOUND in body text)")
md_lines.append("- Supplementary: ESM_1, ESM_2, ESM_3 (all FOUND in body text via the new sentence appended to Section 4.3)")
md_lines.append("- Critical comparator citations retained in Section 2: Yang et al. (2022), Ma and Wu (2023), An et al. (2024), Li et al. (2025), Zhu et al. (2024) (all FOUND)")
md_lines.append("- Critical empirical numbers preserved: SCCD = 8.444\*\*\*, SCCD squared = - 8.091\*\*\*, turning point 0.522 / 0.570, F = 61.55 / 80.44 / 380.73 / 16.24 / 81.33, N = 4544/4514/4375/139, SCCD range, etc. (verified unchanged vs. original manuscript)")
md_lines.append("- Total paragraph count: 227 (original) -> 206 (revised)")
md_lines.append("- Section 2 word count: ~1697 words (well within the 3,000-word target)")
md_lines.append("- Abstract word count: 212 (original) -> 186 (revised)")
md_lines.append("- References: 33 entries unchanged, alphabetically ordered")
md_lines.append("")
md_lines.append("## F. Files written")
md_lines.append("")
md_lines.append("- `manuscript_espr_revised.docx` - Revised manuscript")
md_lines.append("- `cover_letter_espr_revised.docx` - Revised cover letter")
md_lines.append("- `revision_log.md` - This file")
md_lines.append("- `_apply_manuscript_revisions.py`, `_apply_cover_revisions.py` - Revision scripts (kept for audit trail)")
md_lines.append("- `_revisions_log_data.json`, `_revisions_log_cover.json` - Machine-readable change entries (intermediate)")
md_lines.append("")

out_path = Path(r"D:\Workplace\SCS\outputs\espr_submission\revision_log.md")
out_path.write_text("\n".join(md_lines), encoding="utf-8")
print("written:", out_path)
print("word count of revision_log:", len("\n".join(md_lines).split()))