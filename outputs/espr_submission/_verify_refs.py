"""Verify reference list is unchanged (no additions or deletions)."""
from pathlib import Path
from docx import Document

SRC_ORIG = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_revised.docx")
SRC_FINAL = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_final.docx")

doc_orig = Document(str(SRC_ORIG))
doc_final = Document(str(SRC_FINAL))

def get_refs(d):
    refs = []
    in_refs = False
    for p in d.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip().lower().startswith("references"):
            in_refs = True
            continue
        if in_refs and p.text.strip():
            refs.append(p.text.strip())
    return refs

orig_refs = get_refs(doc_orig)
final_refs = get_refs(doc_final)

print(f"Original references: {len(orig_refs)}")
print(f"Final references:    {len(final_refs)}")
print()

if len(orig_refs) == len(final_refs):
    print("Reference COUNT matches.")
else:
    print(f"REFERENCE COUNT MISMATCH: {len(orig_refs)} vs {len(final_refs)}")

# Check each reference matches (allowing for the EC year change)
mismatches = []
for i, (o, f) in enumerate(zip(orig_refs, final_refs)):
    if o != f:
        # Expected: only the EC 2025->2026 change
        if "European Commission" in o and "JointResearch" in o:
            print(f"  [{i}] EXPECTED CHANGE (EC year):")
            print(f"    ORIG:  {o}")
            print(f"    FINAL: {f}")
        else:
            mismatches.append((i, o, f))
            print(f"  [{i}] UNEXPECTED CHANGE:")
            print(f"    ORIG:  {o[:150]}")
            print(f"    FINAL: {f[:150]}")
else:
    if not mismatches:
        print("All references match (except expected EC year correction).")

# Check for any added/removed refs
if len(orig_refs) != len(final_refs):
    orig_set = set(orig_refs)
    final_set = set(final_refs)
    removed = orig_set - final_set
    added = final_set - orig_set
    if removed:
        print(f"REMOVED: {removed}")
    if added:
        print(f"ADDED: {added}")
