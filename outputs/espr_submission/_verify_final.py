"""Verify final manuscript changes."""
from pathlib import Path
from docx import Document

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_final.docx")
doc = Document(str(SRC))

print("=== TASK 1 VERIFICATION: Sustainability residuals ===")
residuals = [
    "the main Sustainability evidence",
    "keeps the Sustainability message",
    "Sustainability reviewers",
    "For Sustainability readers",
    "aligns with Sustainability's applied policy audience",
]
full_text = "\n".join(p.text for p in doc.paragraphs)
for r in residuals:
    status = "STILL PRESENT" if r in full_text else "REMOVED"
    print(f"  [{status}] '{r}'")

# Check replacements are present
replacements = [
    "the main evidence",
    "keeps the paper's message",
    "clear to reviewers who",
    "For readers, this matters",
    "aligns with the applied policy audience",
]
print()
print("=== Replacement verification ===")
for r in replacements:
    status = "PRESENT" if r in full_text else "MISSING"
    print(f"  [{status}] '{r}'")

print()
print("=== Other Sustainability mentions (should be journal refs or concept) ===")
for i, p in enumerate(doc.paragraphs):
    if "Sustainability" in p.text:
        idx = p.text.find("Sustainability")
        s = max(0, idx - 50)
        e = min(len(p.text), idx + 70)
        print(f"  [{i:03d}] ...{p.text[s:e]}...")

print()
print("=== TASK 3 VERIFICATION: European Commission citation ===")
for i, p in enumerate(doc.paragraphs):
    if "European Commission" in p.text and "JointResearch" in p.text:
        print(f"  [{i:03d}] {p.text}")

print()
print("=== TASK 2 VERIFICATION: Section 1 end ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("The paper also examines associated restructuring"):
        print(f"  [{i:03d}] ({len(p.text.split())}w) {p.text[:200]}...")
    if p.text.strip().startswith("The rest of the paper proceeds as follows"):
        print(f"  [{i:03d}] ({len(p.text.split())}w) {p.text[:200]}...")

print()
print("=== TASK 2 VERIFICATION: Section 5.1 ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 2" and "5.1." in p.text:
        # print this and next few
        for j in range(i, min(i+4, len(doc.paragraphs))):
            pp = doc.paragraphs[j]
            print(f"  [{j:03d}] <{pp.style.name}> ({len(pp.text.split())}w) {pp.text[:150]}...")
        break

print()
print("=== TASK 2 VERIFICATION: Section 5.3 ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 2" and "5.3." in p.text:
        for j in range(i, min(i+4, len(doc.paragraphs))):
            pp = doc.paragraphs[j]
            print(f"  [{j:03d}] <{pp.style.name}> ({len(pp.text.split())}w) {pp.text[:150]}...")
        break

print()
print("=== EMPIRICAL NUMBERS CHECK ===")
checks = [
    "8.444", "-8.091", "0.522", "0.570", "61.55", "80.44", "380.73",
    "4,514", "4,544", "4,375", "139", "0.442", "0.601",
]
for c in checks:
    status = "PRESENT" if c in full_text else "MISSING"
    print(f"  [{status}] {c}")

print()
print("=== LIMITATIONS 5.4 CHECK ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 2" and "5.4." in p.text:
        for j in range(i, min(i+10, len(doc.paragraphs))):
            pp = doc.paragraphs[j]
            if pp.style.name.startswith("Heading"):
                break
            print(f"  [{j:03d}] ({len(pp.text.split())}w) {pp.text[:120]}...")
        break

print()
total = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"=== TOTAL WORDS: {total} ===")
