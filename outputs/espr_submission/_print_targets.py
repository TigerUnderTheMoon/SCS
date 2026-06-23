"""Print full text of target paragraphs for precise editing."""
from pathlib import Path
from docx import Document

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_revised.docx")
doc = Document(str(SRC))

targets = [14, 15, 16, 17, 18, 19, 60, 63, 83, 125, 126, 127, 128, 129, 130, 132, 135, 136, 137, 138, 183]
for i in targets:
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        wc = len(p.text.split())
        print(f"===== [{i:03d}] <{p.style.name}> ({wc} words) =====")
        print(p.text)
        print()
