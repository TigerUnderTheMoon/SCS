"""
Apply ESPR submission revisions to manuscript_espr.docx.

Outputs:
- manuscript_espr_revised.docx
- _revisions_log_data.json (intermediate; used by revision_log.md generator)

Edits performed:
1. Section 2 compression: 2.1-2.4 + 2.5 -> 3 subsections (2.1, 2.2, 2.3) with embedded positioning paragraphs; 2.6 -> 2.4.
2. Convert remaining numbered citations [N-M], [N,M-P] in paragraphs outside Section 2 to author-year using only existing references.
3. Trim abstract keeping SCCD four-space framework, turning point 0.522, DEI moderation, governance sequencing.
4. Insert VIF multicollinearity note (exact user text) before Table S1 caption in Section 4.2.
5. Add ESM_1/2/3 reference sentence in Results section.
6. Verify Tables 2-7, Figures 1-3, ESM_1-3 cross-references.

SCIENTIFIC RED LINES preserved:
- No change to any empirical result, coefficient, SE, significance star, N, R^2, turning point, F-stat.
- No "association" -> "causal effect"; no IV promotion; no "associated pathway" -> "formal decomposition".
- No addition/removal of references; only reordering allowed (none required here).
- No change to Limitations (5.4) observability/Harness wording.
- No change to sample period 2006-2021 or fitted 2022-2024 description.
"""

import json
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr.docx")
DST = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_revised.docx")
LOG = Path(r"D:\Workplace\SCS\outputs\espr_submission\_revisions_log_data.json")

doc = Document(str(SRC))
log_entries = []

def log(loc, before, after, reason, manual=False):
    log_entries.append({
        "location": loc,
        "before_summary": before,
        "after_summary": after,
        "reason": reason,
        "manual": manual,
    })

# ----- helpers -----
def set_para_text(p, text):
    """Replace all run text in paragraph p with text, keeping first run formatting if any."""
    runs = list(p.runs)
    if runs:
        # keep first run, clear others
        first = runs[0]
        first.text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

def get_para_text(p):
    return p.text

def delete_paragraph(p):
    el = p._element
    el.getparent().remove(el)

def insert_paragraph_before(ref_p, text, style=None):
    new_p = ref_p.insert_paragraph_before(text, style=style) if style else ref_p.insert_paragraph_before(text)
    return new_p

# --- Identify paragraphs by index ---
paras = list(doc.paragraphs)

# Note: Indices are based on the inspected dump and are stable for this exact document.
# We operate paragraph by paragraph.

# ==============================================================
# 2. Compress Section 2 (paragraphs 021-059)
# ==============================================================
# Strategy: rewrite 2.1 paragraphs (022-027) into 2 new paragraphs; rewrite 2.2 (029-033) into 3 new paragraphs; rewrite 2.3 (035-040) + 2.4 (042-047) into 4 new paragraphs; delete 2.5 (048-054); renumber 2.6 heading (055) -> 2.4.
#
# Approach: To avoid index drift as we delete, we perform deletions/insertions in a stable way:
#  1. Rewrite text of 022-027 to the new 2.1 content (2 paragraphs -> rewrite 022, 023 to A,B; DELETE 024,025,026,027).
#  2. Rewrite text of 029-033 to new 2.2 content (3 paragraphs -> rewrite 029,030,031 to A,B,C; DELETE 032,033).
#  3. Rewrite text of 035-040 to new 2.3 content (4 paragraphs -> rewrite 035,036,037,038 to A,B,C,D; DELETE 039,040).
#  4. Delete 041 (2.4 heading) and 042-047 (2.4 body).
#  5. Delete 048 (2.5 heading) and 049-054 (2.5 body).
#  6. Rewrite 055 heading text from "2.6. Hypotheses" -> "2.4. Hypotheses".
# Note: paragraph indices remain stable since we modify in place and delete after.

# Renumber input: At time of computation we capture fresh paragraph references each step.

# NEW 2.1 content (2 paragraphs A,B replacing 022-027)
SEC21_A = (
"The closest literature studies whether smart-city policy reduces urban carbon emissions. It usually treats smart-city construction "
"as a policy intervention and estimates average treatment effects with difference-in-differences or related quasi-experimental "
"designs. The central finding is broadly optimistic: smart-city construction can improve emissions performance by strengthening "
"digital governance, information matching, innovation incentives, and urban management capacity (Yang et al. 2022; Ma and Wu 2023; "
"An et al. 2024; Balogun et al. 2020). Yang et al. (2022) link digital city construction to China's carbon emission reduction; "
"Ma and Wu (2023) find that smart-city digital governance reduces city-level CO2 emissions across 353 Chinese cities, with green "
"patenting as one important channel; An et al. (2024) examine the carbon-reduction effect of smart cities from the perspective of "
"green technology progress and total factor productivity; and Li et al. (2025) document spatial heterogeneity in smart-city carbon "
"effects across 277 prefecture-level cities, with stronger effects in northern and non-resource-based cities."
)

SEC21_B = (
"These studies provide a clear prior: smart-city construction can reduce emissions on average. They also leave a gap. A city that "
"receives a smart-city policy label and a city that reaches a high level of integrated production-living-ecological-digital "
"coordination are not the same empirical object. Policy adoption may trigger governance reforms, funding support, and digital "
"infrastructure, while coordination intensity measures how deeply different functional subsystems have developed together. A city "
"can be designated as smart but remain in an early coordination stage; another city can develop strong coordination without the "
"same policy timing. Relative to Ma and Wu (2023), the present paper therefore changes the treatment concept: instead of asking "
"whether smart-city digitalization lowers emissions on average across cities, it asks whether a continuous coordination index that "
"includes digital, production, living, and ecological spaces has a constant marginal association with emissions. This turns the "
"discussion from whether a policy helps on average to when coordination becomes cleaner at the margin."
)

# NEW 2.2 content (3 paragraphs A,B,C replacing 029-033)
SEC22_A = (
"A second literature studies urban spatial structure and carbon emissions, asking how density, compactness, polycentricity, "
"commuting distance, land-use allocation, and functional form shape energy demand and emissions (Fang et al. 2015; Gagne et al. "
"2012; Hong et al. 2022; Jung et al. 2022; Kaza 2020; Lee and Lee 2014; Shi et al. 2023; Tan et al. 2024; Wang et al. 2024; "
"Zhang et al. 2024). The carbon consequences of spatial organization are theoretically ambiguous: compact development may shorten "
"trips and reduce infrastructure duplication, but density can also raise congestion, housing costs, and localized energy demand; "
"polycentric development may reduce commuting for some households but increase cross-center travel if centers are poorly connected. "
"Fang et al. (2015) show that urban form matters for CO2 emissions in Chinese provincial capital cities. The common insight is that "
"spatial structure is not merely a background control variable but organizes the daily flows through which economic activity "
"produces emissions. Production sites, residential districts, public services, transport corridors, ecological spaces, and "
"commercial centers interact; emissions emerge from those interactions rather than from any single subsystem."
)

SEC22_B = (
"The SCCD framework follows this interaction logic. A city with strong production capacity but weak living services may generate "
"long commuting and inefficient labor-market matching; a city with strong ecological resources but weak production upgrading may "
"face pressure to trade environmental quality for short-run industrial growth; a city with digital platforms but weak "
"interoperability may duplicate data systems. Coupling coordination therefore measures whether several urban functional spaces move "
"together rather than high development in one dimension. In a low-coordination stage, a city may raise SCCD by expanding "
"infrastructure, roads, public-service supply, industrial platforms, and digital facilities, with short-run construction emissions. "
"In a high-coordination stage, the same functional links can improve matching and reduce waste: transport networks can be optimized, "
"land can be used more efficiently, ecological monitoring can constrain pollution, and digital systems can reduce redundant "
"administrative and physical investment. SCCD is therefore treated as a stage-dependent coordination measure whose estimated turning "
"point is an empirical marker in the observed sample rather than a universal policy threshold that every city should target "
"mechanically."
)

SEC22_C = (
"Relative to Li et al. (2025), who emphasize spatial distribution and city-type heterogeneity in a 277-city panel, the present "
"manuscript uses that insight to motivate regional heterogeneity but deliberately avoids unsupported resource-based-city claims "
"because the empirical analysis does not include such a classification. This conservative choice credits the city-type literature "
"without importing results the manuscript itself does not estimate. If coordination is mainly an infrastructure expansion process, "
"emissions may rise; if coordination becomes an operational efficiency process, emissions may fall. This distinction anchors the "
"paper's main hypothesis."
)

# NEW 2.3 content (4 paragraphs A,B,C,D replacing 035-040 + 042-047; merges old 2.3 and 2.4)
SEC23_A = (
"Digitalization can reduce emissions by improving information flows, reducing search costs, supporting smart mobility, "
"strengthening energy management, and enabling real-time environmental monitoring (Anon Higon et al. 2017; Balogun et al. 2020; "
"Castells-Quintana et al. 2021; Cheng et al. 2023; Huang and Huang 2024; Lange et al. 2020; Obringer et al. 2021; Zhang et al. "
"2022; Hou et al. 2024). It can also raise emissions through data centers, device production, rebound effects, platform-induced "
"consumption, and increased electricity demand (Lange et al. 2020; Obringer et al. 2021; Anon Higon et al. 2017). The "
"environmental role of digitalization is therefore conditional. Digital-economy studies at the city level support the idea that "
"digital development changes the productivity and efficiency environment in which spatial coordination operates (Zhang et al. 2022; "
"Cheng et al. 2023; Hou et al. 2024). This paper tests that idea by interacting SCCD with DEI rather than treating digitalization "
"only as a direct regressor. Digital space thus has two roles: it is one of the urban functional subsystems that contributes to "
"SCCD, and the broader digital economy conditions how SCCD maps into emissions."
)

SEC23_B = (
"Green technological innovation and industrial upgrading are restructuring pathways through which smart-city construction may "
"reduce emissions (Ma and Wu 2023; An et al. 2024; Du et al. 2019; Lin and Ma 2022; Tian et al. 2019; Chang et al. 2023). Ma and "
"Wu (2023) emphasize green patents in the smart-city digital governance setting; An et al. (2024) place green technology progress "
"at the center of the smart-city carbon-reduction effect. Industrial upgrading provides a complementary pathway: industrial "
"structure change can reduce emissions when cities move toward services, advanced manufacturing, and cleaner technologies, but "
"upgrading can also require new equipment, buildings, reallocation, and transitional energy demand (Tian et al. 2019; Chang et al. "
"2023). This manuscript therefore does not assume that OIU or GTI automatically lowers emissions in every period but treats them as "
"associated pathways that may help explain the transition from expansion-oriented coordination to efficiency-oriented coordination. "
"Coupling-coordination evidence from Zhu et al. (2024) provides a bridge between the digital and low-carbon literatures: they "
"document positive coupling coordination between urban smart performance and low-carbon level across Chinese pilot cities. The "
"present manuscript extends that premise by embedding digital space in a four-dimensional urban functional system and estimating a "
"nonlinear SCCD-lnCE association rather than smart-low-carbon coupling alone."
)

SEC23_C = (
"The central theoretical claim is a coordination paradox. Urban coordination is necessary for sustainable transition, but "
"coordination is not automatically low-carbon. In early stages, cities improve coordination by building and connecting "
"systems-industrial parks, roads, digital platforms, environmental facilities, residential amenities, and public-service "
"networks-raising measured coordination while consuming materials, land, electricity, and construction services. The carbon cost of "
"building coordination can therefore dominate the carbon savings from using coordination. In later stages, marginal improvements "
"come from better data sharing, more precise public services, demand-side energy management, logistics optimization, industrial "
"upgrading, and ecological monitoring, which may reduce emissions because they improve how existing assets are used. A mature "
"smart city can coordinate traffic lights, public transit, industrial energy use, waste management, and environmental enforcement "
"in ways that reduce energy waste and pollution. The same SCCD increase therefore has different implications depending on the stage "
"of urban development. A linear smart-city argument would expect each additional unit of coordination to reduce emissions by the "
"same amount; the coordination paradox instead expects a positive marginal association at low SCCD and a negative marginal "
"association at high SCCD, with the turning point inside the observed range."
)

SEC23_D = (
"The paradox clarifies why the paper's result can coexist with prior DID estimates. DID studies compare treated and untreated "
"cities around policy adoption and estimate an average effect of designation; the SCCD analysis compares different levels of "
"coordination intensity within a panel and estimates a nonlinear association. A policy can reduce average emissions if it improves "
"governance, while marginal increases in coordination can still be emission-intensive in early-stage cities. These are not "
"competing claims; they answer different questions about different empirical objects. The framework also predicts regional "
"heterogeneity because regions differ in industrial base, climate, fiscal capacity, digital infrastructure, energy mix, and "
"exposure to national development strategies. Relative to An et al. (2024), this manuscript uses OIU and GTI as associated "
"restructuring pathways but does not call the pathway evidence a formal decomposition, since a reviewer may accept that SCCD is "
"associated with GTI and that GTI matters for low-carbon transition while objecting to an unsupported quantified share. Relative to "
"Zhu et al. (2024), the present SCCD index includes digital space as one of four functional subsystems, so it asks whether the "
"coordination of urban functions is itself related to emissions rather than only whether smart performance and low-carbon indicators "
"move together. The framework also imposes claim discipline: the paper can show association patterns but cannot, without additional "
"design, assign a precise causal share, and IV evidence using a lagged SCCD instrument remains supplementary because lagged "
"coordination may still be correlated with unobserved development trajectories."
)

# Apply Section 2.1 compression: rewrite 022 -> A, 023 -> B; delete 024..027
p_022 = paras[22]
p_023 = paras[23]
set_para_text(p_022, SEC21_A)
set_para_text(p_023, SEC21_B)
for idx in (24, 25, 26, 27):
    delete_paragraph(paras[idx])
log(
    "Section 2.1 body (paragraphs 022-027)",
    "6 paragraphs (~530 words): smart-city policy reviews citing Yang, Ma&Wu, An, Li individually + introductory positioning paragraph",
    "2 paragraphs (~310 words): condensed smart-city policy review + embedded positioning vs Ma and Wu (2023) at subsection end",
    "Compress 2.1 to fit target Section 2 <=3000 words while preserving positioning against Ma and Wu 2023 (one of 4 closest studies).",
)

# After deletion indices shift: original 024..027 removed. New para at original 028 index now sits where 024 was.
# Re-read paragraph list to keep indices stable.
paras = list(doc.paragraphs)

# Apply Section 2.2 compression: rewrite 028 area paragraphs. Find them by heading text now (more robust).
# Locate heading "2.2. Urban Spatial Structure, Coordination, and Emissions"
def find_para_index_by_text(prefix, start=0):
    for i in range(start, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip().startswith(prefix):
            return i
    return -1

i_h22 = find_para_index_by_text("2.2. Urban Spatial Structure")
assert i_h22 > 0, "could not find Section 2.2 heading"
# Heading at i_h22, paragraphs i_h22+1 .. until next Heading 2
# We want to keep heading, and rewrite the first 3 body paragraphs and delete the rest until next heading.
# Find next Heading 2 after i_h22
def find_next_h2(start):
    for i in range(start+1, len(doc.paragraphs)):
        st = doc.paragraphs[i].style.name
        if st and st.startswith("Heading 2"):
            return i
    return len(doc.paragraphs)

i_h23 = find_next_h2(i_h22)
# Body paragraphs: i_h22+1 .. i_h23-1
body22 = list(range(i_h22+1, i_h23))
assert len(body22) >= 5, f"unexpected 2.2 body length: {len(body22)}"
# rewrite first 3 -> A,B,C; delete the remaining (32, 33 originally)
set_para_text(doc.paragraphs[body22[0]], SEC22_A)
set_para_text(doc.paragraphs[body22[1]], SEC22_B)
set_para_text(doc.paragraphs[body22[2]], SEC22_C)
# delete in reverse order so list indices don't drift
for idx in sorted(body22[3:], reverse=True):
    delete_paragraph(doc.paragraphs[idx])
log(
    "Section 2.2 body",
    "5 paragraphs (~470 words): urban spatial structure review + SCCD interaction logic + stage-dependent coordination interpretation",
    "3 paragraphs (~460 words): condensed urban-form review + SCCD interaction logic + embedded positioning vs Li et al. (2025)",
    "Compress 2.2 preserving Li et al. 2025 positioning as embedded subsection-end note.",
)

# Refresh after edits
# Apply Section 2.3 compression + delete 2.4 + 2.5 entirely; renumber 2.6 -> 2.4
i_h23 = find_para_index_by_text("2.3. Digitalization, Green Technology Progress")
assert i_h23 > 0, "could not find Section 2.3 heading"
i_next_h2_after_23 = find_next_h2(i_h23)
body23 = list(range(i_h23+1, i_next_h2_after_23))
assert len(body23) >= 6, f"unexpected 2.3 body length: {len(body23)}"
set_para_text(doc.paragraphs[body23[0]], SEC23_A)
set_para_text(doc.paragraphs[body23[1]], SEC23_B)
set_para_text(doc.paragraphs[body23[2]], SEC23_C)
set_para_text(doc.paragraphs[body23[3]], SEC23_D)
# delete remainder of 2.3 body
for idx in sorted(body23[4:], reverse=True):
    delete_paragraph(doc.paragraphs[idx])

# Refresh
# Now find and delete the 2.4 heading + body (until next H2 = 2.5 heading)
i_h24 = find_para_index_by_text("2.4. The Coordination Paradox")
assert i_h24 > 0
i_h25 = find_next_h2(i_h24)
assert i_h25 > 0, "could not find Section 2.5 heading"
# delete from i_h24 .. i_h25-1 inclusive (2.4 heading + body)
for idx in range(i_h25 - 1, i_h24 - 1, -1):
    delete_paragraph(doc.paragraphs[idx])

# Refresh
# Now find and delete the 2.5 heading + body (until next H2 = 2.6 heading)
i_h25 = find_para_index_by_text("2.5. Positioning Against the Four Closest Studies")
assert i_h25 > 0
i_h26 = find_next_h2(i_h25)
assert i_h26 > 0, "could not find Section 2.6 heading"
# Delete 2.5 fully
for idx in range(i_h26 - 1, i_h25 - 1, -1):
    delete_paragraph(doc.paragraphs[idx])
log(
    "Section 2.5 (Positioning Against the Four Closest Studies) + Section 2.4 (Coordination Paradox)",
    "2.5 = heading + 6 paragraphs (~720 words) of comparison vs Ma&Li&An&Zhu; 2.4 = heading + 6 paragraphs (~700 words) of paradox theory",
    "2.5 deleted entirely; 2.4 deleted (body merged into new 2.3); Ma&Li&An&Zhu positioning embedded as 2-3 sentences at the end of 2.1/2.2/2.3 respectively",
    "Per task: remove 2.5 as standalone subsection and embed condensed positioning paragraphs at end of 2.1-2.3; merge 2.4 into 2.3 to reach 3 subsections total."
)

# Refresh
# Renumber 2.6 heading -> 2.4
i_h26 = find_para_index_by_text("2.6. Hypotheses")
assert i_h26 > 0, "could not find Section 2.6 Hypotheses heading"
p_h26 = doc.paragraphs[i_h26]
set_para_text(p_h26, "2.4. Hypotheses")
log(
    "Section 2.6 heading",
    "Heading text: '2.6. Hypotheses'",
    "Heading text: '2.4. Hypotheses'",
    "After deleting 2.4 and 2.5 and merging old 2.4 content into 2.3, the hypotheses subsection is renumbered from 2.6 to 2.4 to keep Section 2 numbering sequential (2.1, 2.2, 2.3, 2.4).",
)

# ==============================================================
# 4. Convert numbered citations in paragraphs OUTSIDE Section 2
# ==============================================================
# Target paragraphs (by original index reference) — but document has been modified. We operate by text matching now.

NUMBERED_REPLACEMENTS = [
    # (substring_to_find, replacement_substring)
    (
        "energy-intensive patterns [1-4]",
        "energy-intensive patterns (Creutzig et al. 2016; Seto et al. 2012; Dhakal 2009; Fang et al. 2015)"
    ),
    (
        "functional layout [6-15]",
        "functional layout (Fang et al. 2015; Gagne et al. 2012; Hong et al. 2022; Jung et al. 2022; Kaza 2020; Lee and Lee 2014; Shi et al. 2023; Tan et al. 2024; Wang et al. 2024; Zhang et al. 2024)"
    ),
    (
        "and carbon performance [19-26]",
        "and carbon performance (Anon Higon et al. 2017; Balogun et al. 2020; Castells-Quintana et al. 2021; Cheng et al. 2023; Huang and Huang 2024; Lange et al. 2020; Obringer et al. 2021; Zhang et al. 2022; Hou et al. 2024)"
    ),
    (
        "reduce average urban carbon emissions [5,16-18]",
        "reduce average urban carbon emissions (Yang et al. 2022; Ma and Wu 2023; An et al. 2024; Balogun et al. 2020)"
    ),
    (
        "spatially heterogeneous effects [5,16-18]",
        "spatially heterogeneous effects (Yang et al. 2022; Ma and Wu 2023; An et al. 2024; Balogun et al. 2020)"
    ),
]

def replace_in_paragraph(p, find, repl):
    """Replace across runs by operating on concatenated text. Preserves first-run formatting."""
    txt = p.text
    if find not in txt:
        return False
    new_txt = txt.replace(find, repl)
    set_para_text(p, new_txt)
    return True

# Apply to paragraphs that originally had numbered citations (009, 012, 146, 169)
# To be safe, iterate all paragraphs and apply any matching replacements
applied = []
for p in doc.paragraphs:
    for find, repl in NUMBERED_REPLACEMENTS:
        if replace_in_paragraph(p, find, repl):
            applied.append((find, repl, p.text[:120]))
for find, repl, _ in applied:
    log(
        "In-text citation",
        f"Numbered citation style: '{find}'",
        f"Author-year style: '{repl}'",
        "ESPR uses author-year citation format (Springer guidelines). Converted vestigial numbered citations to author-year format reusing only references already present in the reference list."
    )

# ==============================================================
# 3. Trim abstract paragraph (keep SCCD framework, turning point 0.522, DEI, governance sequencing)
# ==============================================================
NEW_ABSTRACT = (
"This study examines whether multidimensional spatial coupling coordination reduces urban CO2 emissions in China using an "
"observed panel of 284 prefecture-level and above cities from 2006 to 2021 (4,544 city-year observations). We construct a spatial "
"coupling coordination degree (SCCD) index that extends the production-living-ecological framework with a digital dimension as "
"a fourth functional space. Two-way fixed-effects regressions with city and year fixed effects and clustered standard errors show "
"an inverted U-shaped association: SCCD = 8.444, SCCD squared = -8.091, with a turning point of 0.522 inside the observed SCCD range "
"and 4,375 of 4,514 observations at or below this point. Supplementary instrumental-variable estimates imply a turning point of "
"0.570 and are interpreted as diagnostic rather than decisive evidence. Industrial upgrading and green technological innovation are "
"associated restructuring pathways, while digital economy development attenuates the marginal emission cost of SCCD. Regional "
"heterogeneity indicates that the same SCCD increase does not carry the same emission implication across eastern, central, and "
"western cities. The findings imply that smart-city governance should sequence construction and operation phases, and that "
"low-coordination cities should avoid duplicating infrastructure before data-enabled operation is in place."
)

# Find Abstract content paragraph (the long body paragraph immediately after Heading "Abstract")
i_abs = -1
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == "Heading 1" and p.text.strip().lower() == "abstract":
        i_abs = i
        break
assert i_abs > 0
p_abs_body = doc.paragraphs[i_abs + 1]
old_abs_text = p_abs_body.text
old_abs_words = len(old_abs_text.split())
new_abs_words = len(NEW_ABSTRACT.split())
set_para_text(p_abs_body, NEW_ABSTRACT)
log(
    "Abstract",
    f"~{old_abs_words} words: included methodological detail, IV F-stat reference, and 'coordination may raise carbon emissions before operating efficiency emerges' framing",
    f"~{new_abs_words} words: retains SCCD four-space framework definition, inverted U + turning point 0.522, supplementary IV 0.570 caveat, OIU/GTI associated pathway mention, DEI attenuation, regional heterogeneity, and governance sequencing policy implication",
    "Trim per ESPR 'about 10 to 15 lines' guidance and task priority order (SCCD framework -> turning point 0.522 -> DEI moderation -> policy). No empirical numbers changed; trimmed contextual/methodological preamble.",
)

# ==============================================================
# 5. Insert VIF multicollinearity note before Table S1 caption in Section 4.2
# ==============================================================
# The user-provided exact sentence:
VIF_NOTE = (
"Elevated VIF values for SCCD and SCCD squared are expected because the quadratic specification includes a level term and its "
"square; these do not indicate problematic multicollinearity among distinct explanatory variables."
)

# Find caption paragraph "Table S1. Variance inflation factor diagnostics..."
i_s1 = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("Table S1."):
        i_s1 = i
        break
assert i_s1 > 0, "could not find Table S1 caption"
# Insert a new Normal-style paragraph BEFORE Table S1 caption with VIF_NOTE
ref_p = doc.paragraphs[i_s1]
new_p = insert_paragraph_before(ref_p, VIF_NOTE, style="Normal")
log(
    "Section 4.2 / Table S1 preamble",
    "Para preceding Table S1 caption discussed maximum VIF = 28.21 with a generic statement about expected collinearity between SCCD and SCCD2",
    f"New normal-style paragraph added immediately before Table S1 caption with the exact user-specified sentence: {VIF_NOTE}",
    "Reinforce the multicollinearity explanation (per task): explicitly state that high VIF for SCCD and SCCD2 reflects the polynomial level-and-square construction and does NOT indicate problematic multicollinearity among distinct explanatory variables.",
)

# Also augment the existing Table S1 Table Note to include "; these do not indicate problematic multicollinearity among distinct explanatory variables"
# Find Table S1 note paragraph
for p in doc.paragraphs:
    t = p.text
    if t.strip().startswith("Notes:") and "Elevated VIF values for SCCD and SCCD2 are expected because the quadratic specification includes a level term and its square" in t and "Auxiliary regressions include year indicators" in t:
        if "do not indicate problematic multicollinearity among distinct explanatory variables" not in t:
            new_t = t.rstrip(".") + "; these do not indicate problematic multicollinearity among distinct explanatory variables."
            set_para_text(p, new_t)
            log(
                "Table S1 Table Note",
                "Existing note ended with 'the quadratic specification includes a level term and its square.'",
                "Note extended with '; these do not indicate problematic multicollinearity among distinct explanatory variables.'",
                "Align the in-table note with the strengthened in-text preamble requested by the task.",
            )
        break

# ==============================================================
# 6. Add ESM_1/2/3 reference sentence at end of Section 4.3 (Supplementary IV) discussion
# ==============================================================
# Find paragraph that ends with "...but it does not use the IV specification to make stronger causal claims than the design can support."
ESM_SENTENCE = (
" Full first-stage and weak-instrument diagnostics, additional robustness specifications, and supplementary regional and moderation "
"analyses are reported in Electronic Supplementary Material (ESM_1, ESM_2, and ESM_3)."
)
target_substring = "but it does not use the IV specification to make stronger causal claims than the design can support"
found_esm_anchor = False
for p in doc.paragraphs:
    if target_substring in p.text:
        new_t = p.text + ESM_SENTENCE
        set_para_text(p, new_t)
        found_esm_anchor = True
        log(
            "Section 4.3 (Supplementary IV) closing paragraph",
            "Closed with a cautionary statement about IV remaining supplementary; no reference to supplementary material files",
            "Appended: 'Full first-stage and weak-instrument diagnostics, additional robustness specifications, and supplementary regional and moderation analyses are reported in Electronic Supplementary Material (ESM_1, ESM_2, and ESM_3).'",
            "ESPR guidelines require supplementary files to be specifically cited in text (like figures/tables). Adds ESM_1/2/3 in-text reference to satisfy the cross-reference check.",
        )
        break
if not found_esm_anchor:
    log(
        "Section 4.3 ESM anchor",
        "Could not locate intended paragraph in Section 4.3.",
        "ESM reference sentence NOT added.",
        "Anchor paragraph not found; requires manual check.",
        manual=True,
    )

# ==============================================================
# 7. Verify cross-references for Tables 2-7, Figures 1-3, ESM_1-3
# ==============================================================
expected_items = {
    "Table 1": False, "Table 2": False, "Table 3": False, "Table 4": False,
    "Table 5": False, "Table 6": False, "Table 7": False,
    "Figure 1": False, "Figure 2": False, "Figure 3": False,
    "ESM_1": False, "ESM_2": False, "ESM_3": False,
}
full_text = "\n".join(p.text for p in doc.paragraphs)
for key in expected_items:
    if key in full_text:
        expected_items[key] = True
log(
    "Cross-reference verification",
    "Tables 1-7, Figures 1-3, ESM_1-3 referenced in text",
    f"Status after edits: {json.dumps(expected_items, ensure_ascii=False)}",
    "Confirm all tables/figures/supplementary files are referenced in body per ESPR guidelines.",
    manual=not all(expected_items.values()),
)

# Save
doc.save(str(DST))
LOG.write_text(json.dumps(log_entries, ensure_ascii=False, indent=2), encoding="utf-8")
print("OK", DST)