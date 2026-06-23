"""Comprehensive ESPR (Springer) submission compliance check."""
from pathlib import Path
from docx import Document
import re

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_final.docx")
doc = Document(str(SRC))
paras = doc.paragraphs
full_text = "\n".join(p.text for p in paras)

print("=" * 70)
print("ESPR SUBMISSION COMPLIANCE CHECK")
print("=" * 70)

# ============================================================
# 1. TITLE PAGE
# ============================================================
print("\n[1] TITLE PAGE")
print("-" * 40)
# Check for title, authors, affiliations, corresponding author
has_title = False
has_authors = False
has_affiliation = False
has_corresponding = False
has_email = False
for i, p in enumerate(paras[:15]):
    t = p.text.strip()
    if not t:
        continue
    style = p.style.name
    if i == 0 or (style == "Title" and t):
        has_title = True
        print(f"  Title: {t[:100]}")
    if "author" in t.lower() or (i <= 3 and len(t.split()) <= 10 and not t.startswith("Smart")):
        if "@" in t or "corresponding" in t.lower():
            has_corresponding = True
    if "@" in t:
        has_email = True
        print(f"  Email found: {t[:80]}")
    if "corresponding" in t.lower():
        has_corresponding = True
        print(f"  Corresponding author: {t[:80]}")
    if "university" in t.lower() or "institute" in t.lower() or "college" in t.lower() or "school" in t.lower():
        has_affiliation = True
        print(f"  Affiliation: {t[:80]}")

# Check title for acronyms and region/country
title_text = paras[0].text if paras else ""
acronyms_in_title = re.findall(r"\b[A-Z]{3,}\b", title_text)
if acronyms_in_title:
    print(f"  WARNING: Acronyms in title: {acronyms_in_title}")
if "China" not in title_text and "china" not in title_text:
    print(f"  NOTE: Title does not mention region/country (local study should indicate)")
    print(f"        Title: {title_text[:120]}")

# ============================================================
# 2. ABSTRACT
# ============================================================
print("\n[2] ABSTRACT")
print("-" * 40)
abs_idx = -1
for i, p in enumerate(paras):
    if p.style.name == "Heading 1" and "abstract" in p.text.lower():
        abs_idx = i
        break
    if p.text.strip().lower().startswith("abstract"):
        abs_idx = i
        break
if abs_idx >= 0:
    # Get abstract body (next paragraph or same)
    abs_body = paras[abs_idx + 1].text if abs_idx + 1 < len(paras) else ""
    abs_words = len(abs_body.split())
    abs_lines = len(abs_body.split("\n"))
    # Estimate lines by character count (~80 chars per line)
    est_lines = max(1, len(abs_body) // 80)
    print(f"  Abstract found at [{abs_idx}]")
    print(f"  Word count: {abs_words}")
    print(f"  Estimated lines: {est_lines} (ESPR requires ~10-15 lines)")
    if est_lines < 10:
        print(f"  WARNING: Abstract may be too short (< 10 lines)")
    elif est_lines > 20:
        print(f"  WARNING: Abstract may be too long (> 15 lines)")
    else:
        print(f"  OK: Abstract length within range")
    print(f"  Preview: {abs_body[:200]}...")
else:
    print("  WARNING: No abstract heading found")

# ============================================================
# 3. KEYWORDS
# ============================================================
print("\n[3] KEYWORDS")
print("-" * 40)
kw_idx = -1
for i, p in enumerate(paras):
    if "keyword" in p.text.lower() and i < 30:
        kw_idx = i
        break
if kw_idx >= 0:
    kw_text = paras[kw_idx].text
    # Extract keywords after colon
    if ":" in kw_text:
        kw_part = kw_text.split(":", 1)[1].strip()
    else:
        kw_part = paras[kw_idx + 1].text.strip() if kw_idx + 1 < len(paras) else ""
    # Count keywords (separated by ; or , or ·)
    kws = re.split(r"[;,·;]", kw_part)
    kws = [k.strip() for k in kws if k.strip()]
    print(f"  Keywords found: {len(kws)}")
    for j, k in enumerate(kws):
        print(f"    {j+1}. {k}")
    if len(kws) < 6:
        print(f"  WARNING: ESPR requires 6-8 keywords (found {len(kws)})")
    elif len(kws) > 8:
        print(f"  WARNING: ESPR requires 6-8 keywords (found {len(kws)})")
    else:
        print(f"  OK: Keyword count within 6-8 range")
else:
    print("  WARNING: No keywords section found")

# ============================================================
# 4. HEADING LEVELS
# ============================================================
print("\n[4] HEADING LEVELS")
print("-" * 40)
heading_levels = set()
for p in paras:
    if p.style.name.startswith("Heading"):
        level = p.style.name.replace("Heading ", "")
        heading_levels.add(level)
print(f"  Heading levels used: {sorted(heading_levels)}")
max_level = max(int(l) for l in heading_levels) if heading_levels else 0
if max_level > 3:
    print(f"  WARNING: ESPR allows max 3 heading levels (found level {max_level})")
else:
    print(f"  OK: Heading levels within limit (max {max_level})")

# ============================================================
# 5. REFERENCES FORMAT
# ============================================================
print("\n[5] REFERENCES FORMAT")
print("-" * 40)
# Find references section
ref_idx = -1
for i, p in enumerate(paras):
    if p.style.name == "Heading 1" and "reference" in p.text.lower():
        ref_idx = i
        break
if ref_idx >= 0:
    ref_paras = [p.text.strip() for p in paras[ref_idx + 1:] if p.text.strip()]
    print(f"  References section found at [{ref_idx}]")
    print(f"  Total references: {len(ref_paras)}")

    # Check citation style: should be author-date in text
    # Check for [1] style citations (wrong for ESPR)
    bracket_citations = re.findall(r"\[\d+\]", full_text[:ref_idx])
    if bracket_citations:
        print(f"  WARNING: Found {len(bracket_citations)} bracket-number citations [N] in text (ESPR uses author-date)")

    # Check alphabetical order
    first_authors = []
    for r in ref_paras:
        # Extract first author surname
        m = re.match(r"^([A-Z][a-z]+)", r)
        if m:
            first_authors.append(m.group(1))
    is_alpha = first_authors == sorted(first_authors, key=str.lower)
    print(f"  Alphabetical order: {'YES' if is_alpha else 'NO - needs reordering'}")
    if not is_alpha:
        print(f"    First authors: {first_authors[:10]}...")

    # Check for "et al" in reference list (not accepted)
    et_al_refs = [r for r in ref_paras if "et al" in r]
    if et_al_refs:
        print(f"  WARNING: {len(et_al_refs)} references use 'et al' (ESPR requires all author names)")
        for r in et_al_refs[:3]:
            print(f"    -> {r[:100]}")
    else:
        print(f"  OK: No 'et al' in reference list")

    # Check for DOIs
    doi_count = sum(1 for r in ref_paras if "doi" in r.lower() or "https://doi" in r.lower())
    no_doi = len(ref_paras) - doi_count
    print(f"  References with DOI: {doi_count}/{len(ref_paras)}")
    if no_doi > 0:
        print(f"  NOTE: {no_doi} references missing DOI (recommended but not mandatory)")

    # Check for wikipedia references
    wiki_refs = [r for r in ref_paras if "wikipedia" in r.lower()]
    if wiki_refs:
        print(f"  WARNING: Wikipedia references found ({len(wiki_refs)}) - NOT acceptable")
    else:
        print(f"  OK: No Wikipedia references")

    # Check self-citations (max 5 per author)
    # Count how many times each first author appears
    from collections import Counter
    author_counts = Counter(first_authors)
    high_self_cite = [(a, c) for a, c in author_counts.items() if c > 5]
    if high_self_cite:
        print(f"  WARNING: Authors with >5 self-citations: {high_self_cite}")
    else:
        print(f"  OK: No author exceeds 5 self-citations")

    # Show first 5 references for format check
    print(f"\n  First 5 references (format check):")
    for r in ref_paras[:5]:
        print(f"    {r[:120]}")
else:
    print("  WARNING: No references section found")

# ============================================================
# 6. STATEMENTS & DECLARATIONS
# ============================================================
print("\n[6] STATEMENTS & DECLARATIONS")
print("-" * 40)
decl_sections = ["funding", "competing interest", "conflict of interest", "author contribution",
                 "data availability", "acknowledgment"]
for decl in decl_sections:
    found = False
    for p in paras:
        if decl in p.text.lower() and p.style.name.startswith("Heading"):
            found = True
            print(f"  FOUND: {p.text.strip()}")
            break
    if not found:
        # Check in body text
        for p in paras:
            if decl in p.text.lower() and len(p.text) < 200:
                found = True
                print(f"  FOUND (inline): {p.text.strip()[:100]}")
                break
    if not found:
        print(f"  MISSING: {decl.title()} section")

# ============================================================
# 7. TABLES
# ============================================================
print("\n[7] TABLES")
print("-" * 40)
table_mentions = re.findall(r"Table\s+\d+", full_text)
table_nums = set(re.findall(r"Table\s+(\d+)", full_text))
print(f"  Tables referenced in text: {sorted(table_nums, key=int)}")
print(f"  Total table mentions: {len(table_mentions)}")
# Check docx tables
print(f"  Tables in docx: {len(doc.tables)}")

# ============================================================
# 8. FIGURES
# ============================================================
print("\n[8] FIGURES")
print("-" * 40)
fig_mentions = re.findall(r"Fig(?:ure|\.)\s+\d+", full_text)
fig_nums = set(re.findall(r"Fig(?:ure|\.)\s+(\d+)", full_text))
print(f"  Figures referenced in text: {sorted(fig_nums, key=int)}")
print(f"  Total figure mentions: {len(fig_mentions)}")
# Check for embedded images
img_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_count += 1
print(f"  Embedded images in docx: {img_count}")

# ============================================================
# 9. SCOPE CHECK
# ============================================================
print("\n[9] SCOPE CHECK")
print("-" * 40)
scope_issues = []
if "diesel" in full_text.lower() and "engine" in full_text.lower():
    scope_issues.append("Diesel/engine focus (ESPR only accepts emissions papers)")
if "tourism" in full_text.lower():
    scope_issues.append("Tourism content (not accepted by ESPR)")
if "pharmacology" in full_text.lower():
    scope_issues.append("Pharmacology content (not accepted by ESPR)")
if scope_issues:
    for s in scope_issues:
        print(f"  WARNING: {s}")
else:
    print(f"  OK: No out-of-scope topics detected")
    print(f"  Topic: Smart-city construction, carbon emissions, spatial coordination - IN SCOPE")

# ============================================================
# 10. WORD COUNT
# ============================================================
print("\n[10] WORD COUNT")
print("-" * 40)
total_words = sum(len(p.text.split()) for p in paras)
print(f"  Total word count: {total_words}")
print(f"  ESPR has no strict page limit for research articles")
print(f"  Typical research article: ~3000-6000 words (Elsevier guideline)")
print(f"  ESPR: 'length in harmony with research area and science presented'")

# ============================================================
# 11. FONT CHECK
# ============================================================
print("\n[11] FONT CHECK")
print("-" * 40)
fonts_used = set()
for p in paras[:50]:
    for run in p.runs:
        if run.font.name:
            fonts_used.add(run.font.name)
        if run.font.size:
            fonts_used.add(f"{run.font.name}({run.font.size.pt}pt)")
print(f"  Fonts detected (first 50 paras): {fonts_used if fonts_used else 'default/inherited'}")
print(f"  ESPR recommends: 10-point Times Roman")

# ============================================================
# 12. ABBREVIATIONS
# ============================================================
print("\n[12] ABBREVIATIONS (first-mention definition check)")
print("-" * 40)
abbrevs = ["SCCD", "TWFE", "DID", "IV", "2SLS", "DEI", "OIU", "GTI", "POLY", "ER", "NTL", "FDI"]
for abbr in abbrevs:
    # Find first mention
    first_mention = -1
    for i, p in enumerate(paras):
        if abbr in p.text:
            first_mention = i
            break
    if first_mention >= 0:
        text = paras[first_mention].text
        # Check if defined (parenthetical expansion nearby)
        idx = text.find(abbr)
        context = text[max(0, idx-80):idx+len(abbr)+80]
        has_def = bool(re.search(rf"\([^{abbr}]*{abbr}\)|{abbr}\s*\(", context))
        print(f"  {abbr}: first at [{first_mention}] {'DEFINED' if has_def else 'check definition'}")
        if not has_def:
            print(f"    Context: ...{context[:100]}...")

print("\n" + "=" * 70)
print("CHECK COMPLETE")
print("=" * 70)
