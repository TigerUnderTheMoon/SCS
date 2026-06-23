"""Inspect manuscript_espr_revised.docx to locate target text for final cleanup."""
from pathlib import Path
from docx import Document

SRC = Path(r"D:\Workplace\SCS\outputs\espr_submission\manuscript_espr_revised.docx")
doc = Document(str(SRC))

print("=== TOTAL PARAGRAPHS:", len(doc.paragraphs), "===")
print()

# Targets for Task 1 (Sustainability residuals)
task1_targets = [
    "the main Sustainability evidence",
    "keeps the Sustainability message",
    "Sustainability reviewers",
    "For Sustainability readers",
    "aligns with Sustainability's applied policy audience",
]

# Also scan for any other "Sustainability" mentions that look like journal refs
print("=== ALL 'Sustainability' MENTIONS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if "Sustainability" in t:
        # Print index and a snippet around the word
        idx = t.find("Sustainability")
        start = max(0, idx - 60)
        end = min(len(t), idx + 80)
        print(f"[{i:03d}] <{p.style.name}> ...{t[start:end]}...")
print()

print("=== TASK 1 TARGET LOCATIONS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    for target in task1_targets:
        if target in t:
            print(f"[{i:03d}] <{p.style.name}> TARGET: '{target}'")
            print(f"      FULL: {t[:300]}")
            print()
print()

# Task 3: European Commission citation
print("=== TASK 3: European Commission citation ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if "European Commission" in t and ("JointResearch" in t or "Joint Research" in t or "2025" in t):
        print(f"[{i:03d}] <{p.style.name}> {t}")
        print()

# Section 1 end - contribution paragraphs
print("=== SECTION 1 END (contribution paragraphs) ===")
in_sec1 = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if p.style.name == "Heading 1" and t.startswith("1. Introduction"):
        in_sec1 = True
        continue
    if in_sec1 and p.style.name == "Heading 1":
        print(f"--- Section 1 ends before [{i:03d}] <{p.style.name}> {t[:80]} ---")
        break
    if in_sec1:
        # Print last several paragraphs of section 1
        pass

# Find the last 6 paragraphs of section 1
sec1_end_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 1" and p.text.strip().startswith("2. Literature"):
        sec1_end_idx = i
        break
if sec1_end_idx > 0:
    print(f"Section 2 starts at [{sec1_end_idx}]")
    print("Last 6 paragraphs of Section 1:")
    for j in range(max(0, sec1_end_idx - 6), sec1_end_idx):
        p = doc.paragraphs[j]
        wc = len(p.text.split())
        print(f"[{j:03d}] <{p.style.name}> ({wc}w) {p.text[:200]}")
    print()

# Section 5.1
print("=== SECTION 5.1 ===")
i_51 = -1
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 2" and "5.1." in p.text:
        i_51 = i
        break
if i_51 > 0:
    # Find next heading 2
    i_next = -1
    for j in range(i_51 + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == "Heading 2":
            i_next = j
            break
    print(f"Section 5.1 heading at [{i_51}], next H2 at [{i_next}]")
    for j in range(i_51, i_next):
        p = doc.paragraphs[j]
        wc = len(p.text.split())
        print(f"[{j:03d}] <{p.style.name}> ({wc}w) {p.text[:250]}")
    print()

# Section 5.3
print("=== SECTION 5.3 ===")
i_53 = -1
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 2" and "5.3." in p.text:
        i_53 = i
        break
if i_53 > 0:
    i_next = -1
    for j in range(i_53 + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == "Heading 2":
            i_next = j
            break
    print(f"Section 5.3 heading at [{i_53}], next H2 at [{i_next}]")
    for j in range(i_53, i_next):
        p = doc.paragraphs[j]
        wc = len(p.text.split())
        print(f"[{j:03d}] <{p.style.name}> ({wc}w) {p.text[:250]}")
    print()

# Total word count
total_words = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"=== TOTAL WORD COUNT (all paragraphs): {total_words} ===")
