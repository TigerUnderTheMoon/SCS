"""Dump manuscript_espr.docx structure for inspection."""
from docx import Document

doc = Document(r'D:\Workplace\SCS\outputs\espr_submission\manuscript_espr.docx')

with open(r'D:\Workplace\SCS\outputs\espr_submission\_inspect_manuscript.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ''
        text = p.text
        f.write(f'[{i:03d}] <{style}> {text}\n')
    f.write('\n=== TABLES ===\n')
    for ti, t in enumerate(doc.tables):
        f.write(f'\n--- Table {ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---\n')
        for ri, row in enumerate(t.rows):
            cells = [c.text.replace('\n', ' | ') for c in row.cells]
            f.write(f'  R{ri}: ' + ' || '.join(cells) + '\n')

print('done')