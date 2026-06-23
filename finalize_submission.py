"""Regenerate manuscript PDF and run final QA verification."""
import os
import win32com.client
import docx
import re
import json
from docx.oxml.ns import qn

OUTPUT_DIR = r'D:\Workplace\SCS\outputs\espr_submission'
MANUSCRIPT_DOCX = os.path.join(OUTPUT_DIR, 'manuscript_espr.docx')
MANUSCRIPT_PDF = os.path.join(OUTPUT_DIR, 'manuscript_espr.pdf')

# Step 1: Regenerate manuscript PDF
print("=== Step 1: Regenerate manuscript PDF ===")
word = win32com.client.Dispatch('Word.Application')
word.Visible = False
doc = word.Documents.Open(MANUSCRIPT_DOCX)
doc.SaveAs(MANUSCRIPT_PDF, FileFormat=17)
doc.Close()
word.Quit()
print(f"  PDF created: {os.path.getsize(MANUSCRIPT_PDF):,} bytes")

# Step 2: Final QA verification
print("\n=== Step 2: Final QA Verification ===")
doc = docx.Document(MANUSCRIPT_DOCX)

checks = {}

# Check 1: Line numbering
sectPr = doc.sections[0]._sectPr
ln = sectPr.find(qn('w:lnNumType'))
checks['line_numbering'] = ln is not None
print(f"  Line numbering: {'ENABLED' if checks['line_numbering'] else 'NOT SET'}")

# Check 2: Statements and Declarations is Heading 1
sd_heading = False
for p in doc.paragraphs:
    if p.text.strip() == 'Statements and Declarations':
        sd_heading = (p.style and p.style.name == 'Heading 1')
        break
checks['statements_heading'] = sd_heading
print(f"  Statements heading is Heading 1: {checks['statements_heading']}")

# Check 3: References have page ranges
ref_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'References' and p.style and 'Heading' in p.style.name:
        ref_idx = i
        break
has_pages = 0
has_issue = 0
ref_count = 0
for i in range(ref_idx + 1, len(doc.paragraphs)):
    t = doc.paragraphs[i].text.strip()
    if t:
        ref_count += 1
        if re.search(r'\d+\s*[-–]\s*\d+', t) or re.search(r':\s*\d+', t):
            has_pages += 1
        if re.search(r'\d+\(\d+\)', t):
            has_issue += 1
checks['ref_count'] = ref_count
checks['refs_with_pages'] = has_pages
checks['refs_with_issue'] = has_issue
print(f"  References: {ref_count} total, {has_pages} with pages, {has_issue} with issue")

# Check 4: Abstract word count
abstract_wc = 0
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == 'Heading 1' and p.text.strip() == 'Abstract':
        abstract = doc.paragraphs[i + 1].text.strip()
        abstract_wc = len(abstract.split())
        break
checks['abstract_words'] = abstract_wc
print(f"  Abstract: {abstract_wc} words (<=220: {abstract_wc <= 220})")

# Check 5: EDGAR year
edgar_fixed = False
for p in doc.paragraphs:
    if 'European Commission JointResearchCentre' in p.text:
        edgar_fixed = '(2025)' in p.text and '(2026)' not in p.text
        break
checks['edgar_year'] = edgar_fixed
print(f"  EDGAR year 2025: {edgar_fixed}")

# Check 6: No green economy
ge_count = 0
for p in doc.paragraphs:
    ge_count += len(re.findall(r'green economy', p.text, re.IGNORECASE))
checks['green_economy'] = ge_count
print(f"  Green economy occurrences: {ge_count}")

# Check 7: No numeric citations in body (before References)
numeric_cites = 0
for i, p in enumerate(doc.paragraphs):
    if i >= ref_idx:
        break
    numeric_cites += len(re.findall(r'\[(\d+)\]', p.text))
checks['numeric_citations'] = numeric_cites
print(f"  Numeric citations in body: {numeric_cites}")

# Check 8: Cover letter title matches manuscript title
cover_doc = docx.Document(os.path.join(OUTPUT_DIR, 'cover_letter_espr.docx'))
manuscript_title = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Title':
        manuscript_title = p.text.strip()
        break

cover_title_match = False
for p in cover_doc.paragraphs:
    if manuscript_title and manuscript_title in p.text:
        cover_title_match = True
        break
checks['cover_title_match'] = cover_title_match
print(f"  Cover letter title matches manuscript: {cover_title_match}")

# Check 9: Cover letter no mediation overclaim
cover_overclaim = False
for p in cover_doc.paragraphs:
    t = p.text.lower()
    if 'mediation and moderation analyses reveal' in t:
        cover_overclaim = True
        break
checks['cover_no_overclaim'] = not cover_overclaim
print(f"  Cover letter no mediation overclaim: {checks['cover_no_overclaim']}")

# Check 10: ESM files exist
esm_files = {}
for esm in ['ESM_1.pdf', 'ESM_2.pdf', 'ESM_3.pdf']:
    fp = os.path.join(OUTPUT_DIR, esm)
    esm_files[esm] = os.path.exists(fp) and os.path.getsize(fp) > 0
    print(f"  {esm}: {'EXISTS' if esm_files[esm] else 'MISSING'} ({os.path.getsize(fp):,} bytes)" if os.path.exists(fp) else f"  {esm}: MISSING")
checks['esm_files'] = esm_files

# Overall verdict
all_pass = (
    checks['line_numbering'] and
    checks['statements_heading'] and
    checks['refs_with_pages'] >= 30 and
    checks['abstract_words'] <= 220 and
    checks['edgar_year'] and
    checks['green_economy'] == 0 and
    checks['numeric_citations'] == 0 and
    checks['cover_title_match'] and
    checks['cover_no_overclaim'] and
    all(esm_files.values())
)

checks['overall_passed'] = all_pass
checks['manuscript_pdf_size'] = os.path.getsize(MANUSCRIPT_PDF)

print(f"\n=== OVERALL: {'ALL CHECKS PASSED' if all_pass else 'SOME CHECKS FAILED'} ===")

# Save updated QA JSON
qa_path = os.path.join(OUTPUT_DIR, 'espr_final_checks.json')
with open(qa_path, 'w', encoding='utf-8') as f:
    json.dump({
        'manuscript': 'outputs/espr_submission/manuscript_espr.docx',
        'manuscript_pdf': 'outputs/espr_submission/manuscript_espr.pdf',
        'cover_letter': 'outputs/espr_submission/cover_letter_espr.docx',
        'cover_letter_pdf': 'outputs/espr_submission/cover_letter_espr.pdf',
        'esm_files': {
            'ESM_1.pdf': 'outputs/espr_submission/ESM_1.pdf',
            'ESM_2.pdf': 'outputs/espr_submission/ESM_2.pdf',
            'ESM_3.pdf': 'outputs/espr_submission/ESM_3.pdf',
        },
        'checks': checks,
        'overall_passed': all_pass,
        'timestamp': '2026-06-23',
    }, f, indent=2, ensure_ascii=False)
print(f"\nQA JSON saved to {qa_path}")
