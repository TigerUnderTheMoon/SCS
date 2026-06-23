"""Fix ESM_3 - convert figures to PDF using img2pdf or fallback to Word COM."""
import os
import win32com.client
from PyPDF2 import PdfWriter, PdfReader
from PIL import Image
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = r'D:\Workplace\SCS\outputs\espr_submission'
SUPP_DIR = r'D:\Workplace\SCS\outputs\sustainability_restructure_20260603_deep\supplementary_materials'
main_analysis_dir = os.path.join(SUPP_DIR, 'main_analysis')
sensitivity_dir = os.path.join(SUPP_DIR, 'sensitivity_analysis')

ARTICLE_TITLE = "Smart City Construction, Sustainable Urban Transition, and Carbon Emissions: Nonlinear Evidence from Chinese Cities"
AUTHORS = "Haoran Ma, Ningning Wang"
CORR_EMAIL = "wangningning@bistu.edu.cn"
JOURNAL = "Environmental Science and Pollution Research"

# Create title page for ESM_3
def create_title_page_3():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Supplementary Material")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ESM_3: Supplementary Figures")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Article:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ARTICLE_TITLE)
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Authors:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(AUTHORS)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Corresponding author:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CORR_EMAIL)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Journal: {JOURNAL}")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    temp_docx = os.path.join(OUTPUT_DIR, 'esm3_title.docx')
    doc.save(temp_docx)

    title_pdf = os.path.join(OUTPUT_DIR, 'esm3_title.pdf')
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    d = word.Documents.Open(temp_docx)
    d.SaveAs(title_pdf, FileFormat=17)
    d.Close()
    word.Quit()
    os.remove(temp_docx)
    return title_pdf

# Create figures PDF by inserting images into a Word doc then converting
def create_figures_pdf():
    """Insert all figures into a Word document, one per page, then convert to PDF."""
    fig_order = [
        (main_analysis_dir, 'fig_nonlinear_sccd_lnce_observed_2006_2021.png', 'Figure S1. Observed SCCD-lnCE relationship and quadratic fit, 2006-2021'),
        (main_analysis_dir, 'fig_sccd_distribution_observed_2006_2021.png', 'Figure S2. Distribution of SCCD in the observed 2006-2021 sample'),
        (main_analysis_dir, 'fig_moderation_dei_observed_2006_2021.png', 'Figure S3. Predicted SCCD-lnCE curves at low and high DEI levels, observed sample'),
        (sensitivity_dir, 'fig_ce_sccd_quadratic_fit_full_2006_2024.png', 'Figure S4. SCCD-lnCE quadratic fit, full 2006-2024 sample'),
        (sensitivity_dir, 'fig_sccd_distribution_full_2006_2024.png', 'Figure S5. Distribution of SCCD, full 2006-2024 sample'),
        (sensitivity_dir, 'fig_yearly_mean_ce_sccd_full_2006_2024.png', 'Figure S6. Yearly mean CE and SCCD, full 2006-2024 sample'),
        (sensitivity_dir, 'fig_regional_mean_ce_trends_full_2006_2024.png', 'Figure S7. Regional mean CE trends, full 2006-2024 sample'),
        (sensitivity_dir, 'fig_ce_by_region_box_full_2006_2024.png', 'Figure S8. CE by region (box plot), full 2006-2024 sample'),
    ]

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    for i, (dirpath, fname, caption) in enumerate(fig_order):
        fp = os.path.join(dirpath, fname)
        if not os.path.exists(fp):
            print(f"  MISSING: {fname}")
            continue

        print(f"  Adding: {fname}")
        if i > 0:
            doc.add_page_break()

        # Add caption
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

        # Add image
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fp, width=Inches(6.0))

    temp_docx = os.path.join(OUTPUT_DIR, 'esm3_figures.docx')
    doc.save(temp_docx)

    fig_pdf = os.path.join(OUTPUT_DIR, 'esm3_figures.pdf')
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    d = word.Documents.Open(temp_docx)
    d.SaveAs(fig_pdf, FileFormat=17)
    d.Close()
    word.Quit()
    os.remove(temp_docx)
    return fig_pdf

print("=== Creating ESM_3: Supplementary Figures ===")
title_pdf = create_title_page_3()
print(f"  Title page created")

fig_pdf = create_figures_pdf()
print(f"  Figures PDF created: {os.path.getsize(fig_pdf):,} bytes")

# Merge title + figures
esm3_path = os.path.join(OUTPUT_DIR, 'ESM_3.pdf')
writer = PdfWriter()
for pdf in [title_pdf, fig_pdf]:
    reader = PdfReader(pdf)
    for page in reader.pages:
        writer.add_page(page)
with open(esm3_path, 'wb') as f:
    writer.write(f)

# Clean up
os.remove(title_pdf)
os.remove(fig_pdf)

# Verify
reader = PdfReader(esm3_path)
pages = len(reader.pages)
sz = os.path.getsize(esm3_path)
print(f"\n  ESM_3.pdf created: {sz:,} bytes, {pages} pages")

# Final summary of all ESM files
print("\n=== ALL ESM FILES ===")
for esm in ['ESM_1.pdf', 'ESM_2.pdf', 'ESM_3.pdf']:
    fp = os.path.join(OUTPUT_DIR, esm)
    if os.path.exists(fp):
        reader = PdfReader(fp)
        print(f"  {esm}: {os.path.getsize(fp):,} bytes, {len(reader.pages)} pages")
    else:
        print(f"  {esm}: MISSING!")
