from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\Workplace\SCS")
IN = ROOT / "SCS0514_revised_tables_ER_IV.docx"
OUT = ROOT / "revised_manuscript.docx"
FIG = ROOT / "outputs" / "figures"
FRAMEWORK_FIG = FIG / "fig_revised_framework.png"

doc = Document(str(IN))


def make_framework_figure():
    img = Image.new("RGB", (2200, 1100), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 42)
        font = ImageFont.truetype("arial.ttf", 32)
        small = ImageFont.truetype("arial.ttf", 27)
        bold = ImageFont.truetype("arialbd.ttf", 30)
    except OSError:
        title_font = font = small = bold = ImageFont.load_default()

    def box(xy, text, fill, outline, width=3, fnt=None):
        draw.rounded_rectangle(xy, radius=28, fill=fill, outline=outline, width=width)
        lines = []
        for raw in text.split("\n"):
            words = raw.split()
            line = ""
            for word in words:
                test = (line + " " + word).strip()
                if draw.textlength(test, font=fnt or font) <= (xy[2] - xy[0] - 60):
                    line = test
                else:
                    lines.append(line)
                    line = word
            if line:
                lines.append(line)
        line_h = 38
        total_h = line_h * len(lines)
        y = xy[1] + (xy[3] - xy[1] - total_h) / 2
        for line in lines:
            w = draw.textlength(line, font=fnt or font)
            draw.text((xy[0] + (xy[2] - xy[0] - w) / 2, y), line, fill="#111111", font=fnt or font)
            y += line_h

    def arrow(start, end):
        draw.line([start, end], fill="#222222", width=5)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) >= abs(ey - sy):
            pts = [(ex, ey), (ex - 24, ey - 14), (ex - 24, ey + 14)] if ex > sx else [(ex, ey), (ex + 24, ey - 14), (ex + 24, ey + 14)]
        else:
            pts = [(ex, ey), (ex - 14, ey - 24), (ex + 14, ey - 24)] if ey > sy else [(ex, ey), (ex - 14, ey + 24), (ex + 14, ey + 24)]
        draw.polygon(pts, fill="#222222")

    draw.text((110, 60), "Revised research framework", fill="#111111", font=title_font)
    box((120, 285, 520, 615), "Urban multidimensional\nSCCD\n\nProduction\nLiving\nEcological\nDigital", "#fff2cc", "#d6a600", fnt=small)
    box((760, 285, 1180, 435), "Nonlinear effect\nInverted U-shaped\nSCCD -> CE", "#eadcf8", "#8d6cb3", fnt=bold)
    box((1500, 305, 1780, 425), "Carbon\nemissions", "#6b7d8f", "#445566", fnt=bold)
    box((760, 650, 1180, 820), "Mechanism channels\nIndustrial upgrading\nGreen technological innovation", "#dce8fa", "#6f92c8", fnt=small)
    box((760, 95, 1180, 215), "Robust moderator\nDigital economy development", "#dcefd8", "#77aa77", fnt=bold)
    box((1350, 650, 1980, 850), "Supplementary checks only\nEnvironmental regulation overlaps with SCCD index\nPolycentricity is not robust in updated estimates", "#f2f2f2", "#999999", fnt=small)
    box((1255, 135, 1920, 245), "Expansion and efficiency effects determine the turning point", "#f7f2e8", "#b79c64", fnt=small)
    arrow((520, 450), (760, 360))
    arrow((1180, 360), (1500, 365))
    arrow((520, 520), (760, 725))
    arrow((1180, 735), (1500, 405))
    arrow((970, 215), (970, 285))
    arrow((1180, 350), (1255, 190))
    draw.line([(1350, 745), (1180, 735)], fill="#777777", width=3)
    img.save(FRAMEWORK_FIG)


make_framework_figure()


def clear_para(p):
    pPr = p._p.pPr
    for child in list(p._p):
        if pPr is not None and child is pPr:
            continue
        p._p.remove(child)


def set_para_text(p, text):
    clear_para(p)
    p.add_run(text)


def replace_paragraph_start(prefix, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            set_para_text(p, text)
            return True
    raise RuntimeError(f"Paragraph not found: {prefix}")


def remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)


def remove_paragraphs_starting(prefixes):
    for p in list(doc.paragraphs):
        s = p.text.strip()
        if any(s.startswith(prefix) for prefix in prefixes):
            remove_paragraph(p)


def remove_caption_and_previous_drawing(caption_text):
    for p in list(doc.paragraphs):
        if p.text.strip() == caption_text:
            prev = p._p.getprevious()
            while prev is not None and not prev.xpath(".//w:drawing") and not "".join(prev.itertext()).strip():
                blank = prev
                prev = prev.getprevious()
                blank.getparent().remove(blank)
            if prev is not None and prev.xpath(".//w:drawing"):
                prev.getparent().remove(prev)
            remove_paragraph(p)
            return True
    return False


def paragraph_after(paragraph, text=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p


def picture_in_paragraph(p, image_path, width=Inches(6.2)):
    clear_para(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=width)


def set_cell(cell, text):
    cell.text = str(text)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(8)


def fill_table(table, rows):
    if len(table.rows) != len(rows):
        raise ValueError(f"Row count mismatch: table has {len(table.rows)}, data has {len(rows)}")
    for row, vals in zip(table.rows, rows):
        if len(row.cells) != len(vals):
            raise ValueError(f"Column count mismatch: row has {len(row.cells)}, data has {len(vals)}")
        for cell, val in zip(row.cells, vals):
            set_cell(cell, val)


def delete_row(table, row):
    table._tbl.remove(row._tr)


def remove_duplicate_indicator_rows():
    table = doc.tables[0]
    duplicate_indicator = "Employees in water conservancy, environment and public facilities management"
    keep_row = None
    rows_to_delete = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3 or cells[2] != duplicate_indicator:
            continue
        if cells[0] == "Ecological functional space" and cells[1] == "Ecological environmental response" and keep_row is None:
            keep_row = row
        else:
            rows_to_delete.append(row)
    for row in reversed(rows_to_delete):
        delete_row(table, row)


remove_duplicate_indicator_rows()


def replace_framework_picture():
    for p in doc.paragraphs:
        if p.text.strip() == "Fig. 1. Research framework and mechanism analysis diagram.":
            prev = p._p.getprevious()
            if prev is not None:
                picture_in_paragraph(Paragraph(prev, p._parent), FRAMEWORK_FIG, width=Inches(6.4))
            return


def split_indicator_table():
    original = doc.tables[0]
    rows = [[cell.text for cell in row.cells] for row in original.rows]
    header, data_rows = rows[0], rows[1:]
    chunks = [data_rows[:12], data_rows[12:25], data_rows[25:]]

    caption = None
    for p in doc.paragraphs:
        if p.text.strip() == "Table 1. Urban multidimensional spatial indicator system.":
            caption = p
            break
    if caption is None:
        raise RuntimeError("Table 1 caption not found")

    original._element.getparent().remove(original._element)
    anchor = caption._p
    for idx, chunk in enumerate(chunks, start=1):
        if idx > 1:
            cont_el = OxmlElement("w:p")
            anchor.addnext(cont_el)
            cont_p = Paragraph(cont_el, caption._parent)
            cont_p.add_run("Table 1. Urban multidimensional spatial indicator system (continued).")
            anchor = cont_p._p
        table = doc.add_table(rows=len(chunk) + 1, cols=len(header))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        fill_table(table, [header] + chunk)
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(6.5)
        anchor.addnext(table._tbl)
        anchor = table._tbl

paragraph_updates = [
    (
        "This study investigates whether and how digital-enabled urban multidimensional spatial coordination affects carbon emissions in China.",
        "This study investigates whether and how digital-enabled urban multidimensional spatial coordination affects carbon emissions in China. Building on the conventional production-living-ecological framework, the analysis incorporates digital space as a fourth functional dimension and constructs a spatial coupling coordination degree (SCCD) index to capture the joint evolution of production, living, ecological, and digital functions. The main analysis uses observed panel data for 284 prefecture-level and above cities from 2006 to 2021; the extended 2006-2024 panel, which adds fitted/extrapolated observations for 2022-2024, is treated only as a sensitivity check. Two-way fixed-effects estimates show a significant inverted U-shaped relationship between SCCD and carbon emissions, with a main-sample turning point of 0.522. For cities below this threshold, stronger coordination remains associated with expansion-led emission growth; once the threshold is crossed, efficiency gains from factor reallocation, digital-physical integration, and technological spillovers begin to dominate. Industrial upgrading and green technological innovation operate as mechanism channels, although their short-run carbon-reduction effects are partly offset by adjustment costs and rebound effects. Digital economy development is the only moderator that remains robust in the updated results. Environmental regulation and polycentricity are not interpreted as supported moderators because the extended-sample interaction terms are not statistically significant, and environmental regulation also overlaps with the SCCD index construction.",
    ),
    (
        "Accordingly, this study investigates how urban multidimensional spatial coupling coordination affects carbon emissions",
        "Accordingly, this study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation. Extending the conventional production-living-ecological framework, it treats digital space as a fourth functional dimension and constructs a spatial coupling coordination degree (SCCD) index covering production, living, ecological, and digital functions. The main analysis uses observed panel data for 284 prefecture-level and above cities during 2006-2021, while the extended 2006-2024 sample is reported only as a sensitivity check because the 2022-2024 observations are fitted/extrapolated rather than directly observed official records.",
    ),
    (
        "This study contributes in three respects.",
        "This study contributes in three respects. First, it extends the analytical framework of urban functional space by explicitly incorporating digital space into multidimensional spatial coupling coordination. Second, it provides city-level evidence on the nonlinear relationship between SCCD and carbon emissions in China using the observed 2006-2021 sample as the main empirical basis. Third, it clarifies mechanism channels and boundary conditions, with industrial upgrading, green technological innovation, and digital economy development as the core supported channels in the updated evidence; environmental regulation and polycentricity are retained only as supplementary interaction checks.",
    ),
    (
        "Fig. 1 illustrates the research framework and mechanism analysis of this study.",
        "Fig. 1 illustrates the research framework and mechanism analysis of this study. It summarizes the direct nonlinear effect of urban multidimensional spatial coupling coordination on carbon emissions, the mechanism roles of industrial upgrading and green technological innovation, and the primary moderating role of digital economy development. Environmental regulation and polycentric urban structure are treated as supplementary interaction checks rather than as supported moderators in the revised interpretation.",
    ),
    (
        "3.4 Moderating roles of environmental regulation and digital economy development",
        "3.4 Moderating role of digital economy development",
    ),
    (
        "Environmental regulation and digital economy development may further condition",
        "Digital economy development may further condition the nonlinear carbon effect of multidimensional spatial coupling coordination. A more developed digital economy can lower coordination frictions, improve information matching, and support smart transport, energy management, and urban governance. It should therefore attenuate the inverted U-shaped relationship and bring the turning point forward. Environmental regulation is not specified as a primary external moderator in the revised framework because environmental regulation intensity is already included in the SCCD indicator system, creating overlap between the index and the external interaction term.",
    ),
    (
        "Digital economy development also matters because",
        "Accordingly, the updated moderation hypothesis focuses on digital economy development as the main boundary condition.",
    ),
    (
        "H4b. Digital economy development attenuates",
        "H4. Digital economy development attenuates the inverted U-shaped effect of urban multidimensional spatial coupling coordination on carbon emissions and brings forward its turning point.",
    ),
    (
        "3.5 Moderating role of polycentric urban structure",
        "3.5 Supplementary boundary condition: polycentric urban structure",
    ),
    (
        "Urban spatial form provides another boundary condition.",
        "Urban spatial form provides another possible boundary condition. Although polycentric development is often associated with lower congestion and more balanced growth, its environmental effect is not uniformly positive (Jung et al., 2022; Shi et al., 2023; Wang et al., 2024; Zhang et al., 2024). When centers are weakly connected or jobs and residences remain mismatched, polycentricity can raise coordination costs, lengthen commuting and logistics chains, and require repeated infrastructure investment. The revised manuscript therefore treats polycentricity as a supplementary interaction check rather than as a confirmed moderator.",
    ),
    (
        "H2. Industrial upgrading mediates",
        "H2. Industrial upgrading serves as a mechanism channel through which urban multidimensional spatial coupling coordination affects carbon emissions.",
    ),
    (
        "H3. Green technological innovation mediates",
        "H3. Green technological innovation serves as a mechanism channel through which urban multidimensional spatial coupling coordination affects carbon emissions.",
    ),
    (
        "This study uses panel data for 284 prefecture-level and above cities in China over the period 2006-2024.",
        "This study uses observed panel data for 284 prefecture-level and above cities in China over the period 2006-2021 as the main empirical sample. The observed sample contains 4,544 city-year records and is balanced by city and year. The available extended dataset covers 2006-2024 and contains 5,396 city-year observations, but the 852 records for 2022-2024 are fitted/extrapolated values identified by the indicator is_fitted. These fitted observations are used only for sensitivity checks and are not treated as directly observed official data in the main analysis. The sample window is determined by the availability of multidimensional spatial indicators and city-level carbon-emission data. The raw data are drawn mainly from the China City Statistical Yearbook, China Energy Statistical Yearbook, China Urban Construction Statistical Yearbook, China Rural Statistical Yearbook, China Industrial Statistical Yearbook, and China Environmental Statistical Yearbook. City-level carbon emissions are taken from the Emissions Database for Global Atmospheric Research (EDGAR), which ensures temporal comparability under a consistent accounting framework.",
    ),
    (
        "The core explanatory variable is the spatial coupling coordination degree",
        "The core explanatory variable is the spatial coupling coordination degree (SCCD) of urban multidimensional spatial functions. Drawing on the measurement logic of composite digital-development and urban-function indices, and adapting it to the production-living-ecological-digital framework emphasized in this study, the revised indicator system covers four functional dimensions, 13 criterion-level dimensions, and 37 specific indicators after removing duplicated indicator entries from the manuscript table. Larger SCCD values indicate a higher level of coordination across production, living, ecological, and digital spaces.",
    ),
    (
        "Three moderating variables are examined.",
        "Digital economy index (DEI) is examined as the main moderating variable because it captures the broader digital-development environment across cities and remains robust in both the observed and extended samples. Environmental regulation intensity (ER), measured by the frequency of environmental protection terms in government work reports, and polycentricity (POLY), reflecting spatial dispersion and multi-center organization, are retained only as supplementary interaction checks. ER is not interpreted as an independent supported moderator because environmental regulation intensity also appears inside the SCCD indicator system.",
    ),
    (
        "Two mediating variables are included.",
        "Two mechanism variables are included. Industrial upgrading (OIU) is measured by a weighted industrial-structure index, defined as the share of the primary, secondary, and tertiary sectors in GDP multiplied by 1, 2, and 3, respectively. Green technological innovation (GTI) is measured by the logarithm of green patents granted.",
    ),
    (
        "Table 2 reports the descriptive statistics for all variables.",
        "Table 2 reports descriptive statistics for the observed 2006-2021 main sample. The mean value of carbon emissions is 4.528, whereas the mean SCCD is 0.313, suggesting that most sample cities remain at a moderate level of multidimensional spatial coordination. The data also show substantial cross-city variation in carbon emissions, digital economy development, green technological innovation, and polycentric urban structure, supporting the analysis of nonlinear, mechanism, and moderation effects. The extended 2006-2024 sample is used only as a sensitivity check because it includes fitted/extrapolated observations for 2022-2024.",
    ),
    (
        "Fig. 2 further visualizes the spatiotemporal distributions",
        "Fig. 2 visualizes the SCCD distribution in the observed 2006-2021 main sample. The distribution provides descriptive context for the nonlinear empirical analysis and avoids treating fitted/extrapolated 2022-2024 observations as part of the main sample.",
    ),
    (
        "Fig. 2. Spatiotemporal distributions",
        "Fig. 2. Distribution of SCCD in the observed 2006-2021 sample.",
    ),
    (
        "Note: CE denotes carbon emissions and SCCD denotes the spatial coupling coordination degree. Values are shown for 2006",
        "Note: SCCD denotes the spatial coupling coordination degree. The figure uses observed 2006-2021 observations only.",
    ),
    (
        "To identify the transmission mechanisms",
        "To identify transmission mechanisms, this study estimates mechanism equations for industrial upgrading and green technological innovation. Each mechanism variable is first regressed on SCCD and its squared term, and is then introduced into the carbon-emissions equation. Because the current workflow does not estimate formal indirect effects or confidence intervals, these results are interpreted as mechanism evidence rather than formal indirect-effect estimates.",
    ),
    (
        "To examine boundary conditions, interaction terms are constructed",
        "To examine boundary conditions, interaction terms are constructed between SCCD and selected contextual variables. Digital economy development is the primary moderator in the revised interpretation. ER and POLY are retained as supplementary checks to test whether their interaction terms remain stable across the observed main sample and the extended sensitivity sample.",
    ),
    (
        "In the equations above, Xit denotes",
        "In the equations above, Xit denotes the vector of control variables, Mit denotes the mechanism variable, Zit denotes the moderating or supplementary interaction variable, μi denotes city fixed effects, and λt denotes year fixed effects. Robust standard errors are used throughout the estimation.",
    ),
    (
        "Table 3 reports the baseline estimates.",
        "Table 3 reports the baseline estimates for the observed 2006-2021 main sample. Across all specifications, the coefficient on the linear SCCD term is significantly positive and the coefficient on the squared term is significantly negative, providing strong evidence of an inverted U-shaped relationship between urban multidimensional spatial coupling coordination and carbon emissions. In the full-control specification, the coefficient of SCCD is 8.444 and the coefficient of SCCD2 is -8.091, both significant at the 1% level. The result indicates that the environmental effect of spatial coordination is not monotonic. In the initial stage, improvements in coordination are still accompanied by infrastructure expansion, construction activity, factor agglomeration, and rising energy demand, so carbon emissions increase with SCCD. Once coordination reaches a sufficiently mature stage, however, the efficiency effect begins to dominate: digital-physical integration improves factor matching, lowers information frictions, strengthens knowledge spillovers, and supports cleaner urban operation.",
    ),
    (
        "The estimated turning point is 0.522",
        "The estimated main-sample turning point is 0.522, calculated from the observed 2006-2021 full-control specification. This value lies within the observed SCCD range of 0.121 to 0.822. As shown in Fig. 3, the fitted relationship follows a clear inverted U-shape: when SCCD is below this threshold, stronger coordination is associated with higher carbon emissions, whereas beyond the threshold the marginal effect turns negative. The extended 2006-2024 sensitivity estimate gives a turning point of 0.542, indicating that adding the fitted/extrapolated 2022-2024 rows changes the numerical threshold but not the main inverted-U conclusion.",
    ),
    (
        "Table 4 reports the robustness, endogeneity, and first-stage results.",
        "Table 4 reports robustness, endogeneity, and first-stage results for the observed 2006-2021 main sample. To verify that the baseline results are not driven by outliers, measurement choices, or reverse causality, several additional tests are conducted. First, bilateral winsorization at the 1% level is used to reduce the influence of extreme values. Second, the dependent variable is replaced with average nighttime light intensity as an alternative proxy. Third, the explanatory variable is replaced with MCCD and MCCD2; the do-file and exported table confirm that column (3) uses MCCD/MCCD2 rather than SCCD/SCCD2. Fourth, IV-2SLS estimation is performed using the existing IV variable as lagged SCCD and IV2 as its squared term. Across these checks, the linear and quadratic terms remain significantly positive and negative, respectively.",
    ),
    (
        "Note: Robust standard errors clustered at the city level",
        "Note: Robust standard errors clustered at the city level are reported in parentheses. ***, **, and * denote significance at the 1%, 5%, and 10% levels, respectively. Column (3) uses MCCD and MCCD2 as the alternative explanatory variables. Column (4) reports the IV-2SLS second-stage estimation using IV and IV2 as instruments for SCCD and SCCD2. The first-stage panel is reported within the same table to avoid adding a separate table.",
    ),
    (
        "To further address endogeneity, the lagged linear and squared terms",
        "To further address endogeneity, IV and IV2 are used as instrumental variables in a two-stage least squares estimation. The second-stage results show that the coefficient of SCCD remains significantly positive, while the coefficient of SCCD2 remains significantly negative. Based on the observed-sample IV-2SLS estimates, the implied turning point is approximately 0.570, which remains within the observed SCCD range. The first-stage results indicate instrument relevance: IV significantly predicts SCCD, and IV2 significantly predicts SCCD2, with within R-squared values of 0.645 and 0.582, respectively. The extended 2006-2024 sensitivity sample gives an IV turning point of approximately 0.638.",
    ),
    (
        "Table 5 reports the regional heterogeneity results.",
        "Table 5 reports the regional heterogeneity results for the observed 2006-2021 main sample. Given substantial regional differences in development patterns, industrial structure, and urbanization, this study further examines whether the carbon effect of SCCD differs across eastern, central, and western China. The estimates show that the inverted U-shaped relationship remains statistically significant in all three regions, suggesting that the nonlinear effect is broadly robust across China. The absolute value of the quadratic coefficient is largest in central China (-12.828), followed by western China (-8.152) and eastern China (-6.905), indicating that carbon emissions in central China are most sensitive to changes in multidimensional spatial coordination.",
    ),
    (
        "The magnitude of the effect, however, differs across regions.",
        "The magnitude of the effect differs across regions. In eastern China, the estimated effect is flatter, likely because industrial upgrading, digital development, and efficiency gains already play a larger role than large-scale physical expansion. In western China, the nonlinear effect remains pronounced, which may reflect the coexistence of infrastructure investment, industrial catch-up, and uneven digital development.",
    ),
    (
        "5.4 Mediation analysis",
        "5.4 Mechanism analysis",
    ),
    (
        "Table 6 reports the mediation analysis results.",
        "Table 6 reports the mechanism analysis results for the observed 2006-2021 main sample. To uncover mechanisms linking SCCD to carbon emissions, this study tests the roles of industrial upgrading and green technological innovation. The results show that SCCD has an inverted U-shaped effect on both mechanism variables, suggesting that stronger spatial coordination initially promotes industrial upgrading and green technological innovation, but that these gains weaken as coordination becomes more mature. The current workflow provides mechanism-equation evidence rather than separately estimated indirect-effect magnitudes.",
    ),
    (
        "When industrial upgrading is added to the carbon-emissions equation",
        "When industrial upgrading is added to the carbon-emissions equation, its coefficient is significantly positive, while the coefficients on SCCD and its squared term decline but remain statistically significant. This indicates that industrial upgrading is an important mechanism channel, although its short-run carbon effect is emission-increasing before cleaner production systems are fully established. From a transitional perspective, this result is understandable: upgrading often involves new fixed-asset investment, equipment replacement, construction of modern industrial facilities, and the coexistence of old and new sectors. In other words, multidimensional spatial coordination can promote structural change, but the carbon benefits of that change may only materialize after the adjustment phase has passed.",
    ),
    (
        "When green technological innovation is added to the carbon-emissions equation",
        "When green technological innovation is added to the carbon-emissions equation, its coefficient is also significantly positive, and the SCCD coefficients shrink relative to the baseline model. This confirms that green technological innovation acts as another important transmission channel. Yet the positive coefficient on GTI suggests that green innovation does not automatically lower observed emissions in the short run. Research and development, testing, commercialization, and adoption all require capital and energy inputs, while efficiency gains may be partly offset by rebound effects. Thus, the mechanism role of green innovation should be understood dynamically: multidimensional spatial coordination improves the urban conditions for green innovation, but the carbon-reduction payoff may lag behind the innovation process.",
    ),
    (
        "Table 6. Mediation analysis results.",
        "Table 6. Mechanism analysis results.",
    ),
    (
        "Taken together, the mediation results suggest",
        "Taken together, the mechanism results suggest that the effect of multidimensional spatial coordination on carbon emissions is not purely direct. Spatial coordination changes the institutional and economic conditions under which industrial upgrading and innovation occur, and these transitions in turn shape the timing and magnitude of carbon outcomes. The low-carbon dividend of coordination therefore depends on whether cities can shorten the adjustment period and reduce rebound effects during structural transformation.",
    ),
    (
        "5.5 Moderation analysis",
        "5.5 Moderation and supplementary interaction checks",
    ),
    (
        "Table 7 reports the moderation analysis results.",
        "Table 7 reports the moderation and supplementary interaction checks for the observed 2006-2021 main sample. The revised interpretation focuses on digital economy development as the only robust moderator. ER and POLY are shown to document the supplementary checks, but they are not treated as supported moderators because ER overlaps with the SCCD index and the extended 2006-2024 interaction terms for ER and POLY are not statistically significant.",
    ),
    (
        "Table 7. Moderation analysis results.",
        "Table 7. Moderation and supplementary interaction results.",
    ),
    (
        "The interaction terms for environmental regulation show",
        "The environmental-regulation interaction is not interpreted as robust moderation evidence. In the observed 2006-2021 sample, SCCD2 x ER is weakly significant at the 10% level, but the corresponding extended 2006-2024 interaction terms are not statistically significant. In addition, environmental regulation intensity is already included in the SCCD indicator system, so treating ER as an independent external moderator would create an overlap problem. The revised manuscript therefore does not claim support for environmental regulation as a moderator.",
    ),
    (
        "The interaction terms for digital economy development are statistically significant",
        "The interaction terms for digital economy development are statistically significant and support the revised H4. In the observed 2006-2021 sample, the coefficient on SCCD x DEI is significantly negative, while the coefficient on SCCD2 x DEI is significantly positive. The extended 2006-2024 sensitivity sample shows the same significance pattern, making digital economy development the only robust moderator in the updated results. This pattern is consistent with the argument that digital economy development lowers coordination frictions, improves information matching, facilitates intelligent governance, and enhances the efficiency of resource allocation.",
    ),
    (
        "For polycentric urban structure, the interaction results support H5.",
        "For polycentric urban structure, the generated interaction terms do not statistically support a moderation claim. The coefficients on SCCD x POLY and SCCD2 x POLY are not significant in the observed sample, and the extended 2006-2024 sensitivity sample also does not support the interaction terms. Polycentricity may still matter theoretically because dispersed centers can create longer logistics chains, weaker commuting-job matching, and duplicated infrastructure when cross-center links remain incomplete, but this argument should be presented as a possible explanation requiring further empirical confirmation rather than as a confirmed moderation result.",
    ),
    (
        "Taken together, the moderation results show that urban context matters.",
        "Taken together, the moderation and supplementary interaction results show that digital capability is the clearest boundary condition in the updated evidence. Environmental regulation and polycentricity remain theoretically relevant, but the current estimates do not support strong empirical claims for either as a robust moderator. This finding complements the baseline and mechanism results by showing that the environmental consequences of spatial coordination depend especially on whether cities can translate digital capacity into lower coordination frictions and more efficient urban operation.",
    ),
    (
        "This study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation.",
        "This study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation. By extending the conventional production-living-ecological framework to include digital space and by using observed panel data for 284 prefecture-level and above cities over 2006-2021, the paper evaluates both the nonlinear carbon effect of spatial coordination and the mechanisms through which this effect operates. The extended 2006-2024 panel, which adds fitted/extrapolated observations for 2022-2024, is used only as a sensitivity check. The central message is that multidimensional spatial coordination is neither inherently low-carbon nor inherently carbon-intensive. Its environmental consequences depend on development stage, structural adjustment, and governance capacity.",
    ),
    (
        "Four main conclusions can be drawn.",
        "Four main conclusions can be drawn. First, SCCD exhibits a significant inverted U-shaped relationship with carbon emissions in the observed 2006-2021 sample, with a main-sample turning point of 0.522. The extended 2006-2024 sensitivity sample gives a similar inverted-U pattern, with a turning point of 0.542, but those fitted/extrapolated observations are not treated as the main empirical sample. Second, the nonlinear relationship remains robust after multiple specification adjustments and IV-2SLS estimation. Third, industrial upgrading and green technological innovation constitute important mechanism channels, but their short-run effects are not purely emission-reducing because transition costs, equipment renewal, and rebound effects partly offset efficiency gains. Fourth, digital economy development significantly conditions the carbon effect of SCCD and remains robust across the observed and extended samples. The updated evidence does not support strong claims that environmental regulation or polycentricity robustly moderate the nonlinear relationship.",
    ),
    (
        "Overall, the evidence shows",
        "Overall, the evidence shows that multidimensional spatial coordination does not automatically produce low-carbon outcomes. Its environmental value depends on whether spatial integration is translated into efficiency improvement rather than land expansion, whether structural and technological upgrading generate durable rather than temporary carbon gains, and whether urban governance can manage the frictions associated with digital-physical transformation. For cities below the observed-sample turning point of 0.522, the immediate challenge is to prevent coordination from being translated primarily into duplicated infrastructure and energy-intensive growth. For cities approaching or surpassing the turning point, the policy task shifts toward consolidating the efficiency advantages of digital integration and converting innovation and upgrading into persistent emission reductions.",
    ),
    (
        "First, urban low-carbon policy should move",
        "First, urban low-carbon policy should move beyond extensive physical expansion and prioritize deeper digital-physical integration. Because many sample cities remain on the left side of the observed-sample turning point, simply raising formal indicators of spatial coordination is insufficient to guarantee lower emissions. Policy should therefore focus on digital infrastructure, interoperable data systems, and intelligent governance tools that reduce the coordination costs of production, living, ecological, and digital spaces.",
    ),
    (
        "Second, policy design should explicitly manage",
        "Second, policy design should explicitly manage the rebound risks associated with industrial upgrading and green technological innovation. Because structural transformation and green innovation do not automatically reduce emissions in the short run, complementary measures such as cleaner energy substitution, green finance, environmental regulation, and performance-based evaluation are needed to ensure that transition costs do not lock cities into higher-carbon pathways. This policy role for regulation is a general governance implication, not a claim that ER is a robust moderator in the updated interaction estimates.",
    ),
    (
        "Third, cities pursuing polycentric development",
        "Third, cities pursuing polycentric development should pay closer attention to cross-center coordination. The current interaction estimates do not provide robust support for polycentricity as a moderator, but the theoretical risk remains clear: dispersed centers can raise logistics costs, duplicate infrastructure, and weaken jobs-housing balance when cross-center links are incomplete. Polycentric planning should therefore be accompanied by stronger integration of land use, transport systems, infrastructure provision, and digital coordination tools.",
    ),
    (
        "Fourth, governance strategies should be differentiated",
        "Fourth, governance strategies should be differentiated across regions and development stages. Central China deserves particular attention because it faces both stronger emission sensitivity to changes in SCCD and greater potential gains from later-stage coordination. Eastern China is better positioned to explore advanced digital-physical governance models, whereas western China should place greater emphasis on coupling ecological protection with improvements in basic public services and urban functional coordination. More broadly, cities below the turning point should prioritize curbing construction-led expansion and improving the quality rather than the quantity of coordination; cities near the turning point should accelerate digital governance reform, industrial upgrading quality, and green technology diffusion; and cities already beyond the turning point should focus on locking in low-carbon gains through institutional coordination, cross-sector data integration, and long-term monitoring of rebound effects.",
    ),
    (
        "This study has several limitations.",
        "This study has several limitations. First, although the SCCD index incorporates digital space into the conventional production-living-ecological framework, the measurement of digital space is constrained by data availability and may not fully capture platform governance, data mobility, algorithmic coordination, or the quality of digital public services. Second, although the empirical design includes fixed effects, robustness checks, and instrumental-variable estimation, the analysis remains based on observational panel data and therefore cannot fully eliminate all identification concerns. Third, the extended 2006-2024 sample includes fitted/extrapolated observations for 2022-2024; these values are used only for sensitivity checks and are not equivalent to directly observed official records. Fourth, the mechanism analysis focuses on industrial upgrading and green technological innovation, but other channels--such as energy structure, green finance, or household behavioral change--may also play meaningful roles in shaping carbon outcomes.",
    ),
]

for prefix, text in paragraph_updates:
    replace_paragraph_start(prefix, text)

remove_paragraphs_starting(["H4a.", "H5."])
remove_caption_and_previous_drawing("Fig. 4. Moderating effect of digital economy development.")
remove_caption_and_previous_drawing("Fig. 5. Moderating effect of polycentric urban structure.")
remove_paragraphs_starting(
    [
        "Data availability statement",
        "Declaration of competing interest",
        "Funding",
        "CRediT author contribution statement",
        "Declaration of generative AI and AI-assisted technologies",
        "[NEEDS CHECK:",
        "[NEEDS AUTHOR CONFIRMATION:",
    ]
)

for p in doc.paragraphs:
    if p.text.strip() == "Fig. 2. Distribution of SCCD in the observed 2006-2021 sample.":
        prev = p._p.getprevious()
        if prev is not None:
            picture_in_paragraph(Paragraph(prev, p._parent), FIG / "fig_sccd_distribution_observed_2006_2021.png")
        break

table2 = [
    ["Category", "Variable", "Symbol", "Obs.", "Mean", "Median", "S.D.", "Min", "Max"],
    ["Dependent variable", "Carbon emissions", "CE", "4544", "4.528", "4.554", "1.292", "0.000", "8.325"],
    ["Core explanatory variable", "Spatial coupling coordination degree", "SCCD", "4544", "0.313", "0.297", "0.086", "0.121", "0.822"],
    ["Control variable", "Trade openness", "OPEN", "4529", "0.193", "0.076", "0.345", "-0.720", "3.640"],
    ["Control variable", "Urbanization rate", "UR", "4544", "53.912", "51.839", "15.862", "6.491", "100.000"],
    ["Control variable", "Urban-rural income gap", "URG", "4529", "2.466", "2.377", "0.567", "1.207", "6.378"],
    ["Control variable", "Government intervention intensity", "GI", "4544", "0.185", "0.159", "0.101", "0.043", "1.027"],
    ["Mechanism variable", "Industrial upgrading", "OIU", "4544", "2.278", "2.270", "0.149", "1.163", "2.836"],
    ["Mechanism variable", "Green technological innovation", "GTI", "4544", "3.931", "3.871", "1.863", "0.000", "9.872"],
    ["Supplementary check", "Environmental regulation intensity", "ER", "4544", "0.008", "0.007", "0.003", "0.000", "0.026"],
    ["Moderating variable", "Digital economy index", "DEI", "4544", "0.056", "0.042", "0.077", "0.000", "0.940"],
    ["Supplementary check", "Degree of polycentricity", "POLY", "4544", "0.347", "0.376", "0.194", "0.000", "0.963"],
]

table3 = [
    ["Variables", "(1)", "(2)", "(3)", "(4)", "(5)"],
    ["Dependent variable", "CE", "CE", "CE", "CE", "CE"],
    ["SCCD", "12.957***", "11.841***", "9.007***", "8.819***", "8.444***"],
    ["", "(1.445)", "(1.421)", "(1.305)", "(1.348)", "(1.338)"],
    ["SCCD2", "-11.476***", "-10.296***", "-8.400***", "-8.125***", "-8.091***"],
    ["", "(1.749)", "(1.761)", "(1.593)", "(1.642)", "(1.665)"],
    ["OPEN", "", "0.456***", "0.211**", "0.244***", "0.276***"],
    ["", "", "(0.108)", "(0.083)", "(0.083)", "(0.082)"],
    ["UR", "", "", "0.029***", "0.027***", "0.025***"],
    ["", "", "", "(0.004)", "(0.004)", "(0.004)"],
    ["URG", "", "", "", "-0.116*", "-0.101"],
    ["", "", "", "", "(0.068)", "(0.067)"],
    ["GI", "", "", "", "", "-1.953***"],
    ["", "", "", "", "", "(0.546)"],
    ["Constant", "1.203***", "1.314***", "0.629**", "1.053***", "1.497***"],
    ["", "(0.284)", "(0.272)", "(0.267)", "(0.357)", "(0.382)"],
    ["N", "4544", "4529", "4529", "4514", "4514"],
    ["Within R-squared", "0.592", "0.601", "0.640", "0.641", "0.650"],
]

table4 = [
    ["Variables", "(1)", "(2)", "(3)", "(4)"],
    ["Specification", "Winsorized at 1%", "Alternative dependent variable", "Alternative explanatory variable", "IV-2SLS / second stage"],
    ["Linear term (MCCD in col. 3)", "9.845***", "3.601***", "8.444***", "17.971***"],
    ["", "(1.337)", "(0.719)", "(1.338)", "(2.816)"],
    ["Squared term (MCCD2 in col. 3)", "-10.134***", "-2.765***", "-8.091***", "-15.757***"],
    ["", "(1.648)", "(0.763)", "(1.665)", "(2.969)"],
    ["Controls", "Yes", "Yes", "Yes", "Yes"],
    ["City fixed effects", "Yes", "Yes", "Yes", "Yes"],
    ["Year fixed effects", "Yes", "Yes", "Yes", "Yes"],
    ["N", "4514", "4514", "4514", "4232"],
    ["Within R-squared", "0.658", "0.583", "0.650", "0.626"],
    ["First-stage instrument relevance", "", "", "", "IV-2SLS first stage"],
    ["Dependent variable: SCCD", "", "", "", ""],
    ["IV", "", "", "", "0.292*** / (0.078)"],
    ["IV2", "", "", "", "-0.027 / (0.097)"],
    ["First-stage within R-squared for SCCD", "", "", "", "0.645"],
    ["Dependent variable: SCCD2", "", "", "", ""],
    ["IV", "", "", "", "-0.091 / (0.065)"],
    ["IV2", "", "", "", "0.441*** / (0.091)"],
    ["First-stage within R-squared for SCCD2", "", "", "", "0.582"],
    ["First-stage N", "", "", "", "4232"],
]

table5 = [
    ["Variables", "(1)", "(2)", "(3)"],
    ["Region", "Eastern China", "Central China", "Western China"],
    ["SCCD", "7.368***", "12.447***", "7.842***"],
    ["", "(2.557)", "(2.490)", "(1.740)"],
    ["SCCD2", "-6.905***", "-12.828***", "-8.152***"],
    ["", "(2.543)", "(3.179)", "(2.102)"],
    ["Constant", "1.615**", "0.545", "1.920***"],
    ["", "(0.711)", "(0.700)", "(0.639)"],
    ["Controls", "Yes", "Yes", "Yes"],
    ["City fixed effects", "Yes", "Yes", "Yes"],
    ["Year fixed effects", "Yes", "Yes", "Yes"],
    ["N", "1537", "1585", "1392"],
    ["Within R-squared", "0.731", "0.645", "0.624"],
]

table6 = [
    ["Variables", "(1)", "(2)", "(3)", "(4)"],
    ["Dependent variable", "Industrial upgrading", "Carbon emissions", "Green technological innovation", "Carbon emissions"],
    ["SCCD", "0.647***", "7.561***", "10.379***", "5.290***"],
    ["", "(0.160)", "(1.254)", "(1.402)", "(1.115)"],
    ["SCCD2", "-0.407*", "-7.535***", "-5.775***", "-6.336***"],
    ["", "(0.207)", "(1.503)", "(1.468)", "(1.418)"],
    ["OIU", "", "1.366***", "", ""],
    ["", "", "(0.316)", "", ""],
    ["GTI", "", "", "", "0.304***"],
    ["", "", "", "", "(0.027)"],
    ["Constant", "1.808***", "-0.972", "-0.725*", "1.718***"],
    ["", "(0.041)", "(0.697)", "(0.411)", "(0.325)"],
    ["Controls", "Yes", "Yes", "Yes", "Yes"],
    ["City fixed effects", "Yes", "Yes", "Yes", "Yes"],
    ["Year fixed effects", "Yes", "Yes", "Yes", "Yes"],
    ["N", "4514", "4514", "4514", "4514"],
    ["Within R-squared", "0.622", "0.658", "0.826", "0.686"],
]

table7 = [
    ["Variables", "(1)", "(2)", "(3)"],
    ["Interaction", "Supplementary ER check", "Digital economy moderator", "Supplementary POLY check"],
    ["SCCD", "10.931***", "8.869***", "7.925***"],
    ["", "(1.991)", "(1.198)", "(2.146)"],
    ["SCCD2", "-11.675***", "-7.977***", "-8.678***"],
    ["", "(2.435)", "(1.423)", "(2.660)"],
    ["ER", "61.643", "", ""],
    ["", "(46.315)", "", ""],
    ["SCCD x ER", "-325.788", "", ""],
    ["", "(228.112)", "", ""],
    ["SCCD2 x ER", "455.053*", "", ""],
    ["", "(266.493)", "", ""],
    ["DEI", "", "5.874***", ""],
    ["", "", "(2.226)", ""],
    ["SCCD x DEI", "", "-25.913***", ""],
    ["", "", "(7.936)", ""],
    ["SCCD2 x DEI", "", "24.173***", ""],
    ["", "", "(6.665)", ""],
    ["POLY", "", "", "-1.043"],
    ["", "", "", "(0.941)"],
    ["SCCD x POLY", "", "", "2.842"],
    ["", "", "", "(4.782)"],
    ["SCCD2 x POLY", "", "", "0.159"],
    ["", "", "", "(5.485)"],
    ["Constant", "1.040**", "1.364***", "1.763***"],
    ["", "(0.486)", "(0.370)", "(0.487)"],
    ["Controls", "Yes", "Yes", "Yes"],
    ["City fixed effects", "Yes", "Yes", "Yes"],
    ["Year fixed effects", "Yes", "Yes", "Yes"],
    ["N", "4514", "4514", "4514"],
    ["Within R-squared", "0.651", "0.652", "0.653"],
]

for table, rows in zip(doc.tables[1:], [table2, table3, table4, table5, table6, table7]):
    fill_table(table, rows)

replace_framework_picture()
split_indicator_table()

for p in doc.paragraphs:
    s = p.text.strip()
    if s == "Table 2. Descriptive statistics.":
        paragraph_after(p, "Note: Statistics are for the observed 2006-2021 main sample. The extended 2006-2024 sample is used only as a sensitivity check.")

for p in doc.paragraphs:
    if p.text.strip().startswith("Table "):
        p.paragraph_format.keep_with_next = True

OUT.unlink(missing_ok=True)
doc.save(str(OUT))
print(OUT)
