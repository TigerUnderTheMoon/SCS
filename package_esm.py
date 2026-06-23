"""Package ESM supplementary PDFs for ESPR submission."""
import os
import win32com.client
from PyPDF2 import PdfWriter, PdfReader
from PIL import Image
import io
import time

OUTPUT_DIR = r'D:\Workplace\SCS\outputs\espr_submission'
SUPP_DIR = r'D:\Workplace\SCS\outputs\sustainability_restructure_20260603_deep\supplementary_materials'

ARTICLE_TITLE = "Smart City Construction, Sustainable Urban Transition, and Carbon Emissions: Nonlinear Evidence from Chinese Cities"
AUTHORS = "Haoran Ma, Ningning Wang"
CORR_EMAIL = "wangningning@bistu.edu.cn"
JOURNAL = "Environmental Science and Pollution Research"

# Create a title page PDF using Word COM
def create_title_page(esm_num, esm_title, output_pdf):
    """Create a title page PDF for an ESM file."""
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # "Supplementary Material" heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Supplementary Material")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'

    # ESM number
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"ESM_{esm_num}: {esm_title}")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    # Article title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Article:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ARTICLE_TITLE)
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Authors:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(AUTHORS)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Corresponding author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Corresponding author:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CORR_EMAIL)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Journal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Journal: {JOURNAL}")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Save as docx first, then convert to PDF
    temp_docx = output_pdf.replace('.pdf', '_titlepage.docx')
    doc.save(temp_docx)

    # Convert to PDF via Word COM
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    d = word.Documents.Open(temp_docx)
    d.SaveAs(output_pdf, FileFormat=17)
    d.Close()
    word.Quit()

    # Clean up temp docx
    os.remove(temp_docx)
    return output_pdf

def rtf_to_pdf(rtf_path, pdf_path):
    """Convert RTF to PDF using Word COM."""
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    d = word.Documents.Open(rtf_path)
    d.SaveAs(pdf_path, FileFormat=17)
    d.Close()
    word.Quit()
    return pdf_path

def merge_pdfs(pdf_list, output_path):
    """Merge multiple PDFs into one."""
    writer = PdfWriter()
    for pdf in pdf_list:
        if os.path.exists(pdf) and os.path.getsize(pdf) > 0:
            reader = PdfReader(pdf)
            for page in reader.pages:
                writer.add_page(page)
    with open(output_path, 'wb') as f:
        writer.write(f)
    return output_path

def images_to_pdf(image_paths, output_path):
    """Convert PNG images to a single PDF."""
    images = []
    for img_path in image_paths:
        if os.path.exists(img_path):
            img = Image.open(img_path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            images.append(img)

    if images:
        images[0].save(output_path, save_all=True, append_images=images[1:], resolution=150.0)
    return output_path

# === ESM_1: Main analysis tables ===
print("=== Creating ESM_1: Main Analysis Tables ===")
main_analysis_dir = os.path.join(SUPP_DIR, 'main_analysis')
table_order_1 = [
    'table02_descriptive_statistics_observed_2006_2021.rtf',
    'table03_baseline_observed_2006_2021.rtf',
    'table04_robustness_endogeneity_observed_2006_2021.rtf',
    'table04_first_stage_observed_2006_2021.rtf',
    'table05_regional_heterogeneity_observed_2006_2021.rtf',
    'table06_mediation_observed_2006_2021.rtf',
    'table07_moderation_observed_2006_2021.rtf',
    'table08_vif_observed_2006_2021.rtf',  # This is actually a txt, skip
    'table08_fixed_effects_joint_tests_observed_2006_2021.rtf',  # Also txt, skip
]

# Filter to only .rtf files that exist
rtf_files_1 = []
for f in table_order_1:
    fp = os.path.join(main_analysis_dir, f)
    if f.endswith('.rtf') and os.path.exists(fp):
        rtf_files_1.append(fp)

# Add table4_iv_diagnostics if exists as rtf - it's txt only
print(f"  Found {len(rtf_files_1)} RTF table files for ESM_1")

# Create title page
title_pdf_1 = os.path.join(OUTPUT_DIR, 'esm1_title.pdf')
create_title_page(1, "Main Analysis Tables (Observed 2006-2021 Sample)", title_pdf_1)
print(f"  Title page created")

# Convert each RTF to PDF
table_pdfs_1 = [title_pdf_1]
for i, rtf in enumerate(rtf_files_1):
    pdf_out = rtf.replace('.rtf', '.pdf')
    print(f"  Converting {os.path.basename(rtf)}...")
    try:
        rtf_to_pdf(rtf, pdf_out)
        table_pdfs_1.append(pdf_out)
    except Exception as e:
        print(f"    ERROR: {e}")

# Merge
esm1_path = os.path.join(OUTPUT_DIR, 'ESM_1.pdf')
merge_pdfs(table_pdfs_1, esm1_path)
print(f"  ESM_1.pdf created: {os.path.getsize(esm1_path)} bytes")

# Clean up temp PDFs
for p in table_pdfs_1:
    if p != esm1_path and os.path.exists(p):
        os.remove(p)

# === ESM_2: Sensitivity analysis tables ===
print("\n=== Creating ESM_2: Sensitivity Analysis Tables ===")
sensitivity_dir = os.path.join(SUPP_DIR, 'sensitivity_analysis')
table_order_2 = [
    'table02_descriptive_statistics_full_2006_2024.rtf',
    'table03_baseline_full_2006_2024.rtf',
    'table04_robustness_endogeneity_full_2006_2024.rtf',
    'table04_first_stage_full_2006_2024.rtf',
    'table05_regional_heterogeneity_full_2006_2024.rtf',
    'table06_mediation_full_2006_2024.rtf',
    'table07_moderation_full_2006_2024.rtf',
]

rtf_files_2 = []
for f in table_order_2:
    fp = os.path.join(sensitivity_dir, f)
    if f.endswith('.rtf') and os.path.exists(fp):
        rtf_files_2.append(fp)

print(f"  Found {len(rtf_files_2)} RTF table files for ESM_2")

title_pdf_2 = os.path.join(OUTPUT_DIR, 'esm2_title.pdf')
create_title_page(2, "Sensitivity Analysis Tables (Full 2006-2024 Sample)", title_pdf_2)
print(f"  Title page created")

table_pdfs_2 = [title_pdf_2]
for i, rtf in enumerate(rtf_files_2):
    pdf_out = rtf.replace('.rtf', '.pdf')
    print(f"  Converting {os.path.basename(rtf)}...")
    try:
        rtf_to_pdf(rtf, pdf_out)
        table_pdfs_2.append(pdf_out)
    except Exception as e:
        print(f"    ERROR: {e}")

esm2_path = os.path.join(OUTPUT_DIR, 'ESM_2.pdf')
merge_pdfs(table_pdfs_2, esm2_path)
print(f"  ESM_2.pdf created: {os.path.getsize(esm2_path)} bytes")

for p in table_pdfs_2:
    if p != esm2_path and os.path.exists(p):
        os.remove(p)

# === ESM_3: Supplementary figures ===
print("\n=== Creating ESM_3: Supplementary Figures ===")
fig_order_3 = [
    (main_analysis_dir, 'fig_nonlinear_sccd_lnce_observed_2006_2021.png'),
    (main_analysis_dir, 'fig_sccd_distribution_observed_2006_2021.png'),
    (main_analysis_dir, 'fig_moderation_dei_observed_2006_2021.png'),
    (sensitivity_dir, 'fig_ce_sccd_quadratic_fit_full_2006_2024.png'),
    (sensitivity_dir, 'fig_sccd_distribution_full_2006_2024.png'),
    (sensitivity_dir, 'fig_yearly_mean_ce_sccd_full_2006_2024.png'),
    (sensitivity_dir, 'fig_regional_mean_ce_trends_full_2006_2024.png'),
    (sensitivity_dir, 'fig_ce_by_region_box_full_2006_2024.png'),
]

fig_paths_3 = []
for dirpath, fname in fig_order_3:
    fp = os.path.join(dirpath, fname)
    if os.path.exists(fp):
        fig_paths_3.append(fp)
        print(f"  Found: {fname}")
    else:
        print(f"  MISSING: {fname}")

title_pdf_3 = os.path.join(OUTPUT_DIR, 'esm3_title.pdf')
create_title_page(3, "Supplementary Figures", title_pdf_3)
print(f"  Title page created")

# Convert figures to PDF
fig_pdf = os.path.join(OUTPUT_DIR, 'esm3_figures.pdf')
images_to_pdf(fig_paths_3, fig_pdf)
print(f"  Figures PDF created: {os.path.getsize(fig_pdf)} bytes")

esm3_path = os.path.join(OUTPUT_DIR, 'ESM_3.pdf')
merge_pdfs([title_pdf_3, fig_pdf], esm3_path)
print(f"  ESM_3.pdf created: {os.path.getsize(esm3_path)} bytes")

# Clean up
for p in [title_pdf_3, fig_pdf]:
    if os.path.exists(p):
        os.remove(p)

# === Final verification ===
print("\n=== FINAL VERIFICATION ===")
from PyPDF2 import PdfReader
for esm in ['ESM_1.pdf', 'ESM_2.pdf', 'ESM_3.pdf']:
    fp = os.path.join(OUTPUT_DIR, esm)
    if os.path.exists(fp):
        sz = os.path.getsize(fp)
        reader = PdfReader(fp)
        pages = len(reader.pages)
        print(f"  {esm}: {sz:,} bytes, {pages} pages")
    else:
        print(f"  {esm}: MISSING!")
