from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\Workplace\SCS")
IN_DOCX = ROOT / "revised_manuscript.docx"
OUT_DOCX = ROOT / "revised_manuscript_2006_2024_noER.docx"
FIG_DIR = ROOT / "outputs" / "figures"
FRAMEWORK = FIG_DIR / "fig_research_framework_2006_2024_noER.png"


doc = Document(str(IN_DOCX))


def clear_para(p):
    p_pr = p._p.pPr
    for child in list(p._p):
        if p_pr is not None and child is p_pr:
            continue
        p._p.remove(child)


def set_para_text(p, text):
    clear_para(p)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def replace_paragraph_start(prefix, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            set_para_text(p, text)
            return p
    raise RuntimeError(f"Paragraph not found: {prefix}")


def remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)


def paragraph_after_element(element, text=None):
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    p = Paragraph(new_p, doc)
    if text:
        set_para_text(p, text)
    return p


def remove_table(table):
    table._tbl.getparent().remove(table._tbl)


def make_framework_figure():
    img = Image.new("RGB", (2200, 1050), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 44)
        font = ImageFont.truetype("arial.ttf", 30)
        small = ImageFont.truetype("arial.ttf", 26)
        bold = ImageFont.truetype("arialbd.ttf", 30)
    except OSError:
        title_font = font = small = bold = ImageFont.load_default()

    def box(xy, text, fill, outline, fnt=None):
        draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=3)
        lines = []
        for raw in text.split("\n"):
            words = raw.split()
            line = ""
            for word in words:
                trial = (line + " " + word).strip()
                if draw.textlength(trial, font=fnt or font) <= xy[2] - xy[0] - 60:
                    line = trial
                else:
                    if line:
                        lines.append(line)
                    line = word
            if line:
                lines.append(line)
        line_h = 36
        y = xy[1] + (xy[3] - xy[1] - line_h * len(lines)) / 2
        for line in lines:
            w = draw.textlength(line, font=fnt or font)
            draw.text((xy[0] + (xy[2] - xy[0] - w) / 2, y), line, fill="#111111", font=fnt or font)
            y += line_h

    def arrow(start, end):
        draw.line([start, end], fill="#222222", width=5)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) >= abs(ey - sy):
            pts = [(ex, ey), (ex - 22, ey - 13), (ex - 22, ey + 13)] if ex > sx else [(ex, ey), (ex + 22, ey - 13), (ex + 22, ey + 13)]
        else:
            pts = [(ex, ey), (ex - 13, ey - 22), (ex + 13, ey - 22)] if ey > sy else [(ex, ey), (ex - 13, ey + 22), (ex + 13, ey + 22)]
        draw.polygon(pts, fill="#222222")

    draw.text((110, 55), "Research framework after noER revision", fill="#111111", font=title_font)
    box((120, 290, 530, 610), "Urban multidimensional\nSCCD\n\nProduction\nLiving\nEcological\nDigital", "#fff2cc", "#d6a600", small)
    box((760, 285, 1180, 440), "Nonlinear effect\nInverted U-shaped\nSCCD -> lnCE", "#eadcf8", "#8d6cb3", bold)
    box((1520, 305, 1810, 425), "Log carbon\nemissions\n(lnCE)", "#dce8fa", "#6f92c8", bold)
    box((760, 655, 1180, 820), "Mechanism channels\nIndustrial upgrading\nGreen technological innovation", "#e1f0dc", "#77aa77", small)
    box((760, 95, 1180, 220), "Core moderator\nDigital economy\ndevelopment", "#dcefd8", "#77aa77", bold)
    box((1310, 655, 1950, 820), "Supplementary check\nPolycentric urban structure\nNo robust moderation claim", "#f2f2f2", "#999999", small)
    box((1275, 120, 1960, 250), "Digital capability reshapes\nmarginal effects and attenuates curvature", "#f7f2e8", "#b79c64", small)
    arrow((530, 450), (760, 365))
    arrow((1180, 365), (1520, 365))
    arrow((530, 520), (760, 735))
    arrow((1180, 735), (1520, 410))
    arrow((970, 220), (970, 285))
    arrow((1180, 355), (1275, 185))
    draw.line([(1310, 735), (1180, 735)], fill="#777777", width=3)
    img.save(FRAMEWORK)


def set_cell(cell, text, size=8):
    cell.text = str(text)
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(0)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
    cell.vertical_alignment = 1


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    for edge in ("left", "right", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def fill_table(table, rows, font_size=8):
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row, vals in zip(table.rows, rows):
        if len(row.cells) != len(vals):
            raise ValueError("Column mismatch while filling table")
        for cell, val in zip(row.cells, vals):
            set_cell(cell, val, font_size)
    set_table_borders(table)
    table.alignment = 1


def replace_table(old_table, rows, font_size=8):
    new_table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    fill_table(new_table, rows, font_size)
    old_table._tbl.addnext(new_table._tbl)
    remove_table(old_table)
    return new_table


def insert_table_after_paragraph(paragraph, rows, font_size=8):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    fill_table(table, rows, font_size)
    paragraph._p.addnext(table._tbl)
    return table


def picture_before_caption(caption_text, image_path, width=Inches(6.3)):
    for p in doc.paragraphs:
        if p.text.strip() == caption_text:
            prev = p._p.getprevious()
            if prev is not None:
                pic_p = Paragraph(prev, p._parent)
                clear_para(pic_p)
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_p.add_run().add_picture(str(image_path), width=width)
            return
    raise RuntimeError(f"Caption not found: {caption_text}")


make_framework_figure()


detailed_indicator_rows = []
for table in doc.tables[:3]:
    for row in table.rows[1:]:
        vals = [cell.text.strip() for cell in row.cells]
        if vals[2] == "Environmental regulation intensity":
            continue
        detailed_indicator_rows.append(vals)

for table in list(doc.tables[:3]):
    remove_table(table)
for p in list(doc.paragraphs):
    if p.text.strip().startswith("Table 1. Urban multidimensional spatial indicator system (continued)."):
        remove_paragraph(p)


paragraph_updates = [
    ("This study investigates whether and how digital-enabled urban multidimensional spatial coordination affects carbon emissions in China.",
     "This study investigates whether and how digital-enabled urban multidimensional spatial coordination affects carbon emissions in China. Building on the conventional production-living-ecological framework, the analysis incorporates digital space as a fourth functional dimension and constructs a spatial coupling coordination degree (SCCD) index to capture the joint evolution of production, living, ecological, and digital functions. The revised main analysis uses panel data for 284 prefecture-level and above cities from 2006 to 2024. Two-way fixed-effects estimates show a significant inverted U-shaped relationship between SCCD and lnCE, with a turning point of 0.542. Formal endpoint tests confirm that the left-side marginal effect is significantly positive and the right-side marginal effect is significantly negative. Industrial upgrading and green technological innovation operate as mechanism channels, while digital economy development reshapes and attenuates the nonlinear carbon effect of multidimensional spatial coupling coordination. Supplementary IV evidence mitigates endogeneity concerns but is not interpreted as fully resolving them."),
    ("Keywords:", "Keywords: Urban spatial coordination; Digital space; lnCE; Industrial upgrading; Green technological innovation; Digital economy; Polycentric urban structure"),
    ("At the same time, digitalization is reshaping the logic of urban development.",
     "At the same time, digitalization is reshaping the logic of urban development. Digital infrastructure, data connectivity, platform coordination, and intelligent governance can improve resource allocation, support cleaner production, and strengthen green governance instruments (Balogun et al., 2020; Añón Higón et al., 2017). Yet digitalization is also associated with electricity use, equipment replacement, internet-related energy demand, and rebound effects, so its environmental consequences remain contested (Lange et al., 2020; Obringer et al., 2021). Recent evidence for China shows that digital city construction and the digital economy can improve carbon performance and generate spatial spillovers, although the effects are often nonlinear and mechanism-dependent (Yang et al., 2022; Zhang et al., 2022; Cheng et al., 2023; Hou et al., 2024; Zhu et al., 2024)."),
    ("Despite these advances, three gaps remain.",
     "Despite these advances, three gaps remain. First, most studies still treat urban space as a predominantly physical system and rarely incorporate digital space into the analytical framework of production, living, and ecological functions. Second, the literature has paid more attention to isolated dimensions of urban form than to the coupling coordination among multiple urban functional spaces, even though recent functional-zone evidence shows that intra-urban functional heterogeneity matters for carbon outcomes (Tan et al., 2024). Third, although industrial upgrading and green technological innovation are widely recognized as important channels of low-carbon transition, their roles in the relationship between multidimensional spatial coordination and carbon emissions remain insufficiently integrated into a unified analytical framework (Tian et al., 2019; Du et al., 2019; Lin and Ma, 2022; Chang et al., 2023)."),
    ("Accordingly, this study investigates how urban multidimensional spatial coupling coordination affects carbon emissions",
     "Accordingly, this study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation. Extending the conventional production-living-ecological framework, it treats digital space as a fourth functional dimension and constructs a spatial coupling coordination degree (SCCD) index covering production, living, ecological, and digital functions. The main empirical sample covers 284 prefecture-level and above cities during 2006-2024."),
    ("This study contributes in three respects.",
     "This study contributes in three respects. First, it extends the analytical framework of urban functional space by explicitly incorporating digital space into multidimensional spatial coupling coordination. Second, it provides city-level evidence on the nonlinear relationship between SCCD and lnCE in China using the unified 2006-2024 sample. Third, it clarifies mechanism channels and boundary conditions, with industrial upgrading, green technological innovation, and digital economy development as the core supported channels; polycentricity is retained only as a supplementary interaction check."),
    ("Urban spatial restructuring is widely regarded",
     "Urban spatial restructuring is widely regarded as a key determinant of urban environmental performance. Existing studies show that urban expansion, land-use change, infrastructure investment, and the reorganization of production and residential activities reshape energy demand and carbon emissions (Seto et al., 2012; Fang et al., 2015; Castells-Quintana et al., 2021; Hong et al., 2022). One strand argues that compact development can lower per-capita emissions through density economies, shorter travel distances, and more efficient infrastructure (Gaigné et al., 2012; Lee and Lee, 2014; Kaza, 2020). Another emphasizes that rapid spatial expansion and functional reconfiguration often entail construction booms, additional transport demand, and higher energy use, especially in developing and industrializing contexts (Dhakal, 2009; Castells-Quintana et al., 2021). Recent functional-zone evidence further shows that the carbon implications of urban form depend on the spatial heterogeneity of land-use and transport functions within cities (Tan et al., 2024)."),
    ("Digital technologies are reshaping the logic of urban development.",
     "Digital technologies are reshaping the logic of urban development. Existing studies show that digitalization can support climate adaptation, resource optimization, and smarter urban governance (Balogun et al., 2020). Research on ICTs further suggests that digital tools can improve coordination efficiency and green governance, while also increasing electricity consumption, equipment demand, and rebound effects (Añón Higón et al., 2017; Lange et al., 2020; Obringer et al., 2021). For China, recent evidence indicates that digital city construction and the digital economy can improve carbon performance through efficiency gains, industrial upgrading, technological innovation, and spatial spillovers (Yang et al., 2022; Zhang et al., 2022; Cheng et al., 2023; Hou et al., 2024; Zhu et al., 2024)."),
    ("Fig. 1 illustrates the research framework",
     "Fig. 1 illustrates the research framework and mechanism analysis of this study. It summarizes the direct nonlinear effect of urban multidimensional spatial coupling coordination on lnCE, the mechanism roles of industrial upgrading and green technological innovation, and the primary moderating role of digital economy development. Polycentric urban structure is retained as a supplementary interaction check rather than as a supported moderator."),
    ("Digital economy development may further condition",
     "Digital economy development may further condition the nonlinear carbon effect of multidimensional spatial coupling coordination. A more developed digital economy can lower coordination frictions, improve information matching, and support smart transport, energy management, and urban governance. The revised hypothesis does not claim a mechanically earlier turning point; instead, it focuses on whether digital capability changes marginal effects and reshapes the curvature of the SCCD-lnCE relationship."),
    ("H4.", "H4. Digital economy development reshapes and attenuates the nonlinear carbon effect of multidimensional spatial coupling coordination."),
    ("This study uses observed panel data",
     "This study uses panel data for 284 prefecture-level and above cities in China over the period 2006-2024 as the main empirical sample. The analysis file contains 5,396 city-year records and is balanced by city and year. The raw file also includes an is_fitted flag for 852 rows in 2022-2024; this revision keeps those years within the unified main sample, while the data-construction procedure for those rows should be documented before submission. The sample window is determined by the availability of multidimensional spatial indicators and city-level carbon-emission data. The raw data are drawn mainly from the China City Statistical Yearbook, China Energy Statistical Yearbook, China Urban Construction Statistical Yearbook, China Rural Statistical Yearbook, China Industrial Statistical Yearbook, and China Environmental Statistical Yearbook. City-level carbon emissions are taken from the Emissions Database for Global Atmospheric Research (EDGAR), which ensures temporal comparability under a consistent accounting framework."),
    ("The dependent variable is carbon emissions",
     "The dependent variable in all revised regressions is lnCE. The available Stata file stores the logarithmic carbon-emission measure under the raw variable name CE; the revised do-files therefore generate lnCE = CE and do not apply an additional logarithmic transformation. Because no separate raw emissions level is available in the data file, the manuscript defines lnCE operationally as the log carbon-emission variable supplied in the analysis dataset. The minimum of the stored series is 0.000, so the authors should confirm before submission whether the original upstream construction used ln(carbon emissions) or ln(carbon emissions + 1)."),
    ("The core explanatory variable is the spatial coupling coordination degree",
     "The core explanatory variable is the spatial coupling coordination degree (SCCD) of urban multidimensional spatial functions. Drawing on the measurement logic of composite digital-development and urban-function indices, and adapting it to the production-living-ecological-digital framework emphasized in this study, the revised indicator system covers four functional dimensions, 13 criterion-level dimensions, and 36 specific indicators after deleting one policy-term frequency indicator. Larger SCCD values indicate a higher level of coordination across production, living, ecological, and digital spaces."),
    ("Using the entropy-weighting method",
     "Using the entropy-weighting method, the study first derives the development level of each functional dimension and then calculates the overall SCCD index. Equal weights are assigned to the four dimensions when constructing the final coordination measure, so that the estimated effect can be interpreted as the joint outcome of multidimensional urban-space integration rather than the dominance of any single spatial function. Table 1 reports a compressed summary of the SCCD indicator system, while Appendix Table A1 lists the detailed indicators after deleting the policy-term frequency item."),
    ("Digital economy index (DEI)",
     "Digital economy index (DEI) is examined as the main moderating variable because it captures the broader digital-development environment across cities and remains statistically supported in the revised 2006-2024 estimates. Polycentricity (POLY), reflecting spatial dispersion and multi-center organization, is retained only as a supplementary interaction check."),
    ("Table 2 reports descriptive statistics",
     "Table 2 reports descriptive statistics for the 2006-2024 main sample. The mean value of lnCE is 4.612, whereas the mean SCCD is 0.318, suggesting that most sample cities remain at a moderate level of multidimensional spatial coordination. The data also show substantial cross-city variation in digital economy development, green technological innovation, and polycentric urban structure, supporting the analysis of nonlinear, mechanism, and moderation effects."),
    ("Note: Statistics are for the observed", "Note: Statistics are for the 2006-2024 main sample. lnCE is generated as lnCE = CE because the available data file already stores the log carbon-emission measure as CE."),
    ("Fig. 2 visualizes", "Fig. 2 visualizes the SCCD distribution in the 2006-2024 main sample. The distribution provides descriptive context for the nonlinear empirical analysis."),
    ("Fig. 2. Distribution", "Fig. 2. Distribution of SCCD in the 2006-2024 main sample."),
    ("Note: SCCD denotes", "Note: SCCD denotes the spatial coupling coordination degree."),
    ("CEit = β0", "lnCE_{it} = β0 + β1SCCD_{it} + β2SCCD²_{it} + β3X_{it} + μi + λt + εit"),
    ("CEit = δ0", "lnCE_{it} = δ0 + δ1SCCD_{it} + δ2SCCD²_{it} + δ3M_{it} + δ4X_{it} + μi + λt + ξit"),
    ("To examine boundary conditions",
     "To examine boundary conditions, interaction terms are constructed between SCCD and selected contextual variables. Digital economy development is the primary moderator in the revised interpretation. POLY is retained as a supplementary check to test whether polycentric urban structure changes the SCCD-lnCE relationship."),
    ("CEit = θ0", "lnCE_{it} = θ0 + θ1SCCD_{it} + θ2SCCD²_{it} + θ3Z_{it} + θ4SCCD_{it} × Z_{it} + θ5SCCD²_{it} × Z_{it} + θ6X_{it} + μi + λt + ωit"),
    ("Table 3 reports the baseline estimates",
     "Table 3 reports the baseline estimates for the 2006-2024 main sample. Across all specifications, the coefficient on the linear SCCD term is significantly positive and the coefficient on the squared term is significantly negative, providing strong evidence of an inverted U-shaped relationship between urban multidimensional spatial coupling coordination and lnCE. In the full-control specification, the coefficient of SCCD is 8.378 and the coefficient of SCCD2 is -7.725, both significant at the 1% level. The result indicates that the carbon effect of spatial coordination is not monotonic. In the initial stage, improvements in coordination are still accompanied by infrastructure expansion, construction activity, factor agglomeration, and rising energy demand, so lnCE increases with SCCD. Once coordination reaches a sufficiently mature stage, however, the efficiency effect begins to dominate: digital-physical integration improves factor matching, lowers information frictions, strengthens knowledge spillovers, and supports cleaner urban operation."),
    ("The estimated main-sample turning point",
     "The estimated turning point is 0.542, calculated from the 2006-2024 full-control specification. This value lies within the SCCD range of 0.121 to 0.822. A formal nonlinear test confirms the inverted-U pattern: the left-endpoint marginal effect is significantly positive (6.502, p < 0.001), the right-endpoint marginal effect is significantly negative (-4.325, p = 0.005), and the nlcom 95% confidence interval for the turning point is [0.450, 0.634]. Fig. 3 provides a descriptive bivariate visualization of the lnCE-SCCD relationship, but the turning-point interpretation is based on the two-way fixed-effects estimates in Table 3."),
    ("Fig. 3. Nonlinear effect", "Fig. 3. Nonlinear effect of SCCD on lnCE."),
    ("Table 4 reports robustness",
     "Table 4 reports robustness and supplementary IV evidence for the 2006-2024 main sample. Panel A shows that the nonlinear pattern remains after winsorizing variables at the 1% level, replacing lnCE with average nighttime light intensity as an alternative dependent variable, and estimating a supplementary IV second stage. The requested purified SCCD robustness check could not be completed because the available data file contains only the composite SCCD variable and does not include the raw indicator-level panel or entropy-weighting script needed to reconstruct SCCD after excluding overlapping indicators. This limitation is reported rather than filled with unsupported results."),
    ("To further address endogeneity",
     "To mitigate endogeneity concerns, IV and IV2 are used as instrumental variables in a supplementary two-stage least squares estimation. This strategy should not be interpreted as fully resolving endogeneity because the instruments are lagged SCCD and lagged SCCD squared. The second-stage results show that the coefficient of SCCD remains significantly positive, while the coefficient of SCCD2 remains significantly negative. The implied IV turning point is approximately 0.638, which remains within the sample SCCD range. First-stage and weak-instrument diagnostics support instrument relevance: the first-stage F statistics are 85.651 for SCCD and 106.982 for SCCD2, the Kleibergen-Paap rk LM statistic is 57.623 (p < 0.001), and the Kleibergen-Paap rk Wald F statistic is 73.061. The overidentification test is not applicable because the model is exactly identified."),
    ("Table 5 reports the regional heterogeneity",
     "Table 5 reports the regional heterogeneity results for the 2006-2024 main sample. Given substantial regional differences in development patterns, industrial structure, and urbanization, this study further examines whether the carbon effect of SCCD differs across eastern, central, and western China. The estimates show that the inverted U-shaped relationship remains statistically significant in all three regions, suggesting that the nonlinear effect is broadly robust across China. The absolute value of the quadratic coefficient is largest in central China (-11.725), followed by western China (-7.151) and eastern China (-6.922), indicating that lnCE in central China is most sensitive to changes in multidimensional spatial coordination."),
    ("Table 6 reports the mechanism",
     "Table 6 reports the mechanism analysis results for the 2006-2024 main sample. To uncover mechanisms linking SCCD to lnCE, this study tests the roles of industrial upgrading and green technological innovation. The results show that SCCD has an inverted U-shaped effect on both mechanism variables, suggesting that stronger spatial coordination initially promotes industrial upgrading and green technological innovation, but that these gains weaken as coordination becomes more mature. The current workflow provides mechanism-equation evidence rather than separately estimated indirect-effect magnitudes."),
    ("When industrial upgrading is added",
     "When industrial upgrading is added to the lnCE equation, its coefficient is significantly positive (1.491), while the coefficients on SCCD and its squared term remain statistically significant. This indicates that industrial upgrading is an important transmission channel, although its short-run carbon effect is emission-increasing before cleaner production systems are fully established. From a transitional perspective, upgrading often involves new fixed-asset investment, equipment replacement, construction of modern industrial facilities, and the coexistence of old and new sectors. The result therefore supports H2 as mechanism evidence, while also suggesting that the quality and timing of industrial upgrading are as important as the direction of upgrading itself."),
    ("When green technological innovation is added",
     "When green technological innovation is added to the lnCE equation, its coefficient is also significantly positive (0.325), and the SCCD coefficients remain statistically significant. This confirms that green technological innovation acts as another important transmission channel. Yet the positive coefficient of the mediator suggests that green innovation does not automatically lower observed emissions in the short run. Research and development, testing, commercialization, and adoption all require capital and energy inputs, while efficiency gains may be partly offset by rebound effects."),
    ("Table 7 reports the moderation",
     "Table 7 reports the moderation and supplementary interaction checks for the 2006-2024 main sample. The revised interpretation focuses on digital economy development as the only core moderator. POLY is retained as a supplementary check, but it is not treated as a supported moderation mechanism because its interaction terms are not statistically significant."),
    ("The environmental-regulation interaction",
     "The interaction terms for digital economy development are statistically significant and support the revised H4. The coefficient on SCCD x DEI is significantly negative (-22.377), while the coefficient on SCCD2 x DEI is significantly positive (21.644). This pattern indicates that DEI reduces the effective linear marginal effect and makes the effective quadratic term less negative. Conditional calculations show that the turning point is 0.574 at low DEI, 0.583 at mean DEI, and 0.587 at high DEI, so the evidence does not support the earlier claim that DEI brings the turning point forward. Instead, DEI mainly attenuates and reshapes the nonlinear relationship."),
    ("The interaction terms for digital economy development",
     "At the sample mean of SCCD, the marginal effect of SCCD on lnCE declines from 3.842 at low DEI to 3.424 at mean DEI and 3.288 at high DEI, with all three effects statistically significant. Substantively, digital economy development changes how spatial coordination translates into emissions by lowering coordination frictions, improving information matching, and strengthening urban governance efficiency. The carbon effect of spatial coordination therefore depends increasingly on cities' digital capabilities."),
    ("For polycentric urban structure",
     "For polycentric urban structure, the generated interaction terms do not statistically support a moderation claim. The coefficients on SCCD x POLY and SCCD2 x POLY are not significant in the 2006-2024 estimates. Polycentricity may still matter theoretically because dispersed centers can create longer logistics chains, weaker commuting-job matching, and duplicated infrastructure when cross-center links are incomplete, but this argument should be presented as a possible explanation requiring further empirical confirmation rather than as a confirmed moderation result."),
    ("Taken together, the moderation",
     "Taken together, the moderation and supplementary interaction results show that digital capability is the clearest boundary condition in the updated evidence. The current estimates do not support a strong empirical claim for POLY as a robust moderator. This finding complements the baseline and mechanism results by showing that the environmental consequences of spatial coordination depend especially on whether cities can translate digital capacity into lower coordination frictions and more efficient urban operation."),
    ("This study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation.",
     "This study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation. By extending the conventional production-living-ecological framework to include digital space and by using panel data for 284 prefecture-level and above cities over 2006-2024, the paper evaluates both the nonlinear carbon effect of spatial coordination and the mechanisms through which this effect operates. The central message is that multidimensional spatial coordination is neither inherently low-carbon nor inherently carbon-intensive. Its environmental consequences depend on development stage, structural adjustment, and governance capacity."),
    ("Four main conclusions can be drawn.",
     "Four main conclusions can be drawn. First, SCCD exhibits a significant inverted U-shaped relationship with lnCE in the 2006-2024 main sample, with a turning point of 0.542 and formal endpoint tests supporting the inverted-U shape. Second, the nonlinear relationship remains robust after winsorization, an alternative dependent-variable check, and supplementary IV estimation, although the IV results are best interpreted as mitigating rather than fully resolving endogeneity concerns. Third, industrial upgrading and green technological innovation constitute important mechanism channels, but their short-run effects are not purely emission-reducing because transition costs, equipment renewal, and rebound effects partly offset efficiency gains. Fourth, digital economy development significantly conditions the carbon effect of SCCD by reshaping marginal effects and attenuating curvature, while the updated evidence does not support a strong POLY moderation claim."),
    ("Overall, the evidence shows",
     "Overall, the evidence shows that multidimensional spatial coordination does not automatically produce low-carbon outcomes. Its environmental value depends on whether spatial integration is translated into efficiency improvement rather than land expansion, whether structural and technological upgrading generate durable rather than temporary carbon gains, and whether urban governance can manage the frictions associated with digital-physical transformation. For cities below the turning point of 0.542, the immediate challenge is to prevent coordination from being translated primarily into duplicated infrastructure and energy-intensive growth. For cities approaching or surpassing the turning point, the policy task shifts toward consolidating the efficiency advantages of digital integration and converting innovation and upgrading into persistent emission reductions."),
    ("First, urban low-carbon policy",
     "First, urban low-carbon policy should move beyond extensive physical expansion and prioritize deeper digital-physical integration. Because many sample cities remain on the left side of the turning point, simply raising formal indicators of spatial coordination is insufficient to guarantee lower emissions. Policy should therefore focus on digital infrastructure, interoperable data systems, and intelligent governance tools that reduce the coordination costs of production, living, ecological, and digital spaces."),
    ("Second, policy design",
     "Second, policy design should explicitly manage the rebound risks associated with industrial upgrading and green technological innovation. Because structural transformation and green innovation do not automatically reduce emissions in the short run, complementary measures such as cleaner energy substitution, green finance, green governance instruments, and performance-based evaluation are needed to ensure that transition costs do not lock cities into higher-carbon pathways."),
    ("Fourth, governance strategies",
     "Fourth, governance strategies should be differentiated across regions and development stages. Central China deserves particular attention because it faces both stronger emission responsiveness to changes in SCCD and greater potential gains from later-stage coordination. Eastern China is better positioned to explore advanced digital-physical governance models, whereas western China should place greater emphasis on coupling ecological protection with improvements in basic public services and urban functional coordination. More broadly, cities below the turning point should prioritize curbing construction-led expansion and improving the quality rather than the quantity of coordination; cities near the turning point should accelerate digital governance reform, industrial upgrading quality, policy coordination, and green technology diffusion; and cities already beyond the turning point should focus on locking in low-carbon gains through institutional coordination, cross-sector data integration, and long-term monitoring of rebound effects."),
    ("This study has several limitations.",
     "This study has several limitations. First, although the SCCD index incorporates digital space into the conventional production-living-ecological framework, the measurement of digital space is constrained by data availability and may not fully capture platform governance, data mobility, algorithmic coordination, or the quality of digital public services. Second, although the empirical design includes fixed effects, robustness checks, and supplementary instrumental-variable estimation, the analysis remains based on observational panel data and therefore cannot fully eliminate all identification concerns. Third, the available data file contains an is_fitted indicator for 2022-2024 rows, but the construction script for those rows is not available in the current project; this should be documented transparently before submission. Fourth, the requested purified SCCD robustness check could not be completed without raw indicator-level data and the entropy-weighting script. Fifth, the mechanism analysis focuses on industrial upgrading and green technological innovation, but other channels--such as energy structure, green finance, or household behavioral change--may also play meaningful roles in shaping carbon outcomes."),
]

for prefix, text in paragraph_updates:
    replace_paragraph_start(prefix, text)

for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("At the sample mean of SCCD, the marginal effect"):
        set_para_text(
            p,
            "The interaction terms for digital economy development are statistically significant and support the revised H4. The coefficient on SCCD x DEI is significantly negative (-22.377), while the coefficient on SCCD2 x DEI is significantly positive (21.644). This pattern indicates that DEI reduces the effective linear marginal effect and makes the effective quadratic term less negative. Conditional calculations show that the turning point is 0.574 at low DEI, 0.583 at mean DEI, and 0.587 at high DEI, so the evidence does not support the earlier claim that DEI brings the turning point forward. Instead, DEI mainly attenuates and reshapes the nonlinear relationship.",
        )
    elif text.startswith("The interaction terms for digital economy development are statistically significant and support the revised H4. In the observed"):
        set_para_text(
            p,
            "At the sample mean of SCCD, the marginal effect of SCCD on lnCE declines from 3.842 at low DEI to 3.424 at mean DEI and 3.288 at high DEI, with all three effects statistically significant. Substantively, digital economy development changes how spatial coordination translates into emissions by lowering coordination frictions, improving information matching, and strengthening urban governance efficiency. The carbon effect of spatial coordination therefore depends increasingly on cities' digital capabilities.",
        )


picture_before_caption("Fig. 1. Research framework and mechanism analysis diagram.", FRAMEWORK, Inches(6.4))
picture_before_caption("Fig. 2. Distribution of SCCD in the 2006-2024 main sample.", FIG_DIR / "fig_sccd_distribution_2006_2024_noER.png", Inches(6.2))
picture_before_caption("Fig. 3. Nonlinear effect of SCCD on lnCE.", FIG_DIR / "fig_nonlinear_sccd_lnce_2006_2024_noER.png", Inches(6.2))


table1_caption = replace_paragraph_start("Table 1. Urban multidimensional spatial indicator system.", "Table 1. Urban multidimensional spatial indicator system (summary).")
summary_rows = [
    ["Functional space", "Criterion layer", "Number of indicators", "Representative indicators", "Attribute direction"],
    ["Production", "Agricultural production function", "4", "Agricultural output; sown area; aquatic products; grain output", "+"],
    ["Production", "Industrial production function", "3", "Secondary industry output; industrial profits; manufacturing employment", "+"],
    ["Production", "Other production functions", "4", "Tertiary output; postal revenue; retail sales; financial employment", "+"],
    ["Living", "Basic living function", "5", "Population density; road area; residential investment; wages; public expenditure", "+"],
    ["Living", "Medical living function", "2", "Hospital beds; licensed physicians", "+"],
    ["Living", "Educational living function", "2", "Education expenditure; teachers", "+"],
    ["Living", "Science and technology living function", "3", "Science expenditure; invention patents; research-service employment", "+"],
    ["Living", "Cultural living function", "1", "Library collection", "+"],
    ["Ecological", "Environmental pressure", "3", "Wastewater; sulfur dioxide; smoke and dust emissions per unit area", "-"],
    ["Ecological", "Environmental response", "1", "Water, environment, and public-facilities employment", "+"],
    ["Ecological", "Environmental status", "2", "Green coverage; green space per capita", "+"],
    ["Digital", "Informatization function", "2", "Broadband users; telecommunication service volume", "+"],
    ["Digital", "Digital-intelligent function", "1", "Digital and smart enterprises", "+"],
    ["Digital", "Networked function", "3", "IoT, industrial internet, and satellite internet enterprises", "+"],
]
summary_table = insert_table_after_paragraph(table1_caption, summary_rows, 7)
table1_note = paragraph_after_element(summary_table._tbl, "Note: The SCCD system contains 36 indicators after deleting one policy-term frequency indicator. Dimension scores are calculated with the entropy-weighting method and then combined with equal weights across production, living, ecological, and digital functional spaces. Detailed indicators are reported in Appendix Table A1.")


table_rows = [
    [
        ["Category", "Variable", "Symbol", "Obs.", "Mean", "Median", "S.D.", "Min", "Max"],
        ["Dependent variable", "Log carbon emissions", "lnCE", "5396", "4.612", "4.654", "1.253", "0.000", "8.325"],
        ["Core explanatory variable", "Spatial coupling coordination degree", "SCCD", "5396", "0.318", "0.304", "0.085", "0.121", "0.822"],
        ["Control variable", "Trade openness", "OPEN", "5381", "0.193", "0.080", "0.337", "-0.720", "3.640"],
        ["Control variable", "Urbanization rate", "UR", "5396", "54.498", "52.606", "15.419", "6.491", "100.000"],
        ["Control variable", "Urban-rural income gap", "URG", "5381", "2.437", "2.342", "0.549", "1.207", "6.378"],
        ["Control variable", "Government intervention intensity", "GI", "5396", "0.186", "0.162", "0.098", "0.043", "1.027"],
        ["Mechanism variable", "Industrial upgrading", "OIU", "5396", "2.288", "2.283", "0.144", "1.163", "2.836"],
        ["Mechanism variable", "Green technological innovation", "GTI", "5396", "4.047", "4.038", "1.795", "0.000", "9.872"],
        ["Moderator", "Digital economy index", "DEI", "5396", "0.060", "0.047", "0.076", "0.000", "0.940"],
        ["Supplementary check", "Degree of polycentricity", "POLY", "5396", "0.346", "0.374", "0.189", "0.000", "0.963"],
    ],
    [
        ["Variables", "(1)", "(2)", "(3)", "(4)", "(5)"],
        ["Dependent variable", "lnCE", "lnCE", "lnCE", "lnCE", "lnCE"],
        ["SCCD", "13.468***", "12.344***", "9.123***", "8.871***", "8.378***"],
        ["", "(1.546)", "(1.521)", "(1.374)", "(1.408)", "(1.388)"],
        ["SCCD2", "-11.472***", "-10.369***", "-8.283***", "-7.882***", "-7.725***"],
        ["", "(1.814)", "(1.816)", "(1.622)", "(1.658)", "(1.665)"],
        ["OPEN", "", "0.486***", "0.204**", "0.223***", "0.259***"],
        ["", "", "(0.107)", "(0.084)", "(0.085)", "(0.084)"],
        ["UR", "", "", "0.032***", "0.030***", "0.028***"],
        ["", "", "", "(0.004)", "(0.004)", "(0.004)"],
        ["URG", "", "", "", "-0.135*", "-0.112"],
        ["", "", "", "", "(0.073)", "(0.070)"],
        ["GI", "", "", "", "", "-2.064***"],
        ["", "", "", "", "", "(0.626)"],
        ["Constant", "1.057***", "1.171***", "0.438", "0.944**", "1.400***"],
        ["", "(0.311)", "(0.299)", "(0.284)", "(0.377)", "(0.412)"],
        ["City fixed effects", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["Year fixed effects", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["N", "5396", "5381", "5381", "5366", "5366"],
        ["Within R-squared", "0.568", "0.579", "0.629", "0.631", "0.641"],
    ],
    [
        ["Variables / diagnostics", "(1)", "(2)", "(3)"],
        ["Panel A. Regression results", "Winsorized lnCE", "Alternative DV: NTL", "Supplementary IV second stage"],
        ["SCCD", "9.888***", "3.536***", "18.658***"],
        ["", "(1.414)", "(0.757)", "(2.683)"],
        ["SCCD2", "-9.886***", "-2.515***", "-14.618***"],
        ["", "(1.698)", "(0.789)", "(2.700)"],
        ["Controls", "Yes", "Yes", "Yes"],
        ["City fixed effects", "Yes", "Yes", "Yes"],
        ["Year fixed effects", "Yes", "Yes", "Yes"],
        ["N", "5366", "5366", "5084"],
        ["Within / centered R-squared", "0.647", "0.573", "0.816"],
        ["Panel B. First-stage and weak-instrument diagnostics", "", "", ""],
        ["First-stage F for SCCD", "", "", "85.651"],
        ["First-stage F for SCCD2", "", "", "106.982"],
        ["Kleibergen-Paap rk LM", "", "", "57.623***"],
        ["Kleibergen-Paap rk Wald F", "", "", "73.061"],
        ["Overidentification test", "", "", "Not applicable"],
    ],
    [
        ["Variables", "(1)", "(2)", "(3)"],
        ["Region", "Eastern China", "Central China", "Western China"],
        ["SCCD", "7.589***", "12.171***", "7.221***"],
        ["", "(2.677)", "(2.747)", "(1.831)"],
        ["SCCD2", "-6.922***", "-11.725***", "-7.151***"],
        ["", "(2.582)", "(3.489)", "(2.180)"],
        ["OPEN", "0.219**", "0.352", "0.266"],
        ["", "(0.093)", "(0.274)", "(0.289)"],
        ["UR", "0.030***", "0.024***", "0.028***"],
        ["", "(0.006)", "(0.007)", "(0.007)"],
        ["URG", "-0.036", "-0.037", "-0.205*"],
        ["", "(0.141)", "(0.138)", "(0.105)"],
        ["GI", "-0.730", "-2.461***", "-2.017**"],
        ["", "(1.130)", "(0.754)", "(0.999)"],
        ["Constant", "1.278*", "0.523", "1.946***"],
        ["", "(0.740)", "(0.763)", "(0.675)"],
        ["City fixed effects", "Yes", "Yes", "Yes"],
        ["Year fixed effects", "Yes", "Yes", "Yes"],
        ["N", "1828", "1885", "1653"],
        ["Within R-squared", "0.720", "0.629", "0.621"],
    ],
    [
        ["Variables", "(1)", "(2)", "(3)", "(4)"],
        ["Dependent variable", "Industrial upgrading", "lnCE with OIU", "Green technological innovation", "lnCE with GTI"],
        ["SCCD", "0.626***", "7.445***", "10.481***", "4.971***"],
        ["", "(0.163)", "(1.297)", "(1.386)", "(1.156)"],
        ["SCCD2", "-0.371*", "-7.172***", "-5.438***", "-5.957***"],
        ["", "(0.213)", "(1.486)", "(1.416)", "(1.418)"],
        ["OIU", "", "1.491***", "", ""],
        ["", "", "(0.337)", "", ""],
        ["GTI", "", "", "", "0.325***"],
        ["", "", "", "", "(0.030)"],
        ["Controls", "Yes", "Yes", "Yes", "Yes"],
        ["City fixed effects", "Yes", "Yes", "Yes", "Yes"],
        ["Year fixed effects", "Yes", "Yes", "Yes", "Yes"],
        ["N", "5366", "5366", "5366", "5366"],
        ["Within R-squared", "0.614", "0.651", "0.812", "0.683"],
    ],
    [
        ["Variables", "(1)", "(2)"],
        ["Interaction", "DEI moderator", "POLY supplementary check"],
        ["SCCD", "8.829***", "7.770***"],
        ["", "(1.289)", "(2.188)"],
        ["SCCD2", "-7.714***", "-8.223***"],
        ["", "(1.473)", "(2.661)"],
        ["DEI", "4.601**", ""],
        ["", "(2.080)", ""],
        ["SCCD x DEI", "-22.377***", ""],
        ["", "(7.346)", ""],
        ["SCCD2 x DEI", "21.644***", ""],
        ["", "(6.145)", ""],
        ["POLY", "", "-1.179"],
        ["", "", "(0.970)"],
        ["SCCD x POLY", "", "3.213"],
        ["", "", "(4.848)"],
        ["SCCD2 x POLY", "", "-0.186"],
        ["", "", "(5.505)"],
        ["Controls", "Yes", "Yes"],
        ["City fixed effects", "Yes", "Yes"],
        ["Year fixed effects", "Yes", "Yes"],
        ["N", "5366", "5366"],
        ["Within R-squared", "0.643", "0.644"],
    ],
]

for idx, rows in enumerate(table_rows, start=1):
    replace_table(doc.tables[idx], rows, 7 if idx in (2, 3, 5) else 8)

notes = [
    ("Note: Robust standard errors", "Note: Clustered standard errors at the city level are reported in parentheses. * p < 0.10, ** p < 0.05, *** p < 0.01."),
]
for old, new in notes:
    try:
        replace_paragraph_start(old, new)
    except RuntimeError:
        pass


existing_refs = {p.text.strip() for p in doc.paragraphs}
new_refs = [
    "Hou, J., Li, W., Zhang, X., 2024. Research on the impacts of digital economy on carbon emission efficiency at China's City level. PLOS ONE 19 (9), e0308001.",
    "Tan, G., Zhang, X., Xiong, S., Sun, Z., Lei, Y., Wang, H., Du, S., 2024. Assessing the impacts of urban functional form on anthropogenic carbon emissions: A case study of 31 major cities in China. Ecological Indicators 167, 112700.",
    "Zhu, X., Li, D., Zhou, S., Zhu, S., Yu, L., 2024. Evaluating coupling coordination between urban smart performance and low-carbon level in China's pilot cities with mixed methods. Scientific Reports 14, 20461.",
]
for ref in new_refs:
    if ref not in existing_refs:
        p = doc.add_paragraph(ref)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


appendix_heading = doc.add_paragraph("Appendix")
appendix_heading.style = doc.styles["Heading 1"]
appendix_caption = doc.add_paragraph("Appendix Table A1. Detailed SCCD indicator system after deleting one policy-term frequency indicator.")
appendix_rows = [["Functional space", "Criterion layer", "Indicator", "Calculation", "Unit", "Attribute"]] + detailed_indicator_rows
appendix_table = doc.add_table(rows=len(appendix_rows), cols=len(appendix_rows[0]))
fill_table(appendix_table, appendix_rows, 6)

for p in doc.paragraphs:
    stripped = p.text.strip()
    if stripped.startswith("Table ") or stripped.startswith("Appendix Table "):
        p.paragraph_format.keep_with_next = True


doc.save(str(OUT_DOCX))
print(OUT_DOCX)
