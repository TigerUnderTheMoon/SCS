from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree
from PIL import Image


ROOT = Path(r"D:\Workplace\SCS")
PACKAGE = ROOT / "outputs" / "submission_package_scs_2026-05-21_clean" / "submission_ready_final (2)"
MANUSCRIPT = PACKAGE / "manuscript.docx"
FIGURES_DIR = PACKAGE / "figures"
SOURCE_FIG2 = ROOT / "outputs" / "figures" / "fig_sccd_distribution_2006_2024_noER.png"
SOURCE_FIG4 = ROOT / "outputs" / "figures" / "fig_nonlinear_sccd_lnce_2006_2024_noER.png"

REORG_DIR = ROOT / "outputs" / "submission_package_scs_2026-05-24_reorganized"


SECTION_REVISIONS = {
    100: "5. Empirical results and discussion",
    101: "5.1 Baseline nonlinear SCCD-lnCE relationship",
    102: (
        "Table 3 presents the baseline estimates for the 2006-2024 main sample. "
        "Across all specifications, SCCD is significantly positive and SCCD2 is significantly negative, "
        "indicating an inverted U-shaped association between multidimensional spatial coordination and lnCE. "
        "In the full-control specification, the SCCD coefficient is 8.378 and the SCCD2 coefficient is -7.725, "
        "both significant at the 1% level. This result supports H1 and sharpens the central coordination paradox: "
        "at lower coordination levels, stronger links among production, living, ecological, and digital spaces "
        "are associated with infrastructure expansion, construction activity, factor agglomeration, and higher energy demand; "
        "at higher coordination levels, efficiency gains and improved allocation begin to offset these pressures."
    ),
    105: (
        "The estimated turning point is 0.542 in the 2006-2024 full-control specification. "
        "This value falls within the observed SCCD range of 0.121 to 0.822, but it should not be read as a universal policy threshold. "
        "In the estimation sample, 5,273 of 5,396 city-year observations (97.72%) are below the turning point, while only 123 observations "
        "(2.28%) are at or above it. The right side of the curve is therefore empirically thinner, and its interpretation should remain cautious. "
        "A formal nonlinear test supports the inverted U-shaped pattern: the left-endpoint marginal effect is significantly positive "
        "(6.502, p < 0.001), the right-endpoint marginal effect is significantly negative (-4.325, p = 0.005), and the nlcom 95% confidence "
        "interval for the turning point is [0.450, 0.634]. Fig. 4 provides a descriptive bivariate visualization; the causal interpretation "
        "continues to rest on the two-way fixed-effects estimates in Table 3."
    ),
    109: (
        "The control-variable estimates are also consistent with an urban-transition interpretation. "
        "Trade openness is positively associated with emissions, although the OPEN-series limitation noted above requires caution. "
        "The urbanization rate is significantly positive, suggesting that population concentration and urban expansion still raise emission pressure "
        "during the sample period. By contrast, the urban-rural income gap is no longer significant in the full specification, while government "
        "intervention intensity is significantly negative. This pattern suggests that public governance capacity may help determine whether stronger "
        "spatial coordination translates into cleaner urban operation."
    ),
    110: (
        "The baseline evidence therefore rejects a simple claim that more coordination automatically lowers emissions. "
        "The environmental payoff from coordination depends on transition stage and governance quality. "
        "Without cleaner energy systems, green standards, and data-enabled management, tighter functional connections may accelerate the circulation "
        "of carbon-intensive activities across urban spaces. Sustainable coordination requires a shift from expansion-oriented integration to "
        "efficiency-oriented operation."
    ),
    111: "5.2 Robustness and supplementary endogeneity evidence",
    112: (
        "Table 4 reports robustness checks and supplementary IV evidence for the 2006-2024 main sample. "
        "Panel A shows that the nonlinear pattern remains after winsorizing variables at the 1% level and after replacing lnCE with average "
        "nighttime light intensity. NTL is not treated as a direct carbon-emissions proxy; it is used as an alternative urban-activity measure "
        "to test whether the nonlinear coordination pattern depends only on the carbon-emissions variable."
    ),
    115: (
        "Lagged SCCD and lagged SCCD squared are used as instruments for SCCD and SCCD2 in a supplementary two-stage least-squares estimation. "
        "This design helps assess persistence-related endogeneity, but it cannot eliminate all identification concerns because the instruments are lagged "
        "values of the core explanatory variables. The second-stage estimates remain consistent with the baseline: SCCD is significantly positive and "
        "SCCD2 is significantly negative. The implied IV turning point is approximately 0.638, which remains within the sample SCCD range and is best "
        "read as supporting evidence. First-stage and weak-instrument diagnostics support instrument relevance: the first-stage F statistics are 85.651 "
        "for SCCD and 106.982 for SCCD2, the Kleibergen-Paap rk LM statistic is 57.623 (p < 0.001), and the Kleibergen-Paap rk Wald F statistic is 73.061. "
        "The overidentification test is not applicable because the model is exactly identified."
    ),
    117: (
        "Table 5 reports regional heterogeneity for the 2006-2024 main sample. "
        "Chinese cities differ in industrial structure, development stage, urbanization pressure, and digital capacity, so the same increase in SCCD may "
        "carry different emission implications across regions. The estimates show that the inverted U-shaped relationship remains statistically significant "
        "in eastern, central, and western China. The absolute value of the quadratic coefficient is largest in central China (-11.725), followed by western "
        "China (-7.151) and eastern China (-6.922), indicating that lnCE in central China is most responsive to changes in multidimensional coordination."
    ),
    120: (
        "The regional pattern fits a staged sustainable-city transition. "
        "In eastern China, the estimated curve is flatter, likely because industrial upgrading, digital development, and operational efficiency already play "
        "a larger role than large-scale physical expansion. In western China, the nonlinear relationship remains pronounced, reflecting the coexistence of "
        "infrastructure investment, industrial catch-up, and uneven digital development. Central China appears especially sensitive to coordination changes, "
        "suggesting that its transition path requires closer sequencing among digital investment, industrial restructuring, and low-carbon governance."
    ),
    121: (
        "These results imply that a uniform national pathway is unlikely to be effective. "
        "Central cities may gain substantially from better coordination, but they also face stronger transition pressure. Eastern cities can focus on consolidating "
        "data-enabled efficiency gains, whereas western cities may need to pair infrastructure and public-service improvements with ecological constraints. "
        "Across regions, digital investment, industrial restructuring, and ecological governance need to advance together rather than through isolated policy pushes."
    ),
    122: "5.4 Restructuring-pathway evidence",
    123: (
        "Table 6 reports restructuring-pathway evidence for the 2006-2024 main sample. "
        "The analysis asks whether industrial upgrading and green technological innovation are associated with the SCCD-lnCE relationship. "
        "SCCD has an inverted U-shaped association with both restructuring variables, suggesting that stronger coordination initially accompanies industrial upgrading "
        "and green innovation but that these associations weaken as coordination matures. These estimates are mechanism-equation evidence, not a formal mediation "
        "or pathway decomposition."
    ),
    126: (
        "When industrial upgrading is added to the lnCE equation, its coefficient is significantly positive (1.491), while the coefficients on SCCD and SCCD2 remain "
        "statistically significant. This result is consistent with a stage-dependent restructuring pathway rather than an immediate emission-reduction channel. "
        "In the short run, upgrading often involves fixed-asset investment, equipment renewal, construction of modern facilities, and the coexistence of old and new "
        "sectors. H2 therefore receives mechanism-equation support, but the result should be read as evidence of transition pressure rather than automatic low-carbon improvement."
    ),
    127: (
        "When green technological innovation is added to the lnCE equation, its coefficient is also significantly positive (0.325), and the SCCD coefficients remain "
        "statistically significant. This provides mechanism-equation evidence that green innovation operates as another stage-dependent pathway. Green innovation may "
        "improve long-run efficiency, but research and development, testing, commercialization, and technology adoption require capital and energy inputs. Efficiency gains "
        "may also be partly offset by rebound effects. H3 therefore receives mechanism-equation support, with the same caution that the observed short-run association is "
        "not an immediate emission-reduction effect."
    ),
    128: (
        "Taken together, the pathway results explain why the baseline relationship is nonlinear. "
        "Multidimensional coordination can activate restructuring processes that are desirable for long-run sustainability, but those processes can also raise short-run "
        "energy demand. The governance challenge is to shorten the adjustment period, guide investment toward cleaner technologies, and reduce rebound effects during "
        "structural transformation."
    ),
    129: "5.5 Digital-economy moderation and supplementary interaction checks",
    130: (
        "Table 7 reports moderation and supplementary interaction checks for the 2006-2024 main sample. "
        "The interpretation focuses on digital economy development because it is theoretically central to the digital-space framework and statistically supported in the estimates. "
        "POLY is retained as a supplementary check, but it is not treated as a supported moderation mechanism because its interaction terms are not statistically significant."
    ),
    136: (
        "Fig. 5 visualizes the corresponding conditional curves. The high-DEI curve is flatter than the low-DEI curve across the observed SCCD range, which is consistent with "
        "attenuation of the marginal emission burden of coordination. The figure supports the interpretation visually, while statistical inference remains based on the interaction "
        "estimates in Table 7."
    ),
    137: (
        "The interaction terms for digital economy development are statistically significant and support H4. "
        "The coefficient on SCCD x DEI is significantly negative (-22.377), while the coefficient on SCCD2 x DEI is significantly positive (21.644). This does not mean that DEI "
        "mechanically delivers low-carbon transition. It means that digital capability reduces the marginal emission cost associated with coordination and flattens the nonlinear curvature. "
        "Conditional calculations show turning points of 0.574, 0.583, and 0.587 at low, mean, and high DEI, respectively. The marginal effect of SCCD at the sample mean declines from "
        "3.842 to 3.424 and 3.288 across the same DEI levels."
    ),
    138: (
        "Substantively, digital economy development changes how coordination translates into emissions. "
        "Digital capability can reduce coordination frictions, improve information matching, support smart public services, and strengthen data-enabled governance. "
        "The key implication is attenuation: digital capability weakens the marginal carbon burden of coordination rather than simply moving cities more quickly to a low-emission stage."
    ),
    140: (
        "For polycentric urban structure, the interaction terms are statistically insignificant in the 2006-2024 estimates. "
        "POLY is therefore retained only as a supplementary interaction check, not as a supported moderator. This distinction is important because polycentricity may still matter for planning, "
        "but the current evidence does not support it as a statistically robust boundary condition in the SCCD-lnCE relationship."
    ),
    141: (
        "Overall, the moderation and supplementary interaction results identify digital capability as the clearest boundary condition in the updated evidence. "
        "This finding complements the baseline and pathway results by showing that the environmental consequences of coordination depend on whether cities can translate digital capacity into "
        "lower coordination frictions, smarter service allocation, and more efficient urban operation."
    ),
    143: "6.1 Main conclusions",
    144: (
        "This study examines how multidimensional urban spatial coordination is associated with carbon emissions under digital transformation. "
        "By extending the production-living-ecological framework to include digital space and using panel data for 284 Chinese prefecture-level and above cities from 2006 to 2024, "
        "the paper identifies a coordination paradox. Spatial coordination can reduce frictions and improve allocation, but it can also intensify infrastructure expansion, industrial "
        "reorganization, and energy demand. Its environmental consequences therefore depend on transition stage and governance capacity."
    ),
    145: (
        "Four conclusions follow from the evidence. First, SCCD has a significant inverted U-shaped association with lnCE, with a turning point of 0.542. "
        "Because 97.72% of city-year observations remain below that point, the efficiency-oriented stage is not yet dominant for most observed cases, and evidence on the right side of the curve "
        "should be interpreted cautiously. Second, the nonlinear pattern remains under winsorization, an NTL urban-activity check, and supplementary IV estimation, although the IV results are used "
        "only as supporting evidence. Third, industrial upgrading and green innovation are stage-dependent restructuring pathways. Fourth, digital economy development is the supported boundary condition, "
        "whereas polycentricity remains a supplementary interaction check rather than a robust moderator."
    ),
    146: (
        "The main message is not that cities should slow coordination. Rather, coordination must be governed differently across stages. "
        "Cities below the turning point should avoid converting coordination into duplicated infrastructure and energy-intensive growth. Cities approaching the turning point should strengthen data-enabled "
        "governance, improve restructuring quality, and manage rebound risks. Cities already beyond the turning point should consolidate efficiency gains through institutional coordination, cross-sector data "
        "integration, and long-term monitoring."
    ),
    148: (
        "First, urban low-carbon policy should prioritize governance quality rather than the simple expansion of coordinated systems. "
        "Because most sample city-year observations remain on the left side of the turning point, raising formal coordination indicators is insufficient. Policy should focus on interoperable data systems, "
        "digital public services, smart transport, energy management, and intelligent governance tools that reduce the marginal emission cost of coordination across production, living, ecological, and digital spaces."
    ),
    149: (
        "Second, policy design should manage the transition costs associated with industrial upgrading and green innovation. "
        "These restructuring processes do not automatically reduce emissions in the short run. Cleaner energy substitution, green finance, performance-based evaluation, and tighter coordination between innovation "
        "policy and industrial policy are needed to prevent transition costs from becoming persistent emission burdens."
    ),
    150: (
        "Third, cities pursuing polycentric development should pay closer attention to cross-center coordination. "
        "The current interaction estimates do not provide robust support for polycentricity as a moderator, but the planning risk remains clear. Dispersed centers can raise logistics costs, duplicate infrastructure, "
        "and weaken jobs-housing balance when cross-center links are incomplete. Polycentric planning should therefore be paired with integrated land use, transport systems, infrastructure provision, and digital coordination tools."
    ),
    151: (
        "Fourth, governance strategies should be differentiated across regions and development stages. "
        "Central China deserves particular attention because it shows stronger emission responsiveness to changes in SCCD. Eastern China is better positioned to consolidate advanced data-enabled governance models, whereas western China "
        "should place greater emphasis on coupling ecological protection with basic public-service improvements and urban functional coordination. Across regions, the priority is to shift from expansion-oriented integration toward "
        "efficiency-oriented governance, a transition that depends on governance capacity rather than coordination levels alone."
    ),
    152: "6.3 Boundary conditions and future research",
    153: (
        "Several boundary conditions should guide interpretation. The SCCD index incorporates digital space into the production-living-ecological framework, but the available indicators cannot fully capture platform governance, "
        "data mobility, algorithmic coordination, or the quality of digital public services. The empirical design includes fixed effects, robustness checks, and supplementary IV estimation, but the analysis remains based on observational "
        "panel data and cannot eliminate all identification concerns. The mechanism-equation analysis focuses on industrial upgrading and green innovation; other pathways, including energy structure, green finance, and household behavior, "
        "may also shape emission outcomes."
    ),
    154: (
        "Future research can extend the framework in three directions. First, more refined measures of digital urban transformation could capture platform governance, data flows, smart service coverage, and the quality of urban digital public goods. "
        "Second, dynamic restructuring deserves closer attention because the emission payoff of industrial upgrading and green innovation may emerge with long lags. Third, city-type heterogeneity could be examined more directly, especially for "
        "resource-based cities, shrinking cities, and metropolitan areas with strong inter-city integration."
    ),
}


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def patch_docx_media(docx_path: Path) -> None:
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    replacements = {
        "word/media/image2.png": SOURCE_FIG2,
        "word/media/image4.png": SOURCE_FIG4,
    }
    tmp = docx_path.with_suffix(".tmp.docx")

    with ZipFile(docx_path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        rels = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
        relmap = {}
        rid_by_media = {}
        for rel in rels:
            target = rel.get("Target")
            relmap[rel.get("Id")] = target
            if target and target.startswith("media/"):
                rid_by_media[f"word/{target}"] = rel.get("Id")

        doc_xml = etree.fromstring(zin.read("word/document.xml"))
        alt_text = {
            "media/image1.png": "Research framework and mechanism analysis diagram.",
            "media/image2.png": "Distribution of SCCD in the 2006-2024 main sample.",
            "media/image3.png": "Spatial distribution of lnCE and SCCD across Chinese cities from 2006 to 2024.",
            "media/image4.png": "Nonlinear relationship between SCCD and lnCE.",
            "media/image5.png": "Conditional SCCD-lnCE curves at low and high digital economy development.",
        }
        for blip in doc_xml.xpath(".//a:blip", namespaces=ns):
            rid = blip.get(f"{{{ns['r']}}}embed")
            target = relmap.get(rid)
            inline = blip
            while inline is not None and inline.tag != f"{{{ns['wp']}}}inline":
                inline = inline.getparent()
            if inline is None:
                continue
            doc_pr = inline.find("wp:docPr", namespaces=ns)
            if doc_pr is not None and target in alt_text:
                doc_pr.set("title", target.replace("media/", "").replace(".png", ""))
                doc_pr.set("descr", alt_text[target])

        for table in doc_xml.xpath(".//w:tbl", namespaces=ns):
            first_row = table.find("w:tr", namespaces=ns)
            if first_row is None:
                continue
            tr_pr = first_row.find("w:trPr", namespaces=ns)
            if tr_pr is None:
                tr_pr = etree.Element(f"{{{ns['w']}}}trPr")
                first_row.insert(0, tr_pr)
            if tr_pr.find("w:tblHeader", namespaces=ns) is None:
                tr_pr.append(etree.Element(f"{{{ns['w']}}}tblHeader"))

        for media_name, source in replacements.items():
            rid = rid_by_media[media_name]
            width_px, height_px = Image.open(source).size
            for blip in doc_xml.xpath(f'.//a:blip[@r:embed="{rid}"]', namespaces=ns):
                inline = blip
                while inline is not None and inline.tag != f"{{{ns['wp']}}}inline":
                    inline = inline.getparent()
                if inline is None:
                    continue
                extent = inline.find("wp:extent", namespaces=ns)
                if extent is not None:
                    cx = int(extent.get("cx"))
                    cy = round(cx * height_px / width_px)
                    extent.set("cy", str(cy))
                for aext in inline.xpath(".//a:xfrm/a:ext", namespaces=ns):
                    cx = int(aext.get("cx"))
                    cy = round(cx * height_px / width_px)
                    aext.set("cy", str(cy))

        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = etree.tostring(doc_xml, xml_declaration=True, encoding="UTF-8", standalone="yes")
            if item.filename in replacements:
                data = replacements[item.filename].read_bytes()
            zout.writestr(item, data)

    tmp.replace(docx_path)


def copy_and_zip_figures() -> None:
    targets = {
        "Figure_2_sccd_distribution.png": SOURCE_FIG2,
        "Figure_4_nonlinear_sccd_lnce.png": SOURCE_FIG4,
    }
    for name, source in targets.items():
        shutil.copy2(source, FIGURES_DIR / name)

    zip_path = PACKAGE / "figures.zip"
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as zf:
        for fig in sorted(FIGURES_DIR.glob("*.png")):
            zf.write(fig, arcname=f"figures/{fig.name}")


def reorganize_package() -> None:
    if REORG_DIR.exists():
        shutil.rmtree(REORG_DIR)
    (REORG_DIR / "01_main_submission").mkdir(parents=True)
    (REORG_DIR / "02_supplementary").mkdir(parents=True)
    (REORG_DIR / "03_figures").mkdir(parents=True)
    (REORG_DIR / "04_archives").mkdir(parents=True)

    main_files = [
        "manuscript.docx",
        "title_page_and_declarations.docx",
        "declarationStatement.docx",
        "cover_letter.docx",
        "highlights.docx",
    ]
    for name in main_files:
        src = PACKAGE / name
        if src.exists():
            shutil.copy2(src, REORG_DIR / "01_main_submission" / name)

    for name in ["appendix_supplement.docx"]:
        src = PACKAGE / name
        if src.exists():
            shutil.copy2(src, REORG_DIR / "02_supplementary" / name)

    for fig in sorted(FIGURES_DIR.glob("*.png")):
        shutil.copy2(fig, REORG_DIR / "03_figures" / fig.name)

    shutil.copy2(PACKAGE / "figures.zip", REORG_DIR / "04_archives" / "figures.zip")

    manifest = [
        "# SCS Submission Package Manifest",
        "",
        "Generated: 2026-05-24",
        "",
        "## 01_main_submission",
        "- manuscript.docx",
        "- title_page_and_declarations.docx",
        "- declarationStatement.docx",
        "- cover_letter.docx",
        "- highlights.docx",
        "",
        "## 02_supplementary",
        "- appendix_supplement.docx",
        "",
        "## 03_figures",
    ]
    for fig in sorted((REORG_DIR / "03_figures").glob("*.png")):
        manifest.append(f"- {fig.name}")
    manifest.extend(
        [
            "",
            "## 04_archives",
            "- figures.zip",
            "",
            "Notes:",
            "- Figure 2 and Figure 4 were regenerated from Stata code with in-image titles removed.",
            "- manuscript.docx was revised in place in the source submission folder and copied here.",
        ]
    )
    (REORG_DIR / "MANIFEST.md").write_text("\n".join(manifest) + "\n", encoding="utf-8")

    reorg_zip = ROOT / "outputs" / "submission_package_scs_2026-05-24_reorganized.zip"
    if reorg_zip.exists():
        reorg_zip.unlink()
    with ZipFile(reorg_zip, "w", ZIP_DEFLATED) as zf:
        for file in sorted(REORG_DIR.rglob("*")):
            if file.is_file():
                zf.write(file, arcname=str(file.relative_to(REORG_DIR.parent)))


def main() -> None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = MANUSCRIPT.with_name(f"manuscript_{timestamp}.bak.docx")
    shutil.copy2(MANUSCRIPT, backup)

    doc = Document(MANUSCRIPT)
    for idx, text in SECTION_REVISIONS.items():
        replace_paragraph_text(doc.paragraphs[idx], text)
    doc.save(MANUSCRIPT)
    patch_docx_media(MANUSCRIPT)
    copy_and_zip_figures()
    reorganize_package()
    print(f"Revised manuscript in place: {MANUSCRIPT}")
    print(f"Backup: {backup}")
    print(f"Reorganized package: {REORG_DIR}")


if __name__ == "__main__":
    main()
