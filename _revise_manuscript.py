from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(r"D:\Workplace\SCS")
IN = ROOT / "SCS0514_revised_tables_ER_IV.docx"
OUT = ROOT / "SCS0514_revised_tables_ER_IV_2006_2024_updated.docx"
FIG = ROOT / "outputs" / "figures"

doc = Document(str(IN))


def clear_para(p):
    pPr = p._p.pPr
    for child in list(p._p):
        if pPr is not None and child is pPr:
            continue
        p._p.remove(child)


def set_para_text(p, text):
    clear_para(p)
    p.add_run(text)


def replace_paragraph_start(prefix, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            set_para_text(p, text)
            return True
    raise RuntimeError(f"Paragraph not found: {prefix}")


def paragraph_after(paragraph, text=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p


def picture_in_paragraph(p, image_path, width=Inches(6.2)):
    clear_para(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=width)


def insert_picture_after(paragraph, image_path, caption, width=Inches(6.2)):
    pic_p = paragraph_after(paragraph)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(image_path), width=width)
    cap_p = paragraph_after(pic_p, caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap_p


def remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)


def remove_caption_and_previous_drawing(caption_text):
    for p in list(doc.paragraphs):
        if p.text.strip() == caption_text:
            prev = p._p.getprevious()
            if prev is not None and prev.xpath(".//w:drawing"):
                prev.getparent().remove(prev)
            remove_paragraph(p)
            return


def set_cell(cell, text):
    cell.text = str(text)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(8)


def fill_table(table, rows):
    if len(table.rows) != len(rows):
        raise ValueError(f"Row count mismatch: table has {len(table.rows)}, data has {len(rows)}")
    for row, vals in zip(table.rows, rows):
        if len(row.cells) != len(vals):
            raise ValueError(f"Column count mismatch: row has {len(row.cells)}, data has {len(vals)}")
        for cell, val in zip(row.cells, vals):
            set_cell(cell, val)


paragraph_updates = [
    ("This study investigates whether and how digital-enabled urban multidimensional spatial coordination affects carbon emissions in China.",
     "This study investigates whether and how digital-enabled urban multidimensional spatial coordination affects carbon emissions in China. Building on the conventional production-living-ecological framework, the analysis incorporates digital space as a fourth functional dimension and constructs a spatial coupling coordination degree (SCCD) index to capture the joint evolution of production, living, ecological, and digital functions. Using panel data for 284 prefecture-level and above cities from 2006 to 2024, including fitted/extrapolated values for 2022-2024, the study estimates two-way fixed-effects models and then examines mediation and moderation mechanisms. The results reveal a significant inverted U-shaped relationship between SCCD and carbon emissions, with a full-sample turning point of 0.542. For cities below this threshold, stronger coordination remains associated with expansion-led emission growth; once the threshold is crossed, efficiency gains from factor reallocation, digital-physical integration, and technological spillovers begin to dominate. Industrial upgrading and green technological innovation operate as important transmission channels, although their short-run carbon-reduction effects are partly offset by adjustment costs and rebound effects. Digital economy development significantly moderates the nonlinear relationship, while the full-sample evidence for environmental regulation and polycentric urban structure as moderators is weaker and requires cautious interpretation. These findings suggest that low-carbon urban transition depends not only on improving spatial coordination itself, but also on raising the quality, maturity, and governance capacity of digital-spatial integration."),
    ("Accordingly, this study investigates how urban multidimensional spatial coupling coordination affects carbon emissions",
     "Accordingly, this study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation. Extending the conventional production-living-ecological framework, it treats digital space as a fourth functional dimension and constructs a spatial coupling coordination degree (SCCD) index covering production, living, ecological, and digital functions. Using panel data for 284 prefecture-level and above cities during 2006-2024, with 2022-2024 values treated as fitted/extrapolated observations, the study estimates two-way fixed-effects models and then examines mediation and moderation mechanisms."),
    ("This study uses panel data for 284 prefecture-level and above cities in China over the period 2006-2024.",
     "This study uses panel data for 284 prefecture-level and above cities in China over the period 2006-2024. The updated analysis dataset contains 5,396 city-year observations and is balanced by city and year. Of these observations, 4,544 are original observed records for 2006-2021, while 852 observations for 2022-2024 are fitted/extrapolated values identified by the indicator is_fitted. These fitted observations are used to extend the empirical window but should not be interpreted as directly observed official data. The sample window is determined by the availability of multidimensional spatial indicators and city-level carbon-emission data. The raw data are drawn mainly from the China City Statistical Yearbook, China Energy Statistical Yearbook, China Urban Construction Statistical Yearbook, China Rural Statistical Yearbook, China Industrial Statistical Yearbook, and China Environmental Statistical Yearbook. City-level carbon emissions are taken from the Emissions Database for Global Atmospheric Research (EDGAR), which ensures temporal comparability under a consistent accounting framework. The exact fitting and extrapolation procedure for 2022-2024 should be documented by the authors before submission [NEEDS AUTHOR CONFIRMATION]."),
    ("The dependent variable is carbon emissions (CE), measured using city-level estimates from the EDGAR database.",
     "The dependent variable is carbon emissions (CE), measured using city-level estimates from the EDGAR database for observed years and extended through the fitted/extrapolated 2022-2024 observations in the analysis dataset. Following the empirical convention in carbon-emissions studies, the variable is used in logarithmic form in the benchmark regressions to reduce heteroskedasticity and improve comparability across cities with different economic scales. The fitted observations are retained in the full-sample estimates and are separately checked against the observed-only 2006-2021 sample."),
    ("Table 2 reports the descriptive statistics for all variables.",
     "Table 2 reports the descriptive statistics for the full 2006-2024 sample. The mean value of carbon emissions is 4.612, whereas the mean SCCD is 0.318, suggesting that most sample cities remain at a moderate level of multidimensional spatial coordination. The data also show substantial cross-city variation in carbon emissions, environmental regulation intensity, digital economy development, green technological innovation, and polycentric urban structure, supporting the analysis of nonlinear, mediating, and moderating effects. Because 2022-2024 values are fitted/extrapolated, these descriptive statistics should be read as full-sample estimates rather than purely observed official statistics."),
    ("Fig. 2 further visualizes the spatiotemporal distributions",
     "Fig. 2 shows the full-sample SCCD distribution, while Fig. 3 reports the annual mean trends of carbon emissions and SCCD over 2006-2024. These figures provide descriptive context for the nonlinear estimates and show that the full sample includes fitted/extrapolated observations for 2022-2024. Updated map-based spatiotemporal figures were not generated in the current workflow [NEEDS AUTHOR CONFIRMATION]."),
    ("Table 3 reports the baseline estimates.",
     "Table 3 reports the baseline estimates for the full 2006-2024 sample. Across all specifications, the coefficient on the linear SCCD term is significantly positive and the coefficient on the squared term is significantly negative, providing strong evidence of an inverted U-shaped relationship between urban multidimensional spatial coupling coordination and carbon emissions. In the full-control specification, the coefficient of SCCD is 8.378 and the coefficient of SCCD2 is -7.725, both significant at the 1% level. The result indicates that the environmental effect of spatial coordination is not monotonic. In the initial stage, improvements in coordination are still accompanied by infrastructure expansion, construction activity, factor agglomeration, and rising energy demand, so carbon emissions increase with SCCD. Once coordination reaches a sufficiently mature stage, however, the efficiency effect begins to dominate: digital-physical integration improves factor matching, lowers information frictions, strengthens knowledge spillovers, and supports cleaner urban operation. This pattern is consistent with the broader literature arguing that the environmental consequences of urban restructuring depend on whether expansion effects or efficiency gains are dominant in a given development stage."),
    ("The estimated turning point is 0.522,",
     "The estimated full-sample turning point is 0.542, calculated as -8.378 / [2 x (-7.725)]. This value lies within the updated SCCD range of 0.121 to 0.822. The fixed-effects estimates therefore imply an inverted U-shaped relationship: when SCCD is below this threshold, stronger coordination is associated with higher carbon emissions, whereas beyond the threshold the marginal effect turns negative. Fig. 4 provides a descriptive bivariate visualization of the CE-SCCD relationship and highlights fitted 2022-2024 observations, but the turning-point interpretation is based on the two-way fixed-effects estimates in Table 3. The observed-only sensitivity estimate for 2006-2021 gives a turning point of 0.522, indicating that the addition of fitted 2022-2024 observations changes the numerical threshold but not the main conclusion. At the left side of the curve, the expansion effect of spatial integration - including infrastructure investment, industrial concentration, and additional energy demand induced by functional reorganization - still dominates. Once SCCD moves beyond the threshold, gains in factor allocation, coordination efficiency, and technology spillovers outweigh the expansion effect, and further improvements in SCCD reduce emissions."),
    ("Table 4 reports the robustness, endogeneity, and first-stage results.",
     "Table 4 reports the robustness, endogeneity, and first-stage results for the full 2006-2024 sample. To verify that the baseline results are not driven by outliers, measurement choices, or reverse causality, several additional tests are conducted. First, bilateral winsorization at the 1% level is used to reduce the influence of extreme values. Second, the dependent variable is replaced with average nighttime light intensity as an alternative proxy. Third, a revised coupling-coordination model is used to reconstruct the explanatory variable. Fourth, IV-2SLS estimation is performed using the existing IV variable as the lagged SCCD instrument and its squared term for the nonlinear specification. Across these checks, the linear and quadratic terms remain significantly positive and negative, respectively. This consistency suggests that the estimated nonlinear relationship is not an artifact of a particular measurement scheme or a few extreme observations, but rather reflects a relatively stable empirical pattern."),
    ("To further address endogeneity, the lagged linear and squared terms of SCCD are used",
     "To further address endogeneity, the lagged linear and squared terms of SCCD are used as instrumental variables in a two-stage least squares estimation. The second-stage results show that the coefficient of SCCD remains significantly positive, while the coefficient of SCCD2 remains significantly negative. Based on the full-sample IV-2SLS estimates, the implied turning point is approximately 0.638, which remains within the observed SCCD range. The first-stage results indicate instrument relevance: IV significantly predicts SCCD, and IV2 significantly predicts SCCD2, with within R-squared values of 0.652 and 0.586, respectively. Taken together, these results reinforce the conclusion that multidimensional spatial coupling coordination has a robust inverted U-shaped association with carbon emissions after accounting for potential reverse causality."),
    ("Table 5 reports the regional heterogeneity results.",
     "Table 5 reports the regional heterogeneity results. Given substantial regional differences in development patterns, industrial structure, and urbanization, this study further examines whether the carbon effect of SCCD differs across eastern, central, and western China. The estimates show that the inverted U-shaped relationship remains statistically significant in all three regions, suggesting that the nonlinear effect is broadly robust across China. The magnitude of the effect, however, differs across regions. The absolute value of the quadratic coefficient is largest in central China (-11.725), followed by western China (-7.151) and eastern China (-6.922), indicating that carbon emissions in central China are most sensitive to changes in multidimensional spatial coordination. In eastern China, the estimated effect is flatter, likely because industrial upgrading, digital development, and efficiency gains already play a larger role than large-scale physical expansion. In western China, the nonlinear effect remains pronounced, which may reflect the coexistence of infrastructure investment, industrial catch-up, and uneven digital development. Fig. 5 and Fig. 6 provide additional regional descriptive evidence. These regional differences imply that the same improvement in SCCD can produce very different carbon consequences depending on the local stage of urban transformation, industrial foundation, and governance capacity."),
    ("Table 6 reports the mediation analysis results.",
     "Table 6 reports the mediation analysis results. To uncover the mechanisms linking SCCD to carbon emissions, this study tests the mediating roles of industrial upgrading and green technological innovation. The results show that SCCD has an inverted U-shaped effect on both mechanism variables, suggesting that stronger spatial coordination initially promotes industrial upgrading and green technological innovation, but that these gains weaken as coordination becomes more mature. The turning points implied by the mediator equations are above the maximum SCCD value in the updated sample, so the mediator-equation curvature should be interpreted cautiously. The current workflow provides mechanism-equation evidence rather than separately estimated indirect-effect magnitudes [NEEDS AUTHOR CONFIRMATION]."),
    ("When industrial upgrading is added to the carbon-emissions equation,",
     "When industrial upgrading is added to the carbon-emissions equation, its coefficient is significantly positive (1.491), while the coefficients on SCCD and its squared term remain statistically significant. This indicates that industrial upgrading is an important transmission channel, although its short-run carbon effect is emission-increasing before cleaner production systems are fully established. From a transitional perspective, this result is understandable: upgrading often involves new fixed-asset investment, equipment replacement, construction of modern industrial facilities, and the coexistence of old and new sectors. In other words, multidimensional spatial coordination can promote structural change, but the carbon benefits of that change may only materialize after the adjustment phase has passed. The result therefore supports H2 as mechanism evidence, while also suggesting that the quality and timing of industrial upgrading are as important as the direction of upgrading itself."),
    ("When green technological innovation is added to the carbon-emissions equation,",
     "When green technological innovation is added to the carbon-emissions equation, its coefficient is also significantly positive (0.325), and the SCCD coefficients remain statistically significant. This confirms that green technological innovation acts as another important transmission channel. Yet the positive coefficient of the mediator suggests that green innovation does not automatically lower observed emissions in the short run. Research and development, testing, commercialization, and adoption all require capital and energy inputs, while efficiency gains may be partly offset by rebound effects. Thus, the mediating role of green innovation should be understood dynamically: multidimensional spatial coordination improves the urban conditions for green innovation, but the carbon-reduction payoff of that innovation is not instantaneous. This interpretation is consistent with the view that technological transition is a gradual process that depends on diffusion speed, market maturity, and institutional support."),
    ("The interaction terms for environmental regulation show that regulation flattens",
     "The full-sample interaction terms for environmental regulation do not provide statistically significant evidence that regulation moderates the inverted U-shaped relationship. Although the coefficient on SCCD2 x ER is positive, it is not statistically significant in the generated full-sample table. Therefore, H4a is not supported by the current full-sample estimates. The observed-only sensitivity result is stronger, but the manuscript should not present environmental regulation as a robust full-sample moderator without additional confirmation [NEEDS AUTHOR CONFIRMATION]."),
    ("The interaction terms for digital economy development are statistically significant",
     "The interaction terms for digital economy development are statistically significant and support H4b. The coefficient on SCCD x DEI is significantly negative, while the coefficient on SCCD2 x DEI is significantly positive. This pattern indicates that digital economy development reshapes the nonlinear relationship between SCCD and carbon emissions. It is consistent with the argument that digital economy development lowers coordination frictions, improves information matching, facilitates intelligent governance, and enhances the efficiency of resource allocation. Under stronger digital conditions, the same increase in SCCD is more likely to translate into operational efficiency, smart transport management, cleaner energy use, and faster diffusion of green technologies, thereby attenuating the emissions peak. Conditional turning-point plots were not generated in the current workflow [NEEDS AUTHOR CONFIRMATION]."),
    ("For polycentric urban structure, the interaction results support H5.",
     "For polycentric urban structure, the generated full-sample interaction terms do not statistically support H5. The coefficients on SCCD x POLY and SCCD2 x POLY are not significant, so the current evidence is insufficient to conclude that polycentricity shifts the turning point to the right. Polycentricity may still matter theoretically because dispersed centers can create longer logistics chains, weaker commuting-job matching, and duplicated infrastructure when cross-center links remain incomplete, but this argument should be presented as a possible explanation requiring further empirical confirmation rather than as a confirmed moderation result [NEEDS AUTHOR CONFIRMATION]."),
    ("Taken together, the moderation results show that urban context matters.",
     "Taken together, the moderation results show that urban context matters, but the strength of the evidence differs across moderators. Digital economy development significantly conditions the SCCD-carbon relationship in the full-sample estimates. By contrast, the full-sample evidence for environmental regulation and polycentricity is not statistically robust in the generated moderation table. This finding complements the baseline and mediation results by demonstrating that digital capability is a clearer boundary condition in the updated full sample, while the roles of regulation and spatial form require more cautious interpretation."),
    ("This study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation.",
     "This study investigates how urban multidimensional spatial coupling coordination affects carbon emissions in China under digital transformation. By extending the conventional production-living-ecological framework to include digital space and by using panel data for 284 prefecture-level and above cities over 2006-2024, the paper evaluates both the nonlinear carbon effect of spatial coordination and the mechanisms through which this effect operates. The dataset contains observed records for 2006-2021 and fitted/extrapolated observations for 2022-2024, so the extended-period findings should be interpreted with this data construction in mind. The central message is that multidimensional spatial coordination is neither inherently low-carbon nor inherently carbon-intensive. Its environmental consequences depend on development stage, structural adjustment, and governance capacity. This perspective shifts the discussion away from a simple more-coordination-is-better logic and toward a more conditional understanding of how spatial integration, digitalization, and carbon transition interact over time."),
    ("Four main conclusions can be drawn.",
     "Four main conclusions can be drawn. First, SCCD exhibits a significant inverted U-shaped relationship with carbon emissions, with a full-sample turning point of 0.542. This indicates that the carbon consequences of spatial coordination are strongly stage-dependent: at lower levels of coordination, integration is still accompanied by construction-led expansion, higher energy demand, and agglomeration pressure, whereas at higher levels, deeper coordination across production, living, ecological, and digital spaces improves factor allocation, operational efficiency, and technological spillovers, thereby reducing emissions. Second, the nonlinear relationship remains robust after multiple specification adjustments and IV-2SLS estimation, suggesting that the main finding is not driven by extreme observations or simple reverse causality. Third, industrial upgrading and green technological innovation constitute important transmission channels, but their short-run effects are not purely emission-reducing because transition costs, equipment renewal, and rebound effects partly offset efficiency gains. Fourth, digital economy development significantly conditions the carbon effect of SCCD in the full-sample estimates. The updated full-sample evidence does not support strong claims that environmental regulation or polycentricity significantly moderate the nonlinear relationship, so these channels should be interpreted cautiously."),
    ("Overall, the evidence shows that multidimensional spatial coordination does not automatically produce low-carbon outcomes.",
     "Overall, the evidence shows that multidimensional spatial coordination does not automatically produce low-carbon outcomes. Its environmental value depends on whether spatial integration is translated into efficiency improvement rather than land expansion, whether structural and technological upgrading generate durable rather than temporary carbon gains, and whether urban governance can manage the frictions associated with digital-physical transformation. For cities below the full-sample turning point of 0.542, the immediate challenge is to prevent coordination from being translated primarily into duplicated infrastructure and energy-intensive growth. For cities approaching or surpassing the turning point, the policy task shifts toward consolidating the efficiency advantages of digital integration and converting innovation and upgrading into persistent emission reductions."),
    ("First, urban low-carbon policy should move beyond extensive physical expansion",
     "First, urban low-carbon policy should move beyond extensive physical expansion and prioritize deeper digital-physical integration. Because many sample cities remain on the left side of the full-sample turning point, simply raising formal indicators of spatial coordination is insufficient to guarantee lower emissions. Policy should therefore focus on digital infrastructure, interoperable data systems, and intelligent governance tools that reduce the coordination costs of production, living, ecological, and digital spaces."),
    ("Fourth, governance strategies should be differentiated across regions and development stages",
     "Fourth, governance strategies should be differentiated across regions and development stages. Central China deserves particular attention because it faces both stronger emission sensitivity to changes in SCCD and greater potential gains from later-stage coordination. Eastern China is better positioned to explore advanced digital-physical governance models, whereas western China should place greater emphasis on coupling ecological protection with improvements in basic public services and urban functional coordination. More broadly, cities below the turning point should prioritize curbing construction-led expansion and improving the quality rather than the quantity of coordination; cities near the turning point should accelerate digital governance reform, industrial upgrading quality, environmental regulation, and green technology diffusion; and cities already beyond the turning point should focus on locking in low-carbon gains through institutional coordination, cross-sector data integration, and long-term monitoring of rebound effects."),
    ("This study has several limitations.",
     "This study has several limitations. First, although the SCCD index incorporates digital space into the conventional production-living-ecological framework, the measurement of digital space is constrained by data availability and may not fully capture platform governance, data mobility, algorithmic coordination, or the quality of digital public services. Second, although the empirical design includes fixed effects, robustness checks, and instrumental-variable estimation, the analysis remains based on observational panel data and therefore cannot fully eliminate all identification concerns. Third, the extended 2006-2024 sample includes fitted/extrapolated observations for 2022-2024; these values improve temporal coverage but are not equivalent to directly observed official records, and their construction procedure should be fully documented before submission [NEEDS AUTHOR CONFIRMATION]. Fourth, the mechanism analysis focuses on industrial upgrading and green technological innovation, but other channels--such as energy structure, green finance, or household behavioral change--may also play meaningful roles in shaping carbon outcomes."),
]

for prefix, text in paragraph_updates:
    replace_paragraph_start(prefix, text)

for p in doc.paragraphs:
    s = p.text.strip()
    if s.startswith("[NEEDS CHECK: The authors must confirm whether the data and code"):
        set_para_text(p, "[NEEDS AUTHOR CONFIRMATION: The authors must confirm whether the updated dataset and Stata workflow can be publicly shared. If shared, the data availability statement should identify the repository or access procedure and explicitly disclose that 2022-2024 observations are fitted/extrapolated rather than directly observed official records.]")
    elif s.startswith("[NEEDS CHECK: The authors must confirm whether there are any financial"):
        set_para_text(p, "[NEEDS AUTHOR CONFIRMATION: The authors must confirm whether there are any financial or personal relationships that could have appeared to influence the work reported in this manuscript.]")
    elif s.startswith("[NEEDS CHECK: The authors must insert verified funding"):
        set_para_text(p, "[NEEDS AUTHOR CONFIRMATION: The authors must insert verified funding sources, grant numbers, and funder roles, or state that no specific funding was received if accurate.]")
    elif s.startswith("[NEEDS CHECK: Insert verified author names"):
        set_para_text(p, "[NEEDS AUTHOR CONFIRMATION: Insert verified author names or initials and CRediT roles. Do not submit until all author contributions have been confirmed.]")
    elif s.startswith("[NEEDS CHECK: The authors must confirm whether generative AI"):
        set_para_text(p, "[NEEDS AUTHOR CONFIRMATION: The authors must confirm whether generative AI or AI-assisted technologies were used. If applicable, identify the tool and purpose of use, and confirm that the authors reviewed and take full responsibility for the final content. If not applicable, replace this placeholder with the journal-appropriate no-use statement.]")

tables = []
tables.append([
["Category","Variable","Symbol","Obs.","Mean","Median","S.D.","Min","Max"],
["Dependent variable","Carbon emissions","CE","5396","4.612","4.654","1.253","0.000","8.325"],
["Core explanatory variable","Spatial coupling coordination degree","SCCD","5396","0.318","0.304","0.085","0.121","0.822"],
["Control variable","Trade openness","OPEN","5381","0.193","0.080","0.337","-0.720","3.640"],
["Control variable","Urbanization rate","UR","5396","54.498","52.606","15.419","6.491","100.000"],
["Control variable","Urban-rural income gap","URG","5381","2.437","2.342","0.549","1.207","6.378"],
["Control variable","Government intervention intensity","GI","5396","0.186","0.162","0.098","0.043","1.027"],
["Mediating variable","Industrial upgrading","OIU","5396","2.288","2.283","0.144","1.163","2.836"],
["Mediating variable","Green technological innovation","GTI","5396","4.047","4.038","1.795","0.000","9.872"],
["Moderating variable","Environmental regulation intensity","ER","5396","0.008","0.007","0.003","0.000","0.026"],
["Moderating variable","Digital economy index","DEI","5396","0.060","0.047","0.076","0.000","0.940"],
["Moderating variable","Degree of polycentricity","POLY","5396","0.346","0.374","0.189","0.000","0.963"],
])
tables.append([
["Variables","(1)","(2)","(3)","(4)","(5)"],["Dependent variable","CE","CE","CE","CE","CE"],
["SCCD","13.468***","12.344***","9.123***","8.871***","8.378***"],["","(1.546)","(1.521)","(1.374)","(1.408)","(1.388)"],
["SCCD2","-11.472***","-10.369***","-8.283***","-7.882***","-7.725***"],["","(1.814)","(1.816)","(1.622)","(1.658)","(1.665)"],
["OPEN","","0.486***","0.204**","0.223***","0.259***"],["","","(0.107)","(0.084)","(0.085)","(0.084)"],
["UR","","","0.032***","0.030***","0.028***"],["","","","(0.004)","(0.004)","(0.004)"],
["URG","","","","-0.135*","-0.112"],["","","","","(0.073)","(0.070)"],
["GI","","","","","-2.064***"],["","","","","","(0.626)"],
["Constant","1.057***","1.171***","0.438","0.944**","1.400***"],["","(0.311)","(0.299)","(0.284)","(0.377)","(0.412)"],
["N","5396","5381","5381","5366","5366"],["Within R²","0.568","0.579","0.629","0.631","0.641"],
])
tables.append([
["Variables","(1)","(2)","(3)","(4)"],["Specification","Winsorized at 1%","Alternative dependent variable","Alternative explanatory variable","IV-2SLS / second stage"],
["Core linear term","9.888***","3.536***","8.378***","18.658***"],["","(1.414)","(0.757)","(1.388)","(2.694)"],
["Core squared term","-9.886***","-2.515***","-7.725***","-14.618***"],["","(1.698)","(0.789)","(1.665)","(2.711)"],
["Controls","Yes","Yes","Yes","Yes"],["City fixed effects","Yes","Yes","Yes","Yes"],["Year fixed effects","Yes","Yes","Yes","Yes"],
["N","5366","5366","5366","5084"],["Within R²","0.647","0.573","0.641","0.602"],
["First-stage instrument relevance","","","","IV-2SLS first stage"],["Dependent variable: SCCD","","","",""],
["IV","","","","0.344*** / (0.064)"],["IV2","","","","-0.071 / (0.079)"],["First-stage within R² for SCCD","","","","0.652"],
["Dependent variable: SCCD2","","","",""],["IV","","","","-0.035 / (0.051)"],["IV2","","","","0.382*** / (0.070)"],
["First-stage within R² for SCCD2","","","","0.586"],["First-stage N","","","","5084"],
])
tables.append([
["Variables","(1)","(2)","(3)"],["Region","Eastern China","Central China","Western China"],
["SCCD","7.589***","12.171***","7.221***"],["","(2.677)","(2.747)","(1.831)"],
["SCCD2","-6.922***","-11.725***","-7.151***"],["","(2.582)","(3.489)","(2.180)"],
["Constant","1.278*","0.523","1.946***"],["","(0.740)","(0.763)","(0.675)"],
["Controls","Yes","Yes","Yes"],["City fixed effects","Yes","Yes","Yes"],["Year fixed effects","Yes","Yes","Yes"],
["N","1828","1885","1653"],["Within R²","0.720","0.629","0.621"],
])
tables.append([
["Variables","(1)","(2)","(3)","(4)"],["Dependent variable","Industrial upgrading","Carbon emissions","Green technological innovation","Carbon emissions"],
["SCCD","0.626***","7.445***","10.481***","4.971***"],["","(0.163)","(1.297)","(1.386)","(1.156)"],
["SCCD2","-0.371*","-7.172***","-5.438***","-5.957***"],["","(0.213)","(1.486)","(1.416)","(1.418)"],
["OIU","","1.491***","",""],["","","(0.337)","",""],["GTI","","","","0.325***"],["","","","","(0.030)"],
["Constant","1.811***","-1.300*","-0.817*","1.666***"],["","(0.043)","(0.738)","(0.423)","(0.355)"],
["Controls","Yes","Yes","Yes","Yes"],["City fixed effects","Yes","Yes","Yes","Yes"],["Year fixed effects","Yes","Yes","Yes","Yes"],
["N","5366","5366","5366","5366"],["Within R²","0.614","0.651","0.812","0.683"],
])
tables.append([
["Variables","(1)","(2)","(3)"],["Moderator","Environmental regulation intensity","Digital economy index","Polycentricity"],
["SCCD","9.569***","8.829***","7.770***"],["","(1.995)","(1.289)","(2.188)"],
["SCCD2","-9.753***","-7.714***","-8.223***"],["","(2.463)","(1.473)","(2.661)"],
["ER","33.864","",""],["","(45.908)","",""],["SCCD x ER","-158.794","",""],["","(230.885)","",""],
["SCCD2 x ER","260.988","",""],["","(277.816)","",""],["DEI","","4.601**",""],["","","(2.080)",""],
["SCCD x DEI","","-22.377***",""],["","","(7.346)",""],["SCCD2 x DEI","","21.644***",""],["","","(6.145)",""],
["POLY","","","-1.179"],["","","","(0.970)"],["SCCD x POLY","","","3.213"],["","","","(4.848)"],
["SCCD2 x POLY","","","-0.186"],["","","","(5.505)"],["Constant","1.157**","1.267***","1.699***"],["","(0.498)","(0.404)","(0.514)"],
["Controls","Yes","Yes","Yes"],["City fixed effects","Yes","Yes","Yes"],["Year fixed effects","Yes","Yes","Yes"],["N","5366","5366","5366"],["Within R²","0.642","0.643","0.644"],
])

for i, rows in enumerate(tables, start=1):
    fill_table(doc.tables[i], rows)

# Replace and add figures.
picture_in_paragraph(doc.paragraphs[62], FIG / "fig_sccd_distribution_full_2006_2024.png")
set_para_text(doc.paragraphs[63], "Fig. 2. Distribution of SCCD in the full 2006-2024 sample.")
doc.paragraphs[63].alignment = WD_ALIGN_PARAGRAPH.CENTER
insert_picture_after(doc.paragraphs[63], FIG / "fig_yearly_mean_ce_sccd_full_2006_2024.png", "Fig. 3. Annual mean carbon emissions and SCCD, 2006-2024.")
for p in list(doc.paragraphs):
    if "base-map source and approval number" in p.text:
        set_para_text(p, "Note: CE denotes carbon emissions and SCCD denotes the spatial coupling coordination degree. The 2022-2024 values shown in the trend figure are fitted/extrapolated observations, not directly observed official records.")
        break

for p in doc.paragraphs:
    if p.text.strip() == "Fig. 3. Nonlinear effect of SCCD on carbon emissions.":
        prev = p._p.getprevious()
        if prev is not None:
            picture_in_paragraph(Paragraph(prev, p._parent), FIG / "fig_ce_sccd_quadratic_fit_full_2006_2024.png")
        set_para_text(p, "Fig. 4. Nonlinear relationship between SCCD and carbon emissions in the full 2006-2024 sample.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        break

for p in doc.paragraphs:
    if p.text.strip().startswith("The heterogeneity findings also imply"):
        c5 = insert_picture_after(p, FIG / "fig_ce_by_region_box_full_2006_2024.png", "Fig. 5. Carbon emissions by region in the full 2006-2024 sample.", width=Inches(6.1))
        insert_picture_after(c5, FIG / "fig_regional_mean_ce_trends_full_2006_2024.png", "Fig. 6. Annual mean carbon emissions by region, 2006-2024.", width=Inches(6.1))
        break

remove_caption_and_previous_drawing("Fig. 4. Moderating effect of digital economy development.")
remove_caption_and_previous_drawing("Fig. 5. Moderating effect of polycentric urban structure.")

notes = {
    "Table 2. Descriptive statistics.": "Note: The full sample covers 2006-2024. Observations for 2022-2024 are fitted/extrapolated values, not directly observed official records.",
    "Table 3. Baseline regression results.": "Note: Standard errors clustered at the city level are reported in parentheses. The dependent variable is CE. The full sample includes fitted/extrapolated observations for 2022-2024.",
    "Table 4. Robustness, endogeneity, and first-stage results.": "Note: Column (4) reports IV-2SLS second-stage estimates. IV is treated as lagged SCCD, and IV2 is its squared term. The full sample includes fitted/extrapolated observations for 2022-2024.",
    "Table 7. Moderation analysis results.": "Note: Full-sample ER and POLY interaction terms are not statistically significant in the generated estimates; DEI interactions remain statistically significant.",
}
for p in list(doc.paragraphs):
    s = p.text.strip()
    if s in notes:
        paragraph_after(p, notes[s])

OUT.unlink(missing_ok=True)
doc.save(str(OUT))
print(OUT)
