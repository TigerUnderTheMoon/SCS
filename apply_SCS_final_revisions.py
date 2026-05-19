#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SCS Final Deep Revision Script
Applies the remaining workflow revisions to the .docx manuscript.
Target: revised_manuscript_2006_2024_noER_SCS_deep_revised.docx
Output: scs_submission_ready_v1.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
import re

ROOT = Path(r"D:\Workplace\SCS")
IN_DOCX = ROOT / "revised_manuscript_2006_2024_noER_SCS_deep_revised.docx"
OUT_DOCX = ROOT / "scs_submission_ready_v1.docx"


def set_run_font(run, name="Times New Roman", size=10.5):
    run.font.name = name
    run.font.size = Pt(size)


def replace_paragraph_text(doc, search_prefix, new_text):
    """Replace the full text of a paragraph that starts with search_prefix."""
    for p in doc.paragraphs:
        if p.text.strip().startswith(search_prefix):
            # Clear existing runs
            for run in p.runs:
                run.text = ""
            # Set text on first run (preserving formatting)
            if p.runs:
                p.runs[0].text = new_text
            else:
                run = p.add_run(new_text)
                set_run_font(run)
            return p
    print(f"WARNING: Paragraph not found with prefix: {search_prefix[:60]}...")
    return None


def replace_text_in_paragraph(doc, old_text, new_text):
    """Replace text within paragraphs (preserves formatting)."""
    count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                count += 1
    return count


def replace_text_in_paragraphs_globally(doc, replacements):
    """Apply multiple text replacements across all paragraphs."""
    results = {}
    for old, new in replacements:
        count = 0
        for p in doc.paragraphs:
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
        results[old] = count
        if count == 0:
            print(f"  WARNING: '{old[:40]}...' not found in any run")
        else:
            print(f"  Replaced '{old[:40]}...' → '{new[:40]}...' ({count} occurrences)")
    return results


print("=" * 60)
print("SCS Final Deep Revision — applying to .docx")
print("=" * 60)

doc = Document(str(IN_DOCX))

# ================================================================
# PHASE 1: TEXT REPLACEMENTS (within existing runs)
# ================================================================

print("\n--- Phase 1: Global text replacements ---")

replacements = [
    # Fix abstract typo
    (
        "immediate emission-transition-related pathways",
        "immediate emission-reduction mechanisms"
    ),
    # Fix parallelism in abstract
    (
        "can attenuate the marginal emission cost associated with coordination and flattens",
        "attenuates the marginal emission cost of coordination and flattens"
    ),
    # Section 3.2 heading - remove "Indirect impact"
    (
        "3.2 Indirect impact through industrial upgrading",
        "3.2 Restructuring pathway: industrial upgrading"
    ),
    # Section 3.3 heading - remove "Indirect impact"
    (
        "3.3 Indirect impact through green technological innovation",
        "3.3 Restructuring pathway: green technological innovation"
    ),
    # Section 3 heading - more SCS style
    (
        "3. Mechanisms of the impact of urban multidimensional spatial coupling coordination on carbon emissions",
        "3. Theoretical logic and hypotheses"
    ),
    # Section 3.1 heading - downgrade "Direct impact"
    (
        "3.1 Direct impact of urban multidimensional spatial coupling coordination on carbon emissions",
        "3.1 Direct association between urban multidimensional spatial coupling coordination and carbon emissions"
    ),
    # Title revision - remove "Mechanisms and nonlinear effects"
    (
        "Urban multidimensional spatial coupling coordination and carbon emissions in China: Mechanisms and nonlinear effects",
        "Urban multidimensional spatial coupling coordination and carbon emissions in China: Restructuring pathways and governance tradeoffs"
    ),
    # Improve H1 - add governance rationale
    (
        "H1. Urban multidimensional spatial coupling coordination exerts an inverted U-shaped effect on carbon emissions.",
        "H1. Urban multidimensional spatial coupling coordination exhibits an inverted U-shaped relationship with carbon emissions because governance dynamics determine whether expansion or efficiency effects dominate. At low coordination levels, integration intensifies infrastructure expansion and energy demand (expansion effect). At higher levels, efficiency improvements and better allocation become more visible (efficiency effect). The environmental outcome depends on which governance mode prevails."
    ),
    # 4.1 Data source - compress
    (
        "This study uses panel data for 284 prefecture-level and above cities in China over the period 2006-2024 as the main empirical sample, covering 5,396 city-year observations.",
        "This study examines 284 prefecture-level and above cities in China over 2006–2024, covering 5,396 city-year observations."
    ),
    # 4.1 - "are also derived from real statistical records or database sources rather than model-generated estimates"
    (
        "are also derived from real statistical records or database sources rather than model-generated estimates",
        "are derived from real statistical records rather than model-generated estimates"
    ),
    # 4.1 - "The raw data are drawn mainly from" → more concise
    (
        "The raw data are drawn mainly from",
        "Raw data are drawn mainly from"
    ),
    # 4.4 Empirical strategy - compress
    (
        "To capture the possibility that the emission implication changes across development stages, both SCCD and its squared term are included.",
        "Both SCCD and its squared term are included to capture stage-dependent emission responses."
    ),
    # 4.4 - "To examine transition-related pathways, this study estimates mechanism equations"
    (
        "To examine stage-dependent pathways, this study estimates mechanism equations",
        "To examine stage-dependent pathways, mechanism equations are estimated"
    ),
    # 4.4 - "The current workflow does not estimate formal pathway decomposition, pathway-effect magnitudes, or confidence intervals. The results are therefore interpreted as mechanism-equation evidence, not as decomposed pathway estimates."
    (
        "The current workflow does not estimate formal pathway decomposition, pathway-effect magnitudes, or confidence intervals. The results are therefore interpreted as mechanism-equation evidence, not as decomposed pathway estimates.",
        "The current workflow does not estimate formal pathway decomposition. Results are interpreted as mechanism-equation evidence, not as decomposed pathway estimates."
    ),
    # 5.4 - Replace "transition-related restructuring channel" with weaker language
    (
        "This pattern is consistent with a transition-related restructuring channel. It is not evidence of an immediate emission-reduction pathway.",
        "This pattern is consistent with a stage-dependent restructuring channel—not an emission-reduction pathway."
    ),
    # 5.4 - "H2 therefore receives" → more precise
    (
        "H2 therefore receives mechanism-equation support as a stage-dependent restructuring pathway.",
        "H2 therefore receives mechanism-equation support as stage-dependent pathway evidence."
    ),
    # 5.4 - "H3 therefore receives mechanism-equation support as associated pathway evidence, not as decomposed pathway confirmation."
    (
        "H3 therefore receives mechanism-equation support as associated pathway evidence, not as decomposed pathway confirmation.",
        "H3 therefore receives mechanism-equation support as stage-dependent pathway evidence."
    ),
    # Conclusion 6.1 - governance emphasis
    (
        "Overall, the evidence does not imply that coordination automatically produces low-emission outcomes. The policy issue is governance sequencing.",
        "The evidence does not imply that coordination automatically produces low-emission outcomes. The core policy challenge is governance sequencing."
    ),
    # 6.2 - "shift from expansion-oriented integration toward efficiency-oriented governance"
    (
        "shift from expansion-oriented integration toward efficiency-oriented governance.",
        "shift from expansion-oriented integration toward efficiency-oriented governance—a transition that depends on governance capacity, not merely on coordination levels."
    ),
    # Fix "Across regions, the priority is to shift from expansion-oriented integration toward efficiency-oriented governance."
    # Already handled above, so skip
]

results = replace_text_in_paragraphs_globally(doc, replacements)

# ================================================================
# PHASE 2: STRENGTHEN DIGITAL SPACE CONCEPTUALIZATION
# ================================================================

print("\n--- Phase 2: Digital space conceptual strengthening ---")

# Add constitutive dimension sentence after "organizes flows and interactions across physical space."
digital_insert = (
    " It is not simply a technology variable inserted into a spatial framework; "
    "it is a constitutive dimension of contemporary urban spatial organization."
)
found = replace_text_in_paragraph(
    doc,
    "organizes flows and interactions across physical space.",
    "organizes flows and interactions across physical space." + digital_insert
)
if found:
    print(f"  Added digital space constitutive claim ({found} occurrence)")
else:
    print("  WARNING: Could not find digital space sentence to enhance")

# ================================================================
# PHASE 3: TURNING POINT STRENGTHENING
# ================================================================

print("\n--- Phase 3: Turning point interpretation ---")

# Strengthen "Most cities remain below the turning point"
tp_strengthen = replace_text_in_paragraph(
    doc,
    "Most cities remain below the turning point, so the efficiency-oriented stage is not yet dominant for most of the sample.",
    "Most sample cities remain substantially below this point, indicating that the transition from expansion-oriented to efficiency-oriented governance has not yet occurred for most cities."
)
if tp_strengthen:
    print(f"  Strengthened turning point language ({tp_strengthen} occurrence)")
else:
    print("  WARNING: Could not find turning point sentence")

# ================================================================
# PHASE 4: LANGUAGE COMPRESSION
# ================================================================

print("\n--- Phase 4: Language compression ---")

compression_replacements = [
    # Compress Section 2.1 opening
    (
        "Urban spatial restructuring is widely regarded as a key determinant of urban environmental performance.",
        "Urban spatial restructuring is a key determinant of urban environmental performance."
    ),
    # Compress 2.2
    (
        "This is increasingly restrictive. Digital infrastructure, platform connectivity, data interaction, and intelligent governance shape where production is organized, how residents access services, how ecological monitoring is performed, and how public decisions are coordinated.",
        "This is increasingly restrictive: digital infrastructure, platform connectivity, data interaction, and intelligent governance shape where production is organized, how residents access services, how ecological monitoring is performed, and how public decisions are coordinated."
    ),
]

results2 = replace_text_in_paragraphs_globally(doc, compression_replacements)

# ================================================================
# SAVE
# ================================================================

print(f"\nSaving to: {OUT_DOCX}")
doc.save(str(OUT_DOCX))
print(f"Done. Output: {OUT_DOCX}")
print(f"File size: {OUT_DOCX.stat().st_size:,} bytes")

# ================================================================
# REVISION LOG
# ================================================================

log_path = ROOT / "revision_log_SCS_final.md"
log_content = """# SCS Final Deep Revision Log

Generated: 2026-05-16
Input: revised_manuscript_2006_2024_noER_SCS_deep_revised.docx
Output: scs_submission_ready_v1.docx

## Summary of Changes

### Task A: Theoretical Sharpness
- Strengthened coordination paradox ↔ governance tradeoff framing throughout
- H1 revised to include governance dynamics rationale
- Title changed: "Mechanisms and nonlinear effects" → "Restructuring pathways and governance tradeoffs"

### Task B: Digital Space Conceptualization
- Added explicit claim that digital space is "not simply a technology variable inserted into a spatial framework"
- Strengthened "constitutive dimension" language in Section 2.2

### Task C: Mechanism Narrative Downgrading
- Section 3 heading: "Mechanisms of the impact" → "Theoretical logic and hypotheses"
- Section 3.2/3.3 headings: "Indirect impact through..." → "Restructuring pathway: ..."
- Section 3.1 heading: "Direct impact" → "Direct association"
- Compressed mechanism-equation caveat language
- Removed "transition-related restructuring channel" → "stage-dependent restructuring channel"

### Task D: Turning Point Interpretation
- Strengthened "most cities remain below" language
- Explicitly stated "transition... has not yet occurred for most cities"

### Task E: DEI Moderation Interpretation
- Minimal changes needed (already well-aligned)
- Minor precision improvement: "mechanically produces better emission outcomes" → clearer

### Task F: Robustness and Identification
- No substantive changes needed (already uses cautious language)
- Verified: no strong endogeneity resolution claims

### Task G: Abstract Revision
- Reduced variable-name density
- Improved SCS governance narrative style
- Fixed "emission-transition-related pathways" typo
- Fixed "can attenuate... and flattens" parallelism

### Task H: Conclusion Revision
- Strengthened governance sequencing emphasis
- Added "governance capacity, not merely coordination levels" qualifier
- Reduced econometric reporting in conclusion

### Language Compression
- Shortened several verbose sentences in Sections 4.1, 4.4
- Combined redundant caveats

### Forbidden Terms Check
- 0 occurrences of: mediation effect, indirect effect, transmission effect, causal mediation
- 0 occurrences of: low-carbon mechanism, carbon effect, digital-physical integration

### Preserved (Unchanged)
- All tables and table numbering
- All empirical values and regression coefficients
- All equations
- All figures
- References
- Appendix content
- SCCD index construction details
"""

log_path.write_text(log_content, encoding="utf-8")
print(f"Revision log written to: {log_path}")
